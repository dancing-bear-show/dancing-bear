"""Command execution resume CLI tests: extract, summarize, align, render, etc."""

from __future__ import annotations

import json
import os
import sys
import tempfile
import unittest

from tests.fixtures import run, temp_yaml_file


def _run_experience_overlay_render(case, tmpdir, profile, template):
    """Render with a profile experience.yaml overlay and assert it replaces the base role.

    Shared by TestResumeCLIRender.test_render_experience_overlay_replaces_base_role
    and the permanently-skipped
    TestResumeCLIOnePageTemplateE2E.test_onepage_render_overlays_experience_from_profile_config,
    which reproduces the same overlay scenario against the (buggy) onepage
    template path documented in ONEPAGE_TEMPLATE_BUG_REASON.
    """
    data_path = os.path.join(tmpdir, "data.json")
    out_path = os.path.join(tmpdir, "out.docx")
    payload = {
        "name": "Overlay Test",
        "experience": [
            {
                "title": "Old Role", "company": "OldCo", "start": "2020",
                "end": "2021", "location": "Remote", "bullets": ["Old bullet"],
            }
        ],
    }
    with open(data_path, "w", encoding="utf-8") as f:
        json.dump(payload, f)

    cfg_path = os.path.join("config", f"experience.{profile}.yaml")
    os.makedirs("config", exist_ok=True)
    with open(cfg_path, "w", encoding="utf-8") as f:
        f.write(
            "experience:\n"
            "  - title: New Role\n"
            "    company: NewCo\n"
            "    start: 2022\n"
            "    end: 2023\n"
            "    location: Remote\n"
            "    bullets:\n"
            "      - New bullet A\n"
            "      - New bullet B\n"
        )
    try:
        proc = run([
            sys.executable, "-m", "resume", "render",
            "--data", data_path,
            "--template", template,
            "--profile", profile,
            "--out", out_path,
        ])
        case.assertEqual(proc.returncode, 0, msg=proc.stderr)
        case.assertTrue(os.path.exists(out_path))

        from docx import Document  # type: ignore

        doc = Document(out_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        case.assertIn("New Role at NewCo", text)
        case.assertNotIn("Old Role at OldCo", text)
    finally:
        try:
            os.unlink(cfg_path)
        except OSError:  # nosec B110 - test cleanup
            pass


class TestResumeCLIExtract(unittest.TestCase):
    """Test resume extract command functionality."""

    def test_extract_missing_args(self):
        """Test extract with no inputs still works (produces empty data)."""
        with tempfile.TemporaryDirectory() as tmpdir:
            out_path = os.path.join(tmpdir, "data.json")
            proc = run([
                sys.executable, "-m", "resume", "extract",
                "--out", out_path,
            ])
            self.assertEqual(proc.returncode, 0, msg=proc.stderr)
            self.assertTrue(os.path.exists(out_path))

    def test_extract_with_nonexistent_file(self):
        """Test extract with nonexistent file reports error."""
        with tempfile.TemporaryDirectory() as tmpdir:
            out_path = os.path.join(tmpdir, "data.json")
            proc = run([
                sys.executable, "-m", "resume", "extract",
                "--linkedin", "/nonexistent/file.txt",
                "--out", out_path,
            ])
            # Command may succeed with empty data or fail - check stderr for error message
            # The io_utils.read_text_any handles missing files gracefully
            # Just verify it doesn't crash unexpectedly
            self.assertIn(proc.returncode, [0, 1])


class TestResumeCLISummarize(unittest.TestCase):
    """Test resume summarize command functionality."""

    def test_summarize_with_minimal_data(self):
        """Test summarize with minimal candidate data."""
        data = {
            "name": "Test User",
            "headline": "Software Engineer",
            "skills": ["Python", "Java"],
            "experience": [
                {
                    "title": "Engineer",
                    "company": "Acme Corp",
                    "bullets": ["Built systems"],
                }
            ],
        }
        with temp_yaml_file(data) as data_path:
            with tempfile.TemporaryDirectory() as tmpdir:
                out_path = os.path.join(tmpdir, "summary.md")
                proc = run([
                    sys.executable, "-m", "resume", "summarize",
                    "--data", data_path,
                    "--out", out_path,
                ])
                self.assertEqual(proc.returncode, 0, msg=proc.stderr)
                self.assertTrue(os.path.exists(out_path))


class TestResumeCLICandidateInit(unittest.TestCase):
    """Test resume candidate-init command functionality."""

    def test_candidate_init_basic(self):
        """Test candidate-init generates skills YAML."""
        data = {
            "name": "Test User",
            "headline": "Software Engineer",
            "email": "test@example.com",
            "skills": ["Python", "Java", "Go"],
        }
        with temp_yaml_file(data) as data_path:
            with tempfile.TemporaryDirectory() as tmpdir:
                out_path = os.path.join(tmpdir, "candidate.yaml")
                proc = run([
                    sys.executable, "-m", "resume", "candidate-init",
                    "--data", data_path,
                    "--out", out_path,
                ])
                self.assertEqual(proc.returncode, 0, msg=proc.stderr)
                self.assertTrue(os.path.exists(out_path))


class TestResumeCLIAlign(unittest.TestCase):
    """Test resume align command functionality."""

    def test_align_basic(self):
        """Test align produces alignment report."""
        candidate = {
            "name": "Test User",
            "skills": ["Python", "AWS", "Docker"],
            "experience": [
                {
                    "title": "Engineer",
                    "company": "Acme",
                    "bullets": ["Built Python services", "Deployed to AWS"],
                }
            ],
        }
        job = {
            "title": "Senior Engineer",
            "required_skills": ["Python", "AWS"],
            "preferred_skills": ["Docker", "Kubernetes"],
        }
        with temp_yaml_file(candidate) as cand_path:
            with temp_yaml_file(job) as job_path:
                with tempfile.TemporaryDirectory() as tmpdir:
                    out_path = os.path.join(tmpdir, "alignment.json")
                    proc = run([
                        sys.executable, "-m", "resume", "align",
                        "--data", cand_path,
                        "--job", job_path,
                        "--out", out_path,
                    ])
                    self.assertEqual(proc.returncode, 0, msg=proc.stderr)
                    self.assertTrue(os.path.exists(out_path))


class TestResumeCLIRender(unittest.TestCase):
    """Test resume render command functionality."""

    def test_render_basic(self):
        """Test render produces a DOCX file."""
        data = {
            "name": "Test User",
            "headline": "Software Engineer",
            "email": "test@example.com",
            "skills": ["Python", "Java"],
            "experience": [
                {
                    "title": "Engineer",
                    "company": "Acme Corp",
                    "start": "2020",
                    "end": "Present",
                    "bullets": ["Built systems", "Led team"],
                }
            ],
        }
        template = {
            "sections": [
                {"key": "header"},
                {"key": "experience", "title": "Experience"},
            ]
        }
        with temp_yaml_file(data) as data_path:
            with temp_yaml_file(template) as template_path:
                with tempfile.TemporaryDirectory() as tmpdir:
                    out_path = os.path.join(tmpdir, "resume.docx")
                    proc = run([
                        sys.executable, "-m", "resume", "render",
                        "--data", data_path,
                        "--template", template_path,
                        "--out", out_path,
                    ])
                    self.assertEqual(proc.returncode, 0, msg=proc.stderr)
                    self.assertTrue(os.path.exists(out_path))

    def test_render_with_seed(self):
        """Test render with seed criteria."""
        data = {
            "name": "Test User",
            "skills": ["Python"],
        }
        with temp_yaml_file(data) as data_path:
            with tempfile.TemporaryDirectory() as tmpdir:
                out_path = os.path.join(tmpdir, "resume.docx")
                proc = run([
                    sys.executable, "-m", "resume", "render",
                    "--data", data_path,
                    "--seed", "keywords=python,aws",
                    "--out", out_path,
                ])
                self.assertEqual(proc.returncode, 0, msg=proc.stderr)

    def test_render_without_template_emits_body_sections(self):
        """render with no --template must use DEFAULT_TEMPLATE, not an empty dict.

        Regression: cmd_render read `load_template(args.template) if args.template
        else {}`. The `else {}` defeated load_template's own None-fallback, so an
        omitted --template resolved to zero sections and the writer emitted only
        the header. It exited 0 and wrote a plausible file, so nothing surfaced
        the failure — the resume was simply missing its entire body.

        Asserting returncode/existence alone cannot catch this; the assertions
        below are on rendered CONTENT.
        """
        from docx import Document

        data = {
            "name": "Test User",
            "headline": "Software Engineer",
            "summary": "Builds reliable systems.",
            "skills": ["Python", "Kubernetes"],
            "experience": [
                {
                    "title": "Engineer",
                    "company": "Acme Corp",
                    "start": "2020",
                    "end": "Present",
                    "bullets": ["Cut latency by 40%"],
                }
            ],
            "education": [
                {"degree": "BSc", "institution": "State University", "year": "2015"}
            ],
        }
        with temp_yaml_file(data) as data_path:
            with tempfile.TemporaryDirectory() as tmpdir:
                out_path = os.path.join(tmpdir, "resume.docx")
                proc = run([
                    sys.executable, "-m", "resume", "render",
                    "--data", data_path,
                    "--out", out_path,
                ])
                self.assertEqual(proc.returncode, 0, msg=proc.stderr)
                self.assertTrue(os.path.exists(out_path))

                text = "\n".join(
                    p.text for p in Document(out_path).paragraphs if p.text.strip()
                )

                # Header alone is what the bug produced; it is not sufficient.
                self.assertIn("Test User", text)

                # Each default section must actually render.
                for heading in ("Summary", "Skills", "Experience", "Education"):
                    self.assertIn(heading, text, f"missing section heading: {heading}")

                # And carry its content, not just a bare heading.
                self.assertIn("Builds reliable systems.", text)
                self.assertIn("Python", text)
                self.assertIn("Acme Corp", text)
                self.assertIn("Cut latency by 40%", text)
                self.assertIn("State University", text)

    def test_render_pdf_raises_error(self):
        """Test render with .pdf output exits with USAGE error and directs user to export-pdf.

        BLOCKER 1b: The stub message no longer contains uppercase "PDF"; assert on stable
        behaviour — exit USAGE (2) and that the message points the user at export-pdf.
        """
        data = {"name": "Test User"}
        with temp_yaml_file(data) as data_path:
            with tempfile.TemporaryDirectory() as tmpdir:
                out_path = os.path.join(tmpdir, "resume.pdf")
                proc = run([
                    sys.executable, "-m", "resume", "render",
                    "--data", data_path,
                    "--out", out_path,
                ])
                # ExitCode.USAGE == 2 (C5 fix replaced RuntimeError with CLIError(..., ExitCode.USAGE))
                self.assertEqual(proc.returncode, 2)
                # The stub message directs the user to the new subcommand; assert on that
                # rather than on a specific casing that may change with future rewording.
                self.assertIn("export-pdf", proc.stderr)

    def test_render_onepage_template_trims_roles_and_bullets(self):
        """E2E: the onepage template caps margins, roles, and bullets per role.

        Note: --template must point at the packaged template under
        src/resume/config/ — a bare "config/template.onepage.yaml" only
        resolves when the process cwd happens to be src/resume/.
        """
        try:
            import docx  # noqa: F401
        except Exception:
            self.skipTest("python-docx not installed")
        with tempfile.TemporaryDirectory() as tmpdir:
            data_path = os.path.join(tmpdir, "data.json")
            out_path = os.path.join(tmpdir, "out.docx")
            exp = []
            for i in range(1, 6):
                exp.append({
                    "title": f"Engineer {i}",
                    "company": f"Co{i}",
                    "start": f"202{i}",
                    "end": f"202{i + 1}",
                    "location": "Remote",
                    "bullets": [f"Did thing {j} for role {i}" for j in range(1, 6)],
                })
            payload = {
                "name": "Test Person",
                "email": "test@example.com",
                "phone": "5550001111",
                "location": "Somewhere, ZZ",
                "summary": "Impactful engineer. Builds reliable systems. Ships results.",
                "skills_groups": [
                    {"title": "Platform", "items": ["AWS", "Kubernetes", "Architecture"]},
                    {"title": "Reliability", "items": ["SLOs", "On-call", "Incidents", "Postmortems"]},
                ],
                "experience": exp,
                "education": [{"degree": "BS CS", "institution": "Uni", "year": "2012"}],
            }
            with open(data_path, "w", encoding="utf-8") as f:
                json.dump(payload, f)

            proc = run([
                sys.executable, "-m", "resume", "render",
                "--data", data_path,
                "--template", "src/resume/config/template.onepage.yaml",
                "--out", out_path,
            ])
            self.assertEqual(proc.returncode, 0, msg=proc.stderr)
            self.assertTrue(os.path.exists(out_path))

            from docx import Document  # type: ignore

            doc = Document(out_path)
            sec = doc.sections[0]

            def emu_to_in(v):
                try:
                    return v.inches
                except Exception:
                    return float(v) / 914400.0

            # Onepage template caps margins at 0.4"
            self.assertLessEqual(emu_to_in(sec.left_margin), 0.5 + 1e-6)
            text = "\n".join(p.text for p in doc.paragraphs)
            # Only the first 4 roles are rendered
            self.assertIn("Engineer 4 at Co4", text)
            self.assertNotIn("Engineer 5 at Co5", text)
            # Bullets per role trimmed to <=3 (spot check role 1)
            self.assertIn("Did thing 3 for role 1", text)
            self.assertNotIn("Did thing 4 for role 1", text)

    def test_render_experience_overlay_replaces_base_role(self):
        """E2E: a profile experience.yaml overlay replaces the base data's role."""
        try:
            import docx  # noqa: F401
        except Exception:
            self.skipTest("python-docx not installed")
        with tempfile.TemporaryDirectory() as tmpdir:
            _run_experience_overlay_render(
                self, tmpdir,
                profile="test_exp_overlay_profile",
                template="src/resume/config/template.onepage.yaml",
            )


class TestResumeCLIStructure(unittest.TestCase):
    """Test resume structure command functionality."""

    # test_structure_help removed — identical body already lives at
    # tests/resume_tests/cli/test_resume_cli_core.py::TestResumeCLISubcommandHelp::test_structure_help
    # which is where all sibling *_help checks are colocated.

    def test_structure_infers_order_from_real_docx(self):
        """E2E: infer section order from an actual rendered reference DOCX.

        Unlike the mocked cmd_structure unit test, this exercises the real
        infer_structure_from_docx heading-detection logic end to end.
        """
        try:
            from docx import Document  # type: ignore
        except Exception:
            self.skipTest("python-docx not installed")
        with tempfile.TemporaryDirectory() as tmpdir:
            ref = os.path.join(tmpdir, "ref.docx")
            out = os.path.join(tmpdir, "struct.json")
            doc = Document()
            doc.add_heading("Profile", level=1)
            doc.add_paragraph("...")
            doc.add_heading("Work History", level=1)
            doc.add_paragraph("...")
            doc.add_heading("Technical Skills", level=1)
            doc.add_paragraph("...")
            doc.add_heading("Education", level=1)
            doc.add_paragraph("...")
            doc.save(ref)

            proc = run([sys.executable, "-m", "resume", "structure", "--source", ref, "--out", out])
            self.assertEqual(proc.returncode, 0, msg=proc.stderr)
            self.assertTrue(os.path.exists(out))
            with open(out, encoding="utf-8") as f:
                data = json.load(f)
            self.assertIn("order", data)
            self.assertIn("experience", data.get("order", []))


class TestResumeCLIStyleBuild(unittest.TestCase):
    """Test resume style build command functionality."""

    def test_style_build_with_corpus(self):
        """Test style build produces style profile."""
        with tempfile.TemporaryDirectory() as corpus_dir:
            # Create sample corpus files
            sample_path = os.path.join(corpus_dir, "sample.txt")
            with open(sample_path, "w") as f:
                f.write("Python developer with experience in AWS and Docker.\n")
                f.write("Built microservices and REST APIs.\n")
            with tempfile.TemporaryDirectory() as out_dir:
                out_path = os.path.join(out_dir, "style.json")
                proc = run([
                    sys.executable, "-m", "resume", "style", "build",
                    "--corpus-dir", corpus_dir,
                    "--out", out_path,
                ])
                self.assertEqual(proc.returncode, 0, msg=proc.stderr)
                self.assertTrue(os.path.exists(out_path))


class TestResumeCLIFilesTidy(unittest.TestCase):
    """Test resume files tidy command functionality."""

    def test_files_tidy_archive(self):
        """Test files tidy archives old files."""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Create test files with different timestamps
            for i in range(4):
                path = os.path.join(tmpdir, f"test_{i}.json")
                with open(path, "w") as f:
                    f.write("{}")
                # Modify mtime to simulate age
                import time
                os.utime(path, (time.time() - i * 86400, time.time() - i * 86400))

            proc = run([
                sys.executable, "-m", "resume", "files", "tidy",
                "--dir", tmpdir,
                "--suffixes", ".json",
                "--keep", "2",
            ])
            self.assertEqual(proc.returncode, 0, msg=proc.stderr)

    def test_files_tidy_delete(self):
        """Test files tidy with --delete flag."""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Create test files
            for i in range(4):
                path = os.path.join(tmpdir, f"test_{i}.json")
                with open(path, "w") as f:
                    f.write("{}")
                import time
                os.utime(path, (time.time() - i * 86400, time.time() - i * 86400))

            proc = run([
                sys.executable, "-m", "resume", "files", "tidy",
                "--dir", tmpdir,
                "--suffixes", ".json",
                "--keep", "2",
                "--delete",
            ])
            self.assertEqual(proc.returncode, 0, msg=proc.stderr)

    def test_files_tidy_purge_temp(self):
        """Test files tidy with --purge-temp flag."""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Create a temp file
            temp_path = os.path.join(tmpdir, "~$test.docx")
            with open(temp_path, "w") as f:
                f.write("temp")

            proc = run([
                sys.executable, "-m", "resume", "files", "tidy",
                "--dir", tmpdir,
                "--purge-temp",
            ])
            self.assertEqual(proc.returncode, 0, msg=proc.stderr)
            # Temp file should be removed
            self.assertFalse(os.path.exists(temp_path))


class TestResumeCLIExperienceExport(unittest.TestCase):
    """Test resume experience export command functionality."""

    def test_experience_export_from_data(self):
        """Test experience export from data file."""
        data = {
            "name": "Test User",
            "experience": [
                {
                    "title": "Engineer",
                    "company": "Acme Corp",
                    "start": "2020",
                    "end": "Present",
                    "bullets": ["Built systems", "Led team", "Optimized performance"],
                }
            ],
        }
        with temp_yaml_file(data) as data_path:
            with tempfile.TemporaryDirectory() as tmpdir:
                out_path = os.path.join(tmpdir, "experience.yaml")
                proc = run([
                    sys.executable, "-m", "resume", "experience", "export",
                    "--data", data_path,
                    "--out", out_path,
                ])
                self.assertEqual(proc.returncode, 0, msg=proc.stderr)
                self.assertTrue(os.path.exists(out_path))

    def test_experience_export_with_max_bullets(self):
        """Test experience export with --max-bullets."""
        data = {
            "name": "Test User",
            "experience": [
                {
                    "title": "Engineer",
                    "company": "Acme",
                    "bullets": ["One", "Two", "Three", "Four", "Five"],
                }
            ],
        }
        with temp_yaml_file(data) as data_path:
            with tempfile.TemporaryDirectory() as tmpdir:
                out_path = os.path.join(tmpdir, "experience.yaml")
                proc = run([
                    sys.executable, "-m", "resume", "experience", "export",
                    "--data", data_path,
                    "--max-bullets", "2",
                    "--out", out_path,
                ])
                self.assertEqual(proc.returncode, 0, msg=proc.stderr)

    def test_experience_export_missing_args(self):
        """Test experience export with no --data or --resume raises error."""
        with tempfile.TemporaryDirectory() as tmpdir:
            out_path = os.path.join(tmpdir, "experience.yaml")
            proc = run([
                sys.executable, "-m", "resume", "experience", "export",
                "--out", out_path,
            ])
            self.assertEqual(proc.returncode, 1)
            self.assertIn("--data", proc.stderr.lower() + proc.stdout.lower())


class TestResumeCLICandidateInitExtended(unittest.TestCase):
    """Test resume candidate-init command extended functionality."""

    def test_candidate_init_with_experience(self):
        """Test candidate-init with --include-experience flag."""
        data = {
            "name": "Test User",
            "headline": "Software Engineer",
            "skills": ["Python", "Java"],
            "experience": [
                {
                    "title": "Engineer",
                    "company": "Acme",
                    "start": "2020",
                    "end": "Present",
                    "location": "Remote",
                    "bullets": ["Built systems", "Led team", "Optimized code", "Deployed apps"],
                }
            ],
        }
        with temp_yaml_file(data) as data_path:
            with tempfile.TemporaryDirectory() as tmpdir:
                out_path = os.path.join(tmpdir, "candidate.yaml")
                proc = run([
                    sys.executable, "-m", "resume", "candidate-init",
                    "--data", data_path,
                    "--include-experience",
                    "--max-bullets", "2",
                    "--out", out_path,
                ])
                self.assertEqual(proc.returncode, 0, msg=proc.stderr)
                self.assertTrue(os.path.exists(out_path))
                # Verify experience is included with limited bullets
                import yaml
                with open(out_path) as f:
                    result = yaml.safe_load(f)
                self.assertIn("experience", result)
                self.assertEqual(len(result["experience"][0]["bullets"]), 2)


class TestResumeCLIOnePageTemplateE2E(unittest.TestCase):
    """E2E coverage for the onepage template's margin/role/bullet trimming.

    Known pre-existing defect, not introduced by this port: render passes
    --template config/template.onepage.yaml as a literal relative path
    straight into resume.templating.load_template -> read_yaml_or_json,
    which resolves it against the subprocess's CWD. The real file lives at
    src/resume/config/template.onepage.yaml, not <repo-root>/config/. These
    tests only ever passed when CWD happened to be src/resume/ (e.g. run
    directly from within the package tests dir) — from the repo root, which
    is how `make test` / unittest discovery actually run, `render` exits 1
    with "Error: config/template.onepage.yaml" before the docx is written.
    Skipped rather than silently dropped or "fixed" by reaching into
    templating/CLI path resolution, which is out of scope for a test-dedup
    pass. See followup-dead-tests.md for the report.
    """

    ONEPAGE_TEMPLATE_BUG_REASON = (
        "render resolves --template config/template.onepage.yaml against CWD; "
        "the file actually lives at src/resume/config/template.onepage.yaml, "
        "so this only passes when CWD is src/resume/, not from the repo root"
    )

    @unittest.skip(ONEPAGE_TEMPLATE_BUG_REASON)
    def test_onepage_render_trims_roles_and_bullets_within_margins(self):
        try:
            import docx  # noqa: F401
        except Exception:
            self.skipTest("python-docx not installed")

        with tempfile.TemporaryDirectory() as tmpdir:
            data_path = os.path.join(tmpdir, "data.json")
            out_path = os.path.join(tmpdir, "out.docx")
            exp = []
            for i in range(1, 6):
                exp.append({
                    "title": f"Engineer {i}",
                    "company": f"Co{i}",
                    "start": f"202{i}",
                    "end": f"202{i + 1}",
                    "location": "Remote",
                    "bullets": [f"Did thing {j} for role {i}" for j in range(1, 6)],
                })
            payload = {
                "name": "Test Person",
                "email": "test@example.com",
                "phone": "5550001111",
                "location": "Somewhere, ZZ",
                "summary": "Impactful engineer. Builds reliable systems. Ships results.",
                "skills_groups": [
                    {"title": "Platform", "items": ["AWS", "Kubernetes", "Architecture"]},
                    {"title": "Reliability", "items": ["SLOs", "On-call", "Incidents", "Postmortems"]},
                ],
                "experience": exp,
                "education": [{"degree": "BS CS", "institution": "Uni", "year": "2012"}],
            }
            with open(data_path, "w", encoding="utf-8") as f:
                json.dump(payload, f)

            proc = run([
                sys.executable, "-m", "resume", "render",
                "--data", data_path,
                "--template", "config/template.onepage.yaml",
                "--out", out_path,
            ])
            self.assertEqual(proc.returncode, 0, msg=proc.stderr)
            self.assertTrue(os.path.exists(out_path))

            from docx import Document  # type: ignore

            doc = Document(out_path)
            sec = doc.sections[0]

            def emu_to_in(v):
                try:
                    return v.inches
                except Exception:
                    return float(v) / 914400.0

            self.assertLessEqual(emu_to_in(sec.left_margin), 0.5 + 1e-6)
            text = "\n".join(p.text for p in doc.paragraphs)
            self.assertIn("Engineer 4 at Co4", text)
            self.assertNotIn("Engineer 5 at Co5", text)
            self.assertIn("Did thing 3 for role 1", text)
            self.assertNotIn("Did thing 4 for role 1", text)

    @unittest.skip(ONEPAGE_TEMPLATE_BUG_REASON)
    def test_onepage_render_overlays_experience_from_profile_config(self):
        try:
            import docx  # noqa: F401
        except Exception:
            self.skipTest("python-docx not installed")

        with tempfile.TemporaryDirectory() as tmpdir:
            _run_experience_overlay_render(
                self, tmpdir,
                profile="test_profile",
                template="config/template.onepage.yaml",
            )


class TestResumeCLIAlignExtended(unittest.TestCase):
    """Test resume align command extended functionality."""

    def test_align_resolves_synonyms_and_categories_from_job_yaml(self):
        """E2E: real job.yaml with weighted tiers, categories, and synonyms.

        Unlike the fully-mocked cmd_align unit tests, this exercises real
        build_keyword_spec + align_candidate_to_job + build_tailored_candidate
        integration, including synonym resolution ("k8s" -> "Kubernetes").
        """
        candidate = {
            "name": "Alex Test",
            "skills": ["Python", "AWS", "Docker", "Kubernetes", "Grafana"],
            "experience": [
                {
                    "title": "SRE", "company": "Acme", "start": "2020", "end": "2024", "location": "Remote",
                    "bullets": ["Managed Kubernetes clusters", "Built dashboards in Grafana", "Automated deployments with Docker"],
                },
                {
                    "title": "DevOps", "company": "Other", "start": "2018", "end": "2020", "location": "Remote",
                    "bullets": ["Provisioned AWS infrastructure", "Wrote Python tooling"],
                },
            ],
        }
        job_yaml = """
title: Senior SRE
company: ExampleCo
keywords:
  required:
    - skill: Kubernetes
      weight: 3
    - skill: AWS
      weight: 3
  preferred:
    - skill: Grafana
      weight: 2
  soft_skills:
    - Communication
    - Collaboration
  tech_skills:
    - Python
    - Docker
  synonyms:
    Kubernetes: [k8s, EKS]
""".strip()
        with tempfile.TemporaryDirectory() as tmpdir:
            cand_path = os.path.join(tmpdir, "cand.json")
            job_path = os.path.join(tmpdir, "job.yaml")
            align_out = os.path.join(tmpdir, "align.json")
            tailored_out = os.path.join(tmpdir, "tailored.json")
            with open(cand_path, "w", encoding="utf-8") as f:
                json.dump(candidate, f)
            with open(job_path, "w", encoding="utf-8") as f:
                f.write(job_yaml)

            proc = run([
                sys.executable, "-m", "resume", "align",
                "--data", cand_path,
                "--job", job_path,
                "--out", align_out,
                "--tailored", tailored_out,
                "--max-bullets", "2",
            ])
            self.assertEqual(proc.returncode, 0, msg=proc.stderr)

            with open(align_out, encoding="utf-8") as f:
                a = json.load(f)
            self.assertIn("matched_keywords", a)
            matched_skills = {x["skill"] for x in a["matched_keywords"]}
            self.assertIn("Kubernetes", matched_skills)
            self.assertIn("AWS", matched_skills)
            self.assertIn("missing_by_category", a)

            with open(tailored_out, encoding="utf-8") as f:
                t = json.load(f)
            self.assertIn("experience", t)
            self.assertTrue(any(
                "Kubernetes" in b or "Grafana" in b or "AWS" in b
                for e in t["experience"] for b in e.get("bullets", [])
            ))

    def test_align_with_tailored_output(self):
        """Test align with --tailored flag produces tailored candidate."""
        candidate = {
            "name": "Test User",
            "skills": ["Python", "AWS"],
            "experience": [
                {
                    "title": "Engineer",
                    "company": "Acme",
                    "bullets": ["Built Python services", "Deployed to AWS", "Extra bullet"],
                }
            ],
        }
        job = {
            "title": "Senior Engineer",
            "required_skills": ["Python"],
        }
        with temp_yaml_file(candidate) as cand_path:
            with temp_yaml_file(job) as job_path:
                with tempfile.TemporaryDirectory() as tmpdir:
                    align_path = os.path.join(tmpdir, "alignment.json")
                    tailored_path = os.path.join(tmpdir, "tailored.yaml")
                    proc = run([
                        sys.executable, "-m", "resume", "align",
                        "--data", cand_path,
                        "--job", job_path,
                        "--tailored", tailored_path,
                        "--max-bullets", "2",
                        "--out", align_path,
                    ])
                    self.assertEqual(proc.returncode, 0, msg=proc.stderr)
                    self.assertTrue(os.path.exists(tailored_path))


class TestResumeCLISummarizeExtended(unittest.TestCase):
    """Test resume summarize command extended functionality."""

    def test_summarize_to_yaml(self):
        """Test summarize with YAML output."""
        data = {
            "name": "Test User",
            "headline": "Software Engineer",
            "skills": ["Python"],
            "experience": [{"title": "Engineer", "company": "Acme", "bullets": ["Built systems"]}],
        }
        with temp_yaml_file(data) as data_path:
            with tempfile.TemporaryDirectory() as tmpdir:
                out_path = os.path.join(tmpdir, "summary.yaml")
                proc = run([
                    sys.executable, "-m", "resume", "summarize",
                    "--data", data_path,
                    "--out", out_path,
                ])
                self.assertEqual(proc.returncode, 0, msg=proc.stderr)
                self.assertTrue(os.path.exists(out_path))

    def test_summarize_with_seed(self):
        """Test summarize with seed criteria."""
        data = {
            "name": "Test User",
            "skills": ["Python", "AWS"],
            "experience": [{"title": "Engineer", "company": "Acme", "bullets": ["Built systems"]}],
        }
        with temp_yaml_file(data) as data_path:
            with tempfile.TemporaryDirectory() as tmpdir:
                out_path = os.path.join(tmpdir, "summary.md")
                proc = run([
                    sys.executable, "-m", "resume", "summarize",
                    "--data", data_path,
                    "--seed", "keywords=python,cloud",
                    "--out", out_path,
                ])
                self.assertEqual(proc.returncode, 0, msg=proc.stderr)


class TestResumeCLIExportPdf(unittest.TestCase):
    """Test resume export-pdf subcommand.

    BLOCKER 4 strategy: Option A (subprocess-based) for all tests.
    All tests drive the CLI as a subprocess (matching the rest of this file).
    The soffice-absent sad path is made hermetic by passing a PATH that
    excludes soffice, which causes convert_docx_to_pdf's FileNotFoundError
    branch to fire deterministically — no parent-process mocking needed.
    """

    # --------------------------------------------------------------------------
    # Helpers
    # --------------------------------------------------------------------------

    @staticmethod
    def _run_export_pdf(args: list[str], *, env: dict | None = None):
        """Run `resume export-pdf <args>` as a subprocess.

        When *env* is None, the inherited environment is used.  Pass a dict to
        override the entire environment (e.g. to scrub PATH of soffice).
        """
        import subprocess  # nosec B404
        cmd = [sys.executable, "-m", "resume", "export-pdf"] + args
        return subprocess.run(  # nosec B603
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            env=env,
        )

    # --------------------------------------------------------------------------
    # Happy path
    # --------------------------------------------------------------------------

    def test_export_pdf_happy_path_output_file_exists(self):
        """Happy path: conversion succeeds and the output PDF exists at the resolved path.

        Uses a real (but minimal) .docx created by python-docx so the file is
        valid.  Requires soffice on PATH; skipped when absent.
        """
        import shutil
        if not shutil.which("soffice"):
            self.skipTest("soffice not installed — happy-path test skipped")
        try:
            from docx import Document  # noqa: F401
        except Exception:
            self.skipTest("python-docx not installed")

        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            docx_path = os.path.join(tmpdir, "myresume.docx")
            out_path = os.path.join(tmpdir, "myresume.pdf")
            Document().save(docx_path)

            proc = self._run_export_pdf([
                "--docx", docx_path,
                "--out", out_path,
            ])
            self.assertEqual(proc.returncode, 0, msg=proc.stderr)
            self.assertTrue(os.path.exists(out_path), msg=f"PDF not found; stderr={proc.stderr}")

    # --------------------------------------------------------------------------
    # Sad path: soffice absent / conversion failure
    # --------------------------------------------------------------------------

    def test_export_pdf_soffice_absent_exits_error_with_libreoffice_message(self):
        """Sad path: soffice absent on PATH → exit ERROR (1) and stderr names LibreOffice.

        BLOCKER 4 (Option A): pass PATH=/nonexistent so subprocess cannot find soffice.
        BLOCKER 5: asserts BOTH returncode and stderr substring.
        BLOCKER 6: conversion failure exits ERROR (1), not USAGE (2).
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            # Create a minimal but real-looking .docx so the CLI gets past the
            # file-existence check and actually attempts conversion.
            docx_path = os.path.join(tmpdir, "resume.docx")
            try:
                from docx import Document
                Document().save(docx_path)
            except Exception:
                # Fallback: write a stub file that passes the existence check.
                with open(docx_path, "wb") as f:
                    f.write(b"PK\x03\x04")  # PKZIP magic — not a valid docx but the file exists

            out_path = os.path.join(tmpdir, "resume.pdf")
            # Scrub PATH so soffice cannot be found — this makes convert_docx_to_pdf
            # raise FileNotFoundError, which is returned as False → CLIError exit 1.
            empty_path_env = dict(os.environ, PATH="/nonexistent")

            proc = self._run_export_pdf([
                "--docx", docx_path,
                "--out", out_path,
            ], env=empty_path_env)

            # BLOCKER 6: conversion failure is ERROR (1)
            self.assertEqual(proc.returncode, 1, msg=f"expected exit 1; stderr={proc.stderr}")
            # BLOCKER 5: stderr must mention LibreOffice (the remedy named in the error message)
            self.assertIn("LibreOffice", proc.stderr, msg=f"expected 'LibreOffice' in stderr; got: {proc.stderr!r}")

    # --------------------------------------------------------------------------
    # Sad path: missing --docx file
    # --------------------------------------------------------------------------

    def test_export_pdf_missing_docx_exits_usage_with_not_found_message(self):
        """Sad path: --docx file does not exist → exit USAGE (2) and stderr says not found.

        BLOCKER 5: asserts BOTH returncode and stderr substring.
        BLOCKER 6: missing input file must exit USAGE (2), not ERROR (1).
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            absent_docx = os.path.join(tmpdir, "ghost.docx")
            out_path = os.path.join(tmpdir, "ghost.pdf")
            # absent_docx deliberately not created

            proc = self._run_export_pdf([
                "--docx", absent_docx,
                "--out", out_path,
            ])

            # BLOCKER 6: missing file → USAGE (2)
            self.assertEqual(proc.returncode, 2, msg=f"expected exit 2; stderr={proc.stderr}")
            # BLOCKER 5: stderr must say something about the file not being found
            self.assertIn("not found", proc.stderr.lower(), msg=f"expected 'not found' in stderr; got: {proc.stderr!r}")

    # --------------------------------------------------------------------------
    # Regression: distinct inputs must not collide on the same output path
    # --------------------------------------------------------------------------

    def test_export_pdf_different_inputs_resolve_to_different_output_paths(self):
        """Regression: two different --docx inputs must not resolve to the same output path.

        BLOCKER 3 follow-up: with kind=docx.stem, omitting --out uses the stem of the
        input file, so 'alpha.docx' → '<data-home>/.../alpha.pdf' and
        'beta.docx' → '<data-home>/.../beta.pdf'.  They must differ.
        This test verifies the _resolve_out behaviour without actually running soffice.
        """
        from resume.cli.main import _resolve_out  # type: ignore[import]
        import argparse

        with tempfile.TemporaryDirectory() as tmpdir:
            # Simulate args with no explicit --out and a custom --out-dir so
            # we don't write to the real data home.
            def make_args() -> argparse.Namespace:
                return argparse.Namespace(out=None, out_dir=tmpdir, profile=None)

            alpha_args = make_args()
            beta_args = make_args()

            # _resolve_out(args, suffix, kind=stem) is the post-BLOCKER-3 call
            alpha_path = _resolve_out(alpha_args, ".pdf", kind="alpha")
            beta_path = _resolve_out(beta_args, ".pdf", kind="beta")

            self.assertNotEqual(
                str(alpha_path),
                str(beta_path),
                msg=(
                    f"Two different inputs resolve to the same output path: {alpha_path!r}. "
                    "This is the silent-overwrite bug from BLOCKER 3."
                ),
            )


if __name__ == "__main__":
    unittest.main(verbosity=2)
