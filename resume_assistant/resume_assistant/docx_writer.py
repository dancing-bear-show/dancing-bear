"""DOCX resume writer.

Renders resume data to DOCX format using templates and styling configuration.
"""
from __future__ import annotations

from typing import Any, Dict, List, Optional

from .io_utils import safe_import
from docx.shared import Pt, Inches, RGBColor  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

# Import from new abstraction modules
from .docx_styles import (
    StyleManager,
    TextFormatter,
    # Backward-compatible aliases
    _parse_hex_color,
    _hex_fill,
    _is_dark,
    _tight_paragraph,
    _compact_bullet,
    _flush_left,
    _apply_paragraph_shading,
    _normalize_present,
    _format_date_location,
    _format_phone_display,
    _format_link_display,
    _clean_inline_text,
    _normalize_bullet_text,
)
from .docx_sections import (
    BulletRenderer,
    HeaderRenderer,
    ListSectionRenderer,
    InterestsSectionRenderer,
    LanguagesSectionRenderer,
    CourseworkSectionRenderer,
    CertificationsSectionRenderer,
    PresentationsSectionRenderer,
)


SECTION_SYNONYMS = {
    "summary": {"summary", "profile", "about"},
    "skills": {"skills", "technical skills"},
    "technologies": {"technologies", "technology", "tools"},
    "experience": {"experience", "work history", "employment"},
    "education": {"education", "academics"},
}


def _match_section_key(title: str) -> Optional[str]:
    t = title.strip().lower()
    for key, names in SECTION_SYNONYMS.items():
        if t in names:
            return key
    return None


# Backward-compatible function aliases that delegate to renderers
def _bold_keywords(paragraph, text: str, keywords: List[str]):
    """Bold keywords in paragraph text."""
    renderer = BulletRenderer.__new__(BulletRenderer)
    renderer._bold_keywords(paragraph, text, keywords)


def _add_bullet_line(doc, text: str, *, sec: Dict[str, Any] | None = None, keywords: List[str] | None = None, glyph: str = "•"):
    renderer = BulletRenderer(doc)
    return renderer.add_bullet_line(text, sec=sec, keywords=keywords, glyph=glyph)


def _add_plain_bullet(doc, text: str, keywords: List[str] | None = None):
    return _add_bullet_line(doc, text, keywords=keywords, glyph="•")


def _add_bullets(
    doc,
    items: List[str],
    *,
    sec: Dict[str, Any] | None = None,
    keywords: List[str] | None = None,
    plain: bool = True,
    glyph: str = "•",
    list_style: str = "List Bullet",
):
    renderer = BulletRenderer(doc)
    renderer.add_bullets(items, sec=sec, keywords=keywords, plain=plain, glyph=glyph, list_style=list_style)


def _render_group_title(doc, title: str, sec: Dict[str, Any] | None = None):
    renderer = HeaderRenderer(doc)
    return renderer.add_group_title(title, sec)


def _add_header_line(
    doc,
    *,
    title_text: str = "",
    company_text: str = "",
    loc_text: str = "",
    span_text: str = "",
    sec: Dict[str, Any] | None = None,
    style: str = "Normal",
):
    renderer = HeaderRenderer(doc)
    return renderer.add_header_line(
        title_text=title_text,
        company_text=company_text,
        loc_text=loc_text,
        span_text=span_text,
        sec=sec,
        style=style,
    )


def _add_named_bullet(
    doc,
    name_text: str,
    desc_text: str,
    *,
    sec: Dict[str, Any] | None = None,
    glyph: str = "•",
    sep: str = ": ",
):
    renderer = BulletRenderer(doc)
    return renderer.add_named_bullet(name_text, desc_text, sec=sec, glyph=glyph, sep=sep)


def _get_header_level(sec: Dict[str, Any] | None, page_cfg: Dict[str, Any] | None) -> int:
    try:
        if sec and isinstance(sec.get("header_level"), int):
            return int(sec.get("header_level"))
        if page_cfg and isinstance(page_cfg.get("header_level"), int):
            return int(page_cfg.get("header_level"))
    except Exception:
        pass
    return 1


def _use_plain_bullets(sec: Dict[str, Any] | None, page_cfg: Dict[str, Any] | None) -> tuple:
    renderer = BulletRenderer.__new__(BulletRenderer)
    renderer.page_cfg = page_cfg or {}
    return renderer.get_bullet_config(sec)


def write_resume_docx(
    data: Dict[str, Any],
    template: Dict[str, Any],
    out_path: str,
    seed: Optional[Dict[str, Any]] = None,
    structure: Optional[Dict[str, Any]] = None,
) -> None:
    docx = safe_import("docx")
    if not docx:
        raise RuntimeError("Rendering DOCX requires python-docx; install python-docx.")

    from docx import Document  # type: ignore

    doc = Document()
    # Page compaction settings
    page_cfg = template.get("page") or {}
    if page_cfg.get("compact"):
        try:
            sec = doc.sections[0]
            m = float(page_cfg.get("margins_in", 0.5))
            sec.top_margin = Inches(m)
            sec.bottom_margin = Inches(m)
            sec.left_margin = Inches(m)
            sec.right_margin = Inches(m)
        except Exception:
            pass
        try:
            body_pt = float(page_cfg.get("body_pt", 10.5))
            h1_pt = float(page_cfg.get("h1_pt", 12))
            title_pt = float(page_cfg.get("title_pt", 14))
            h1_color = page_cfg.get("h1_color") or page_cfg.get("heading_color")
            h1_bg = page_cfg.get("h1_bg") or page_cfg.get("heading_bg")
            title_color = page_cfg.get("title_color")
            normal = doc.styles["Normal"]
            normal.font.size = Pt(body_pt)
            if "Heading 1" in doc.styles:
                doc.styles["Heading 1"].font.size = Pt(h1_pt)
                doc.styles["Heading 1"].font.bold = True
                rgb = _parse_hex_color(h1_color)
                bg = _parse_hex_color(h1_bg)
                # If no explicit h1_color, auto-contrast against background
                if (not rgb) and bg:
                    rgb = (255, 255, 255) if _is_dark(bg) else (0, 0, 0)
                if rgb:
                    doc.styles["Heading 1"].font.color.rgb = RGBColor(*rgb)
            if "Title" in doc.styles:
                doc.styles["Title"].font.size = Pt(title_pt)
                # keep title bold for prominence
                doc.styles["Title"].font.bold = True
                rgbt = _parse_hex_color(title_color)
                if rgbt:
                    doc.styles["Title"].font.color.rgb = RGBColor(*rgbt)
        except Exception:
            pass
    # Core metadata
    name = data.get("name") or ""
    # Contact fields (accept top-level or under contact)
    contact = data.get("contact") or {}
    headline = data.get("headline") or contact.get("headline") or ""
    email = data.get("email") or contact.get("email") or ""
    phone = data.get("phone") or contact.get("phone") or ""
    display_phone = _format_phone_display(phone) if phone else ""
    location = data.get("location") or contact.get("location") or ""
    website = data.get("website") or contact.get("website") or ""
    linkedin = data.get("linkedin") or contact.get("linkedin") or ""
    github = data.get("github") or contact.get("github") or ""
    links_list = data.get("links") or contact.get("links") or []
    try:
        cp = doc.core_properties
        contact_line = " | ".join([p for p in [email, phone, location] if p])
        cp.title = " - ".join([p for p in [name, contact_line] if p]) or "Resume"
        cp.subject = "Resume"
        if name:
            cp.author = name
        # include phone/location in keywords for easy search
        kw = [k for k in [name, email, phone, location] if k]
        # Optionally include job locations in metadata keywords/category
        try:
            include_exp_locs = bool((template.get("page") or {}).get("metadata_include_locations", True))
        except Exception:
            include_exp_locs = True
        if include_exp_locs:
            locs = []
            for e in (data.get("experience") or []):
                loc_str = str(e.get("location") or "").strip()
                if loc_str:
                    locs.append(loc_str)
            # de-dup preserve order
            uniq_locs = list(dict.fromkeys(locs))
            kw.extend(uniq_locs)
            if uniq_locs:
                try:
                    cp.category = "; ".join(uniq_locs)
                except Exception:
                    pass
        cp.keywords = "; ".join(kw)
    except Exception:
        pass

    # Header
    if name:
        doc.add_heading(name, level=0)
        _tight_paragraph(doc.paragraphs[-1], after_pt=2)
        try:
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = doc.paragraphs[-1].paragraph_format
            pf.left_indent = Pt(0)
            pf.first_line_indent = Pt(0)
        except Exception:
            pass
    # Optional headline line directly under the name
    if headline:
        p_head = doc.add_paragraph(str(headline))
        _tight_paragraph(p_head, after_pt=2)
        try:
            p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p_head.paragraph_format
            pf.left_indent = Pt(0)
            pf.first_line_indent = Pt(0)
        except Exception:
            pass
    # Build compact contact line with auto-included extras
    extras = []
    for val in [website, linkedin, github]:
        if isinstance(val, str) and val.strip():
            extras.append(_format_link_display(val))
    for val in (links_list if isinstance(links_list, list) else []):
        if isinstance(val, str) and val.strip():
            extras.append(_format_link_display(val))
    subtitle_parts = [p for p in [email, display_phone, location] if p]
    if extras:
        subtitle_parts.extend(extras)
    if subtitle_parts:
        p = doc.add_paragraph(" | ".join(subtitle_parts))
        _tight_paragraph(p, after_pt=6)
        try:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.left_indent = Pt(0)
            pf.first_line_indent = Pt(0)
        except Exception:
            pass

    sections = template.get("sections") or []
    keywords = []
    if seed and isinstance(seed.get("keywords"), list):
        keywords = [str(k) for k in seed.get("keywords", [])]

    # Optionally replace order/titles from structure reference
    if structure and isinstance(structure.get("order"), list):
        # Map from key to title in structure; fall back to template
        order_keys: List[str] = structure.get("order", [])
        key_to_title: Dict[str, str] = structure.get("titles", {})
        # rebuild sections preserving configuration from template by key
        tpl_by_key = {s.get("key"): s for s in sections if s.get("key")}
        sections = [
            {**tpl_by_key.get(k, {"key": k, "title": key_to_title.get(k, k.title())})}
            for k in order_keys
            if k in tpl_by_key or key_to_title.get(k)
        ]

    for sec in sections:
        key = sec.get("key")
        title = sec.get("title") or (key.title() if isinstance(key, str) else "")
        if not key:
            continue
        if title:
            doc.add_heading(str(title), level=1)
            _tight_paragraph(doc.paragraphs[-1], before_pt=6, after_pt=2)
            _flush_left(doc.paragraphs[-1])
            # Optional shading for section headings
            try:
                page_h1_bg = (template.get("page") or {}).get("h1_bg") or (template.get("page") or {}).get("heading_bg")
            except Exception:
                page_h1_bg = None
            bg_rgb = _parse_hex_color(page_h1_bg)
            if bg_rgb:
                _apply_paragraph_shading(doc.paragraphs[-1], bg_rgb)

        if key == "summary":
            summary = data.get("summary") or data.get("headline") or ""
            if isinstance(summary, list) and summary:
                items: List[str] = []
                for it in summary:
                    if isinstance(it, dict):
                        s = str(it.get("text") or it.get("line") or it.get("desc") or "").strip()
                        if s:
                            items.append(s)
                    else:
                        s = str(it).strip()
                        if s:
                            items.append(s)
                if items:
                    # Normalize bullets: remove terminal period for fragment style
                    norm_items = [_normalize_bullet_text(it) for it in items]
                    plain, glyph = _use_plain_bullets(sec, page_cfg)
                    _add_bullets(doc, norm_items, sec=sec, keywords=keywords, plain=plain, glyph=glyph, list_style="List Bullet")
            elif isinstance(summary, str) and summary.strip():
                text = summary
                if sec.get("bulleted"):
                    # simple split into bullets by sentence/newline
                    raw_items = [s.strip() for s in text.replace("\n", " ").split(".")]
                    items = [s for s in raw_items if s]
                    # optional cap on sentence count
                    try:
                        max_sent = int(sec.get("max_sentences", 0) or 0)
                    except Exception:
                        max_sent = 0
                    if max_sent and max_sent > 0:
                        items = items[:max_sent]
                    norm_items = [_normalize_bullet_text(it) for it in items]
                    plain, glyph = _use_plain_bullets(sec, page_cfg)
                    _add_bullets(doc, norm_items, sec=sec, keywords=keywords, plain=plain, glyph=glyph, list_style="List Bullet")
                else:
                    p = doc.add_paragraph()
                    _tight_paragraph(p, after_pt=2)
                    if keywords:
                        _bold_keywords(p, text, keywords)
                    else:
                        p.add_run(text)
        elif key == "skills":
            groups = data.get("skills_groups") or []
            skills = [_clean_inline_text(str(s)) for s in (data.get("skills") or [])]
            if groups:
                columns = int(sec.get("columns", 2))
                # Force single-column rendering for reliability
                columns = 1
                compact = bool(sec.get("compact", True))
                as_bullets = bool(sec.get("bullets", False))
                sep = sec.get("separator") or " • "
                max_groups = int(sec.get("max_groups", 999))
                max_items_per_group = int(sec.get("max_items_per_group", 999))
                show_desc = bool(sec.get("show_desc", True))
                for g in (groups[:max_groups]):
                    title = str(g.get("title") or "").strip()
                    raw_items = (g.get("items") or [])
                    norm_items: List[str] = []
                    desc_sep = str(sec.get("desc_separator") or " — ")
                    for x in raw_items:
                        if isinstance(x, dict):
                            name = x.get("name") or x.get("title") or x.get("label") or ""
                            desc = x.get("desc") or x.get("description") or ""
                            s = name.strip()
                            if desc and show_desc:
                                s = f"{s}{desc_sep}{desc.strip()}"
                            norm_items.append(_clean_inline_text(s))
                        else:
                            norm_items.append(_clean_inline_text(str(x)))
                    items = norm_items[:max_items_per_group]
                    if not items:
                        continue
                    if columns > 1:
                        if as_bullets:
                            # Optional group title above the grid
                            if title:
                                _render_group_title(doc, title, sec)
                            # Build a grid and place one item per cell as a bulleted paragraph
                            rows = (len(items) + columns - 1) // columns
                            table = doc.add_table(rows=rows, cols=columns)
                            table.autofit = True
                            for r in range(rows):
                                for c in range(columns):
                                    idx = r + c * rows
                                    cell = table.cell(r, c)
                                    if idx >= len(items):
                                        continue
                                    it = items[idx]
                                    bp = cell.add_paragraph(style="List Bullet")
                                    _tight_paragraph(bp, after_pt=0)
                                    bp.add_run(it)
                        else:
                            # One row, each column contains a compact inline list
                            table = doc.add_table(rows=1, cols=columns)
                            table.autofit = True
                            n = (len(items) + columns - 1) // columns
                            for c in range(columns):
                                start = c * n
                                end = min(start + n, len(items))
                                col_items = items[start:end]
                                cell = table.cell(0, c)
                                if not col_items:
                                    continue
                                text = sep.join(col_items) if compact else "\n".join(col_items)
                                cp = cell.add_paragraph()
                                _tight_paragraph(cp, after_pt=0)
                                cp.add_run(text)
                    else:
                        # One column: either bulleted list or compact inline
                        if as_bullets:
                            if title:
                                _render_group_title(doc, title, sec)
                            plain, glyph = _use_plain_bullets(sec, page_cfg)
                            desc_sep = str(sec.get("desc_separator") or ": ")
                            for it in items:
                                if plain and sec.get("show_desc") and desc_sep in it:
                                    left, right = it.split(desc_sep, 1)
                                    _add_named_bullet(doc, left, right, sec=sec, glyph=glyph, sep=desc_sep)
                                elif plain:
                                    _add_bullet_line(doc, it, sec=sec, glyph=glyph)
                                else:
                                    bp = doc.add_paragraph(style="List Bullet")
                                    _tight_paragraph(bp, after_pt=0)
                                    _compact_bullet(bp)
                                    bp.add_run(it)
                        else:
                            # "Title: item • item" as a single tight paragraph
                            if title:
                                text = f"{title}: {sep.join(items)}" if compact else (title + ":\n" + "\n".join(items))
                            else:
                                text = sep.join(items) if compact else "\n".join(items)
                            p = doc.add_paragraph(text)
                            _tight_paragraph(p, after_pt=0)
            elif skills:
                columns = int(sec.get("columns", 1))
                columns = 1
                as_bullets = bool(sec.get("bullets", False))
                compact = bool(sec.get("compact", True))
                sep = sec.get("separator") or " • "
                max_items = int(sec.get("max_items", 999))
                skills = skills[:max_items]
                if columns > 1:
                    if compact:
                        # One row table, columns contain compact inline lists
                        table = doc.add_table(rows=1, cols=columns)
                        table.autofit = True
                        n = (len(skills) + columns - 1) // columns
                        for c in range(columns):
                            start = c * n
                            end = min(start + n, len(skills))
                            col_items = skills[start:end]
                            cell = table.cell(0, c)
                            if not col_items:
                                continue
                            text = sep.join(col_items)
                            cp = cell.add_paragraph()
                            _tight_paragraph(cp, after_pt=0)
                            cp.add_run(text)
                    else:
                        # build a grid with each item as its own paragraph
                        rows = (len(skills) + columns - 1) // columns
                        table = doc.add_table(rows=rows, cols=columns)
                        table.autofit = True
                        for r in range(rows):
                            for c in range(columns):
                                idx = r + c * rows
                                cell = table.cell(r, c)
                                if idx >= len(skills):
                                    continue
                                item = skills[idx]
                                if as_bullets:
                                    cp = cell.add_paragraph(style="List Bullet")
                                    _tight_paragraph(cp, after_pt=0)
                                    cp.add_run(item)
                                else:
                                    cp = cell.add_paragraph(item)
                                    _tight_paragraph(cp, after_pt=0)
                else:
                    if as_bullets:
                        plain, glyph = _use_plain_bullets(sec, page_cfg)
                        desc_sep = str(sec.get("desc_separator") or ": ")
                        for s in skills:
                            if plain and sec.get("show_desc") and desc_sep in s:
                                left, right = s.split(desc_sep, 1)
                                _add_named_bullet(doc, left, right, sec=sec, glyph=glyph, sep=desc_sep)
                            elif plain:
                                _add_bullet_line(doc, s, sec=sec, glyph=glyph)
                            else:
                                p = doc.add_paragraph(s, style="List Bullet")
                                _tight_paragraph(p, after_pt=0)
                    else:
                        p = doc.add_paragraph(sep.join(skills))
                        _tight_paragraph(p, after_pt=2)
        elif key == "technologies":
            # Build a flat technologies list from either a dedicated list or the Technology group
            tech_items: List[str] = []
            desc_sep = str(sec.get("desc_separator") or ": ")
            # From data.technologies if present
            for t in (data.get("technologies") or []):
                if isinstance(t, dict):
                    nm = t.get("name") or t.get("title") or t.get("label") or ""
                    ds = t.get("desc") or t.get("description") or ""
                    s = nm.strip()
                    if sec.get("show_desc", False) and ds:
                        s = f"{s}{desc_sep}{ds.strip()}"
                    if s:
                        tech_items.append(_clean_inline_text(s))
                else:
                    tech_items.append(_clean_inline_text(str(t)))
            # Or from skills_groups under a Technology/Technologies/Tooling group
            if not tech_items and (data.get("skills_groups") or []):
                for g in data.get("skills_groups") or []:
                    title = str(g.get("title") or "").strip().lower()
                    if title in {"technology", "technologies", "tooling", "tools"}:
                        for x in g.get("items") or []:
                            if isinstance(x, dict):
                                nm = x.get("name") or x.get("title") or x.get("label") or ""
                                ds = x.get("desc") or x.get("description") or ""
                                s = nm.strip()
                                if sec.get("show_desc", False) and ds:
                                    s = f"{s}{desc_sep}{ds.strip()}"
                                if s:
                                    tech_items.append(_clean_inline_text(s))
                            else:
                                tech_items.append(_clean_inline_text(str(x)))
                        break
            # Render technologies similar to skills (no groups)
            if tech_items:
                # optional cap on item count
                try:
                    max_items = int(sec.get("max_items", 0) or 0)
                except Exception:
                    max_items = 0
                if max_items and max_items > 0:
                    tech_items = tech_items[:max_items]
                columns = int(sec.get("columns", 1))
                columns = 1
                as_bullets = bool(sec.get("bullets", True))
                compact = bool(sec.get("compact", True))
                sep = sec.get("separator") or " • "
                # If technologies came from dicts, allow desc separator too
                desc_sep = str(sec.get("desc_separator") or " — ")
                if bool(sec.get("show_desc", False)):
                    # ensure any dict-derived items get the desired separator if not already
                    normed = []
                    for s in tech_items:
                        s2 = _clean_inline_text(str(s))
                        # leave as-is if already contains separator glyphs
                        normed.append(s2)
                    tech_items = normed
                if columns > 1:
                    if as_bullets:
                        rows = (len(tech_items) + columns - 1) // columns
                        table = doc.add_table(rows=rows, cols=columns)
                        table.autofit = True
                        for r in range(rows):
                            for c in range(columns):
                                idx = r + c * rows
                                if idx >= len(tech_items):
                                    continue
                                cell = table.cell(r, c)
                                bp = cell.add_paragraph(style="List Bullet")
                                _tight_paragraph(bp, after_pt=0)
                                bp.add_run(tech_items[idx])
                    else:
                        table = doc.add_table(rows=1, cols=columns)
                        table.autofit = True
                        n = (len(tech_items) + columns - 1) // columns
                        for c in range(columns):
                            start = c * n
                            end = min(start + n, len(tech_items))
                            col_items = tech_items[start:end]
                            cell = table.cell(0, c)
                            if not col_items:
                                continue
                            text = sep.join(col_items)
                            cp = cell.add_paragraph()
                            _tight_paragraph(cp, after_pt=0)
                            cp.add_run(text)
                else:
                    if as_bullets:
                        plain, glyph = _use_plain_bullets(sec, page_cfg)
                        desc_sep = str(sec.get("desc_separator") or ": ")
                        for s in tech_items:
                            if plain and sec.get("show_desc") and desc_sep in s:
                                left, right = s.split(desc_sep, 1)
                                _add_named_bullet(doc, left, right, sec=sec, glyph=glyph, sep=desc_sep)
                            elif plain:
                                _add_bullet_line(doc, s, sec=sec, glyph=glyph)
                            else:
                                p = doc.add_paragraph(s, style="List Bullet")
                                _tight_paragraph(p, after_pt=0)
                    else:
                        p = doc.add_paragraph(sep.join(tech_items))
                        _tight_paragraph(p, after_pt=2)
        elif key == "interests":
            renderer = InterestsSectionRenderer(doc, page_cfg)
            renderer.render(data, sec)
        elif key == "presentations":
            renderer = PresentationsSectionRenderer(doc, page_cfg)
            renderer.render(data, sec)
        elif key == "languages":
            renderer = LanguagesSectionRenderer(doc, page_cfg)
            renderer.render(data, sec)
        elif key == "coursework":
            renderer = CourseworkSectionRenderer(doc, page_cfg)
            renderer.render(data, sec)
        elif key == "certifications":
            renderer = CertificationsSectionRenderer(doc, page_cfg)
            renderer.render(data, sec)
        elif key == "experience":
            items = data.get("experience") or []
            max_items = int(sec.get("max_items", 999))
            inline_dates = bool(sec.get("inline_dates", True))
            role_style = str(sec.get("role_style", "Normal"))
            bullet_style = str(sec.get("bullet_style", "List Bullet"))
            max_bullets = int(sec.get("max_bullets", 999))
            item_color_hex = sec.get("item_color") or sec.get("header_color")
            # Per-recency bullet controls
            recent_roles_count = int(sec.get("recent_roles_count", 0) or 0)
            recent_max_bullets = int(sec.get("recent_max_bullets", max_bullets))
            prior_max_bullets = int(sec.get("prior_max_bullets", max_bullets))
            for e in items[:max_items]:
                header_parts = [p for p in [e.get("title"), e.get("company")] if p]
                header = " at ".join(header_parts) if header_parts else ""
                # Build distinct pieces
                start_txt = str(e.get("start", "") or "")
                end_txt = str(e.get("end", "") or "")
                loc_txt = str(e.get("location", "") or "").strip()
                # Compose date span
                span = ""
                if start_txt and end_txt:
                    span = f"{_normalize_present(start_txt)} – {_normalize_present(end_txt)}"
                elif start_txt and not end_txt:
                    span = f"{_normalize_present(start_txt)} – Present"
                elif end_txt and not start_txt:
                    span = _normalize_present(end_txt)
                if header:
                    # Compose span (duration)
                    span = span
                    # Use shared header line builder for exact consistency
                    _add_header_line(
                        doc,
                        title_text=str(e.get("title") or "").strip(),
                        company_text=str(e.get("company") or "").strip(),
                        loc_text=loc_txt,
                        span_text=span,
                        sec=sec,
                        style=role_style,
                    )
                if (not inline_dates) and date_loc.strip():
                    p_dt = doc.add_paragraph(date_loc)
                    _tight_paragraph(p_dt, after_pt=0)
                plain_bul, glyph = _use_plain_bullets(sec, page_cfg)
                # Determine per-role bullet limit based on recency (assumes most-recent first)
                idx = items.index(e)
                per_role_limit = max_bullets
                if recent_roles_count and idx < recent_roles_count:
                    per_role_limit = min(max_bullets, recent_max_bullets)
                elif recent_roles_count and idx >= recent_roles_count:
                    per_role_limit = min(max_bullets, prior_max_bullets)
                # Collect normalized bullet texts first for efficient rendering
                norm_bullets: List[str] = []
                for b in (e.get("bullets") or [])[:per_role_limit]:
                    if isinstance(b, dict):
                        bt = str(b.get("text") or b.get("line") or b.get("name") or "").strip()
                    else:
                        bt = str(b).strip()
                    if bt:
                        norm_bullets.append(_normalize_bullet_text(bt))
                if norm_bullets:
                    _add_bullets(
                        doc,
                        norm_bullets,
                        sec=sec,
                        keywords=keywords,
                        plain=plain_bul,
                        glyph=glyph,
                        list_style=bullet_style,
                    )
        elif key == "education":
            for ed in (data.get("education") or []):
                degree = str(ed.get("degree") or "").strip()
                institution = str(ed.get("institution") or "").strip()
                year = str(ed.get("year") or "").strip()
                if degree or institution or year:
                    _add_header_line(
                        doc,
                        title_text=degree,
                        company_text=institution,
                        loc_text="",
                        span_text=year,
                        sec=sec,
                        style="Normal",
                    )
        else:
            # Unknown section: noop
            pass

    doc.save(out_path)
