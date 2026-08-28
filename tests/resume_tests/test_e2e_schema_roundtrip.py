"""End-to-end pipeline test for the consolidated resume schema.

Walks the real production path a ``resume render`` invocation takes::

    fixture dict
      -> Resume.from_dict            (schema.py)
      -> align_candidate_to_job      (aligner.py)
      -> FilterPipeline              (pipeline.py: overlays, skills, experience, priority)
      -> write_resume_docx           (docx_writer.py)
      -> re-extract text from the .docx
      -> assert every section that went in comes out

Why the re-extraction step carries the weight
---------------------------------------------
Unit tests assert that a section *types* correctly. They cannot see a section
that types cleanly and then renders to nothing, because nothing in the typed
domain is wrong -- the loss happens in the renderer, one layer down. That exact
bug class appeared repeatedly during this migration: ``label``-keyed entries
rendering blank, alias-keyed items vanishing, an empty scalar summary emitting a
bare heading, ``line``/``priority`` key-salad. Every one of those shipped past a
green unit suite.

So the assertions here compare *rendered prose* against the fixture's own
strings. Asserting the document is non-empty, or that it has the right paragraph
count, would not catch any of the five. Asserting that "Certified Fictional
Operator" appears as text does.

Entry point
-----------
Rendering goes through ``write_resume_docx``, because that is what
``resume.cli.main.cmd_render`` calls and what the golden harness renders with.
``create_resume_writer``/``StandardResumeWriter`` is a second, *differently
behaved* entry point -- it suppresses headings for empty sections and
``write_resume_docx`` does not -- so the empty-section expectations below are
pinned per entry point rather than assumed common. See
``test_empty_section_renders_without_body_through_writer_factory``.

Fixture data
------------
Entirely synthetic. Names, hosts and phone numbers use reserved/example forms
(``example.com``, ``.invalid``, the 555-01xx block), matching the convention in
``golden/golden_fixtures.py``. This file is committed, so a fixture seeded from
real candidate data would publish PII permanently.

Experience entries are **flat**. Multi-role promotion (``roles`` nesting,
``group_id``) is not a shape this schema models -- it was deferred at the design
gate. Two positions at one company are two independent flat entries, which is
what ``two_roles_same_company`` below exercises.

No external binaries. DOCX rendering is pure python-docx; only PDF conversion
would need ``soffice``, and this test never converts.
"""

from __future__ import annotations

import json
import logging
import tempfile
import unittest
from pathlib import Path
from typing import Any

from resume.aligner import align_candidate_to_job
from resume.docx_base import create_resume_writer
from resume.docx_writer import write_resume_docx
from resume.io_utils import read_yaml_or_json
from resume.job import build_keyword_spec
from resume.pipeline import FilterPipeline
from resume.schema import (
    CertificationItem,
    CourseworkItem,
    Education,
    ExperienceEntry,
    NamedLevelItem,
    Presentation,
    PriorityItem,
    Resume,
    SkillGroup,
    SkillGroupItem,
)

from .fixtures import make_education_entry, make_keyword_spec

# --- fixture strings -------------------------------------------------------
# Declared as constants so a round-trip assertion and the fixture that feeds it
# cannot drift apart: asserting a literal that no longer appears in the fixture
# would silently stop testing anything.

SUMMARY_TEXT = "Runs invented systems at an invented scale."
BULLET_RECENT = "Reduced imaginary Kubernetes latency by a fictional margin."
BULLET_PRIOR = "Built a make-believe Terraform deployment pipeline."
COMPANY = "Nonexistent Systems"
TITLE_RECENT = "Staff Engineer"
TITLE_PRIOR = "Senior Engineer"
DEGREE = "BSc Imaginary Computing"
INSTITUTION = "Invented University"
INTEREST = "Fictional cartography"
PRESENTATION = "Scaling Nothing In Particular"
PRESENTATION_EVENT = "MadeUpCon"
LANGUAGE = "Esperanto"
LANGUAGE_LEVEL = "Fluent"
COURSE = "Invented Algorithms"
COURSE_DESC = "Graduate seminar"
CERTIFICATION = "Certified Fictional Operator"
CERTIFICATION_YEAR = "2022"
TECHNOLOGY = "Terraform"
SKILL_GROUP_TITLE = "Platform"
SKILL_GROUP_ITEM = "Kubernetes"
TEACHING_ENTRY = "Intro to Nothing"
UNTYPED_SKILL = "Python"
NAME = "Ada Placeholder"
HEADLINE = "Staff Platform Engineer"
EMAIL = "ada@example.com"
PHONE = "+1-555-0142"
LOCATION = "Fictional City, ZZ"
# Nested under ``contact``; never promoted to ``Resume.website`` -- the
# renderer reads it live from the dict via ``_get_contact_field``.
CONTACT_WEBSITE = "ada.example.invalid"


def make_full_candidate() -> dict[str, Any]:
    """Every section the schema models, in the shapes real data uses.

    Deliberately mixes shapes the way production data does: ``summary`` as a
    list of priority items, ``skills_groups`` items carrying ``priority``, and
    the three untyped sections (``skills``, ``teaching``, ``contact``) in their
    raw form.
    """
    return {
        "name": NAME,
        "headline": HEADLINE,
        "email": EMAIL,
        "phone": PHONE,
        "location": LOCATION,
        "summary": [{"text": SUMMARY_TEXT, "priority": 1.0}],
        "skills_groups": [
            {
                "title": SKILL_GROUP_TITLE,
                "items": [{"name": SKILL_GROUP_ITEM, "priority": 1.0}],
            }
        ],
        "experience": two_roles_same_company(),
        "education": [make_education_entry(degree=DEGREE, institution=INSTITUTION, year="2015")],
        "interests": [{"text": INTEREST, "priority": 1.0}],
        "presentations": [
            {
                "title": PRESENTATION,
                "event": PRESENTATION_EVENT,
                "year": "2024",
                "priority": 1.0,
            }
        ],
        "languages": [{"name": LANGUAGE, "level": LANGUAGE_LEVEL}],
        "coursework": [{"name": COURSE, "desc": COURSE_DESC}],
        "certifications": [{"name": CERTIFICATION, "year": CERTIFICATION_YEAR}],
        "technologies": [{"name": TECHNOLOGY, "priority": 1.0}],
        # Deliberately untyped sections (schema-design.md §1).
        "skills": [UNTYPED_SKILL, "Go"],
        "teaching": [{"name": TEACHING_ENTRY, "desc": "Guest lecture"}],
        "contact": {"website": CONTACT_WEBSITE},
    }


def two_roles_same_company() -> list[dict[str, Any]]:
    """Two positions at one employer, as two FLAT entries.

    This is the closest real data comes to a promotion. There is no ``roles``
    nesting and no ``group_id``: that shape does not exist in this schema, and
    a fixture inventing one would test a contract nobody implements.
    """
    return [
        {
            "title": TITLE_RECENT,
            "company": COMPANY,
            "start": "2021",
            "end": "2025",
            "location": LOCATION,
            "bullets": [{"text": BULLET_RECENT, "priority": 1.0}],
        },
        {
            "title": TITLE_PRIOR,
            "company": COMPANY,
            "start": "2019",
            "end": "2021",
            "location": LOCATION,
            "bullets": [{"text": BULLET_PRIOR, "priority": 1.0}],
        },
    ]


def make_template() -> dict[str, Any]:
    """Single-column template listing every renderable section key."""
    return {
        "page": {"compact": True, "body_pt": 10, "meta_pt": 9},
        "sections": [
            {"key": "summary", "title": "Summary"},
            {"key": "skills", "title": "Skills"},
            {"key": "technologies", "title": "Technologies"},
            {"key": "experience", "title": "Experience"},
            {"key": "education", "title": "Education"},
            {"key": "presentations", "title": "Presentations"},
            {"key": "languages", "title": "Languages"},
            {"key": "coursework", "title": "Coursework"},
            {"key": "certifications", "title": "Certifications"},
            {"key": "interests", "title": "Interests"},
            {"key": "teaching", "title": "Teaching"},
        ],
    }


def make_job_config() -> dict[str, Any]:
    """A job spec whose keywords the fixture genuinely matches.

    The keywords appear in the fixture's *bullets* and untyped ``skills`` list
    on purpose. ``KeywordMatcher.collect_matches_from_candidate`` reads
    ``summary``, ``skills`` and experience bullets -- it does not read
    ``skills_groups`` -- so seeding the keywords only inside the groups would
    match nothing, the skills filter would drop every group, and the pipeline
    assertions would pass vacuously against an empty document.
    """
    return {
        "keywords": {
            "required": [SKILL_GROUP_ITEM],
            "preferred": [TECHNOLOGY],
            "nice": [UNTYPED_SKILL, "Go"],
            "tech_skills": [SKILL_GROUP_ITEM, TECHNOLOGY],
        }
    }


def extract_paragraphs(path: str) -> list[str]:
    """Read the non-blank prose back out of a rendered .docx.

    The round-trip half of the test: what a reader of the finished document
    would actually see, as opposed to what the typed domain believed it held.
    """
    from docx import Document

    return [p.text for p in Document(path).paragraphs if p.text.strip()]


def render_text(resume: Resume, template: dict[str, Any], out_dir: str, tag: str = "resume") -> list[str]:
    """Render through ``write_resume_docx`` -- the path the CLI takes."""
    path = str(Path(out_dir) / f"{tag}.docx")
    write_resume_docx(resume=resume, template=template, out_path=path)
    return extract_paragraphs(path)


def render_text_via_factory(resume: Resume, template: dict[str, Any], out_dir: str) -> list[str]:
    """Render through ``create_resume_writer``, the empty-section-aware path."""
    path = str(Path(out_dir) / "factory.docx")
    create_resume_writer(resume, template).write(path)
    return extract_paragraphs(path)


class E2EPipelineTestBase(unittest.TestCase):
    """Shared temp-dir lifecycle and prose assertion helper."""

    def setUp(self) -> None:  # NOSONAR - required unittest lifecycle method name
        self._tmp = tempfile.TemporaryDirectory()
        self.addCleanup(self._tmp.cleanup)
        self.out_dir = self._tmp.name

    def assertProseIn(self, needle: str, paragraphs: list[str]) -> None:
        """Assert ``needle`` appears in some rendered paragraph.

        Substring rather than equality because renderers legitimately decorate
        a value -- an experience header composes title, company and span into
        one line, and the bullet path strips a terminal period. Equality would
        fail on formatting that is not the thing under test; a bare
        ``assertTrue(any(...))`` would report only "False is not true" and
        leave the reader to guess which section vanished.
        """
        if not any(needle in text for text in paragraphs):
            self.fail(
                f"{needle!r} did not survive the pipeline into the rendered DOCX.\n"
                f"Rendered paragraphs were: {paragraphs}"
            )

    def assertProseNotIn(self, needle: str, paragraphs: list[str]) -> None:
        """Assert no rendered paragraph contains ``needle``."""
        matches = [text for text in paragraphs if needle in text]
        self.assertEqual([], matches, f"{needle!r} unexpectedly rendered")

    def run_pipeline(
        self,
        candidate: dict[str, Any],
        job_cfg: dict[str, Any] | None = None,
        min_priority: float | None = 0.5,
    ) -> tuple[Resume, dict[str, Any]]:
        """Load, align, and filter -- the whole pre-render half of the path.

        Returns the filtered ``Resume`` and the alignment report, so callers
        can assert against either.
        """
        resume = Resume.from_dict(candidate)
        spec, synonyms = build_keyword_spec(job_cfg or make_job_config())
        alignment = align_candidate_to_job(resume.to_dict(), spec, synonyms)

        alignment_path = Path(self.out_dir) / "alignment.json"
        alignment_path.write_text(json.dumps(alignment), encoding="utf-8")

        filtered = (
            FilterPipeline(resume)
            .with_skill_filter(str(alignment_path))
            .with_experience_filter(str(alignment_path))
            .with_priority_filter(min_priority)
            .execute()
        )
        return filtered, alignment


class SchemaLoadTests(E2EPipelineTestBase):
    """Step 2: every section survives ``from_dict`` with the right type."""

    def setUp(self) -> None:  # NOSONAR - required unittest lifecycle method name
        super().setUp()
        self.resume = Resume.from_dict(make_full_candidate())

    def test_scalar_identity_fields_load(self) -> None:
        self.assertEqual(NAME, self.resume.name)
        self.assertEqual(HEADLINE, self.resume.headline)
        self.assertEqual(EMAIL, self.resume.email)
        self.assertEqual(PHONE, self.resume.phone)
        self.assertEqual(LOCATION, self.resume.location)

    def test_typed_list_sections_load_with_element_types(self) -> None:
        """Each typed section is a list of its declared element type."""
        expected: dict[str, type] = {
            "summary": PriorityItem,
            "interests": PriorityItem,
            "skills_groups": SkillGroup,
            "experience": ExperienceEntry,
            "presentations": Presentation,
            "technologies": SkillGroupItem,
            "languages": NamedLevelItem,
            "coursework": CourseworkItem,
            "certifications": CertificationItem,
            "education": Education,
        }
        for section, item_type in expected.items():
            with self.subTest(section=section):
                value = getattr(self.resume, section)
                self.assertIsInstance(value, list)
                self.assertGreater(len(value), 0, f"{section} loaded empty")
                for item in value:
                    self.assertIsInstance(item, item_type)

    def test_nested_items_load_as_typed_items(self) -> None:
        """Nesting is converted too, not left as raw dicts."""
        group = self.resume.skills_groups[0]
        self.assertEqual(SKILL_GROUP_TITLE, group.title)
        self.assertIsInstance(group.items[0], SkillGroupItem)
        self.assertEqual(SKILL_GROUP_ITEM, group.items[0].name)

        bullet = self.resume.experience[0].bullets[0]
        self.assertIsInstance(bullet, PriorityItem)
        self.assertEqual(BULLET_RECENT, bullet.text)

    def test_untyped_sections_pass_through_unconverted(self) -> None:
        """``skills``/``teaching``/``contact`` stay raw by design."""
        self.assertEqual([UNTYPED_SKILL, "Go"], self.resume.skills)
        self.assertEqual([{"name": TEACHING_ENTRY, "desc": "Guest lecture"}], self.resume.teaching)
        self.assertEqual({"website": CONTACT_WEBSITE}, self.resume.contact)

    def test_two_roles_at_one_company_stay_two_flat_entries(self) -> None:
        """No promotion, no nesting: two positions are two entries."""
        self.assertEqual(2, len(self.resume.experience))
        self.assertEqual([TITLE_RECENT, TITLE_PRIOR], [e.title for e in self.resume.experience])
        self.assertEqual([COMPANY, COMPANY], [e.company for e in self.resume.experience])
        for entry in self.resume.experience:
            self.assertFalse(hasattr(entry, "roles"), "flat entries must not gain a roles field")

    def test_roundtrips_key_for_key(self) -> None:
        """``from_dict`` then ``to_dict`` is the identity for canonical input."""
        candidate = make_full_candidate()
        self.assertEqual(candidate, Resume.from_dict(candidate).to_dict())


class AlignmentTests(E2EPipelineTestBase):
    """Step 3: the alignment report is structurally valid."""

    def test_alignment_report_has_expected_structure(self) -> None:
        _, alignment = self.run_pipeline(make_full_candidate())
        self.assertEqual(
            {"matched_keywords", "missing_required", "missing_by_category", "experience_scores"},
            set(alignment),
        )
        self.assertIsInstance(alignment["matched_keywords"], list)
        self.assertIsInstance(alignment["missing_required"], list)
        self.assertIsInstance(alignment["missing_by_category"], dict)
        self.assertIsInstance(alignment["experience_scores"], list)

    def test_matched_keywords_carry_full_metadata(self) -> None:
        _, alignment = self.run_pipeline(make_full_candidate())
        self.assertGreater(len(alignment["matched_keywords"]), 0)
        for match in alignment["matched_keywords"]:
            self.assertEqual(
                {"skill", "count", "weight", "tier", "category"}, set(match)
            )

    def test_alignment_matches_the_fixtures_keywords(self) -> None:
        """The spec's keywords are genuinely found, not vacuously empty."""
        _, alignment = self.run_pipeline(make_full_candidate())
        matched = {m["skill"] for m in alignment["matched_keywords"]}
        self.assertIn(SKILL_GROUP_ITEM, matched)
        self.assertIn(UNTYPED_SKILL, matched)

    def test_experience_scores_cover_every_entry(self) -> None:
        _, alignment = self.run_pipeline(make_full_candidate())
        indices = sorted(index for index, _score in alignment["experience_scores"])
        self.assertEqual([0, 1], indices)


class FilterPipelineTests(E2EPipelineTestBase):
    """Step 4: no section is silently dropped by the filters."""

    def test_no_typed_section_is_dropped_by_the_pipeline(self) -> None:
        filtered, _ = self.run_pipeline(make_full_candidate())
        sections = [
            "summary",
            "skills_groups",
            "experience",
            "education",
            "interests",
            "presentations",
            "languages",
            "coursework",
            "certifications",
            "technologies",
        ]
        for section in sections:
            with self.subTest(section=section):
                self.assertGreater(
                    len(getattr(filtered, section)),
                    0,
                    f"{section} was silently emptied by the filter pipeline",
                )

    def test_untyped_sections_survive_the_pipeline(self) -> None:
        filtered, _ = self.run_pipeline(make_full_candidate())
        self.assertEqual([{"name": TEACHING_ENTRY, "desc": "Guest lecture"}], filtered.teaching)
        self.assertGreater(len(filtered.skills), 0)

    def test_pipeline_preserves_both_roles_at_the_same_company(self) -> None:
        filtered, _ = self.run_pipeline(make_full_candidate())
        self.assertEqual(2, len(filtered.experience))
        self.assertEqual([COMPANY, COMPANY], [e.company for e in filtered.experience])

    def test_pipeline_returns_a_typed_resume(self) -> None:
        filtered, _ = self.run_pipeline(make_full_candidate())
        self.assertIsInstance(filtered, Resume)
        self.assertIsInstance(filtered.experience[0], ExperienceEntry)

    def test_priority_filter_drops_below_threshold_items(self) -> None:
        """The threshold is load-bearing, not a no-op on this fixture."""
        candidate = make_full_candidate()
        candidate["interests"] = [
            {"text": INTEREST, "priority": 1.0},
            {"text": "Low priority pastime", "priority": 0.1},
        ]
        filtered, _ = self.run_pipeline(candidate, min_priority=0.5)
        kept = [item.text for item in filtered.interests]
        self.assertIn(INTEREST, kept)
        self.assertNotIn("Low priority pastime", kept)


class RenderRoundTripTests(E2EPipelineTestBase):
    """Step 6: every section that went in comes back out as prose.

    The load-bearing assertions. A section that types cleanly but renders to
    nothing fails here and nowhere else.
    """

    def setUp(self) -> None:  # NOSONAR - required unittest lifecycle method name
        super().setUp()
        filtered, _ = self.run_pipeline(make_full_candidate())
        self.resume_under_test = filtered
        self.paragraphs = render_text(filtered, make_template(), self.out_dir)

    def test_identity_header_round_trips(self) -> None:
        self.assertProseIn(NAME, self.paragraphs)
        self.assertProseIn(HEADLINE, self.paragraphs)
        self.assertProseIn(EMAIL, self.paragraphs)
        self.assertProseIn(PHONE, self.paragraphs)
        self.assertProseIn(LOCATION, self.paragraphs)

    def test_nested_contact_website_reaches_the_header(self) -> None:
        """The website lives only under ``contact`` and still renders.

        Nothing promotes it: ``Resume.website`` stays ``""`` because
        ``_promote_contact`` fills the identity scalars (name/email/phone/
        location) and not the link fields. The header gets it because
        ``_collect_link_extra_items`` resolves each link field through
        ``get_contact_field``, which falls back to the nested dict.

        That fallback is the whole mechanism, and it is a single lookup away
        from silently returning nothing. Contact handling has already produced
        two bugs in this migration -- a missing ``name`` fallback and a
        non-dict ``contact`` crashing the render -- so the untested path is
        pinned here rather than assumed.
        """
        self.assertEqual("", self.resume_under_test.website)
        self.assertEqual({"website": CONTACT_WEBSITE}, self.resume_under_test.contact)
        self.assertProseIn(CONTACT_WEBSITE, self.paragraphs)

    def test_contact_website_shares_the_contact_line(self) -> None:
        """It is a link extra on the contact line, not a stray paragraph.

        Pins placement as well as presence: a regression that rendered the
        website into its own orphan paragraph would still satisfy a bare
        ``assertProseIn`` while visibly breaking the header.
        """
        lines = [p for p in self.paragraphs if CONTACT_WEBSITE in p]
        self.assertEqual(1, len(lines))
        self.assertIn(EMAIL, lines[0])

    def test_every_section_body_round_trips(self) -> None:
        """The prose of each section, not merely its heading.

        Each string is a value from the fixture. A renderer that emits the
        heading and drops the body fails here -- which is exactly the shape of
        every bug this stage exists to catch.
        """
        expected = {
            "summary": SUMMARY_TEXT.rstrip("."),
            "skills_groups": SKILL_GROUP_ITEM,
            "technologies": TECHNOLOGY,
            "experience_recent_title": TITLE_RECENT,
            "experience_prior_title": TITLE_PRIOR,
            "experience_company": COMPANY,
            "experience_recent_bullet": BULLET_RECENT.rstrip("."),
            "experience_prior_bullet": BULLET_PRIOR.rstrip("."),
            "education_degree": DEGREE,
            "education_institution": INSTITUTION,
            "presentations": PRESENTATION,
            "presentations_event": PRESENTATION_EVENT,
            "languages": LANGUAGE,
            "languages_level": LANGUAGE_LEVEL,
            "coursework": COURSE,
            "coursework_desc": COURSE_DESC,
            "certifications": CERTIFICATION,
            "certifications_year": CERTIFICATION_YEAR,
            "interests": INTEREST,
            "teaching": TEACHING_ENTRY,
        }
        for section, prose in expected.items():
            with self.subTest(section=section):
                self.assertProseIn(prose, self.paragraphs)

    def test_both_roles_at_one_company_render_separately(self) -> None:
        """Two flat entries produce two distinct headers, not one merged line."""
        recent = [p for p in self.paragraphs if TITLE_RECENT in p]
        prior = [p for p in self.paragraphs if TITLE_PRIOR in p]
        self.assertEqual(1, len(recent))
        self.assertEqual(1, len(prior))
        self.assertNotEqual(recent[0], prior[0])

    def test_date_spans_round_trip(self) -> None:
        joined = "\n".join(self.paragraphs)
        for year in ("2021", "2025", "2019"):
            with self.subTest(year=year):
                self.assertIn(year, joined)

    def test_section_headings_render(self) -> None:
        for heading in ("Summary", "Skills", "Experience", "Education", "Certifications"):
            with self.subTest(heading=heading):
                self.assertProseIn(heading, self.paragraphs)


class EmptyAndMissingSectionTests(E2EPipelineTestBase):
    """Sad paths: degradation, not exceptions. Validation is advisory."""

    def test_empty_section_renders_without_crashing(self) -> None:
        """An empty section must not break the render."""
        candidate = {"name": "Empty Placeholder", "summary": SUMMARY_TEXT, "interests": []}
        resume = Resume.from_dict(candidate)
        paragraphs = render_text(resume, make_template(), self.out_dir)
        self.assertProseIn("Empty Placeholder", paragraphs)
        self.assertProseIn(SUMMARY_TEXT.rstrip("."), paragraphs)

    def test_empty_section_renders_without_body_through_writer_factory(self) -> None:
        """``create_resume_writer`` suppresses the heading of an empty section.

        Pinned against this entry point specifically. ``write_resume_docx``
        emits the heading unconditionally -- the two paths differ, and this
        asserts the guard where the guard actually lives
        (``docx_standard._section_has_data``).
        """
        candidate = {"name": "Empty Placeholder", "summary": SUMMARY_TEXT, "interests": []}
        resume = Resume.from_dict(candidate)
        paragraphs = render_text_via_factory(resume, make_template(), self.out_dir)
        self.assertProseIn(SUMMARY_TEXT.rstrip("."), paragraphs)
        self.assertProseNotIn("Interests", paragraphs)

    def test_empty_scalar_summary_emits_no_summary_heading(self) -> None:
        """The regression that shipped once: a bare heading over no body.

        An empty scalar ``summary`` normalizes to ``[PriorityItem(text="")]``,
        which is truthy where ``""`` was falsy. ``_data_key_is_truthy``
        special-cases the scalar origin so the section stays suppressed.
        """
        resume = Resume.from_dict({"name": "Empty Placeholder", "summary": "", "headline": ""})
        paragraphs = render_text_via_factory(resume, make_template(), self.out_dir)
        self.assertProseIn("Empty Placeholder", paragraphs)
        self.assertProseNotIn("Summary", paragraphs)

    def test_all_optional_sections_absent_renders_minimal_document(self) -> None:
        """Name only: a minimal document, no crash."""
        resume = Resume.from_dict({"name": "Minimal Placeholder"})
        paragraphs = render_text_via_factory(resume, make_template(), self.out_dir)
        self.assertEqual(["Minimal Placeholder"], paragraphs)

    def test_all_optional_sections_absent_survives_the_full_pipeline(self) -> None:
        """The same minimal input through align + filter + render."""
        filtered, alignment = self.run_pipeline({"name": "Minimal Placeholder"})
        self.assertEqual([], alignment["matched_keywords"])
        paragraphs = render_text(filtered, make_template(), self.out_dir)
        self.assertProseIn("Minimal Placeholder", paragraphs)

    def test_experience_entry_missing_title_still_renders_via_company(self) -> None:
        """Documented fallback: the header renders on company alone.

        ``ExperienceSectionRenderer._render_experience_entry`` gates the header
        on ``title or company``, so a title-less entry keeps its company, span
        and bullets rather than disappearing.
        """
        candidate = {
            "name": "Untitled Placeholder",
            "experience": [
                {
                    "company": COMPANY,
                    "start": "2020",
                    "end": "2024",
                    "bullets": [{"text": BULLET_RECENT, "priority": 1.0}],
                }
            ],
        }
        resume = Resume.from_dict(candidate)
        self.assertEqual("", resume.experience[0].title)

        paragraphs = render_text(resume, make_template(), self.out_dir)
        self.assertProseIn(COMPANY, paragraphs)
        self.assertProseIn(BULLET_RECENT.rstrip("."), paragraphs)


class InvalidShapeTests(E2EPipelineTestBase):
    """Wrong-shaped input degrades with a warning; it never raises."""

    def test_invalid_section_shape_falls_back_and_warns_naming_the_key(self) -> None:
        """A wrong-typed section becomes ``[]`` and names itself in the warning."""
        candidate = {
            "name": "Malformed Placeholder",
            "summary": SUMMARY_TEXT,
            "experience": "not-a-list",
            "interests": {"unexpected": "mapping"},
        }
        with self.assertLogs("resume.schema", level=logging.WARNING) as captured:
            resume = Resume.from_dict(candidate)

        self.assertEqual([], resume.experience)
        self.assertEqual([], resume.interests)

        messages = "\n".join(captured.output)
        self.assertIn("Resume.experience", messages)
        self.assertIn("Resume.interests", messages)

    def test_invalid_section_shape_still_renders_the_rest(self) -> None:
        """The offending section is lost; every other section still renders."""
        candidate = {
            "name": "Malformed Placeholder",
            "summary": SUMMARY_TEXT,
            "experience": "not-a-list",
            "education": [make_education_entry(degree=DEGREE, institution=INSTITUTION)],
        }
        with self.assertLogs("resume.schema", level=logging.WARNING):
            resume = Resume.from_dict(candidate)

        paragraphs = render_text(resume, make_template(), self.out_dir)
        self.assertProseIn("Malformed Placeholder", paragraphs)
        self.assertProseIn(SUMMARY_TEXT.rstrip("."), paragraphs)
        self.assertProseIn(DEGREE, paragraphs)

    def test_rejects_invalid_top_level_document(self) -> None:
        """A non-mapping *document* raises rather than falling back.

        This is the one carve-out from the advisory policy the tests above
        exercise. A wrong-typed argument is a programming error at an API
        boundary, not malformed candidate data: there is no document to
        salvage. It previously returned an empty ``Resume``, which rendered a
        blank document with nothing reporting a failure.
        """
        with self.assertRaises(TypeError) as ctx:
            Resume.from_dict(["not", "a", "mapping"])

        self.assertIn("list", str(ctx.exception))
        self.assertIn("mapping", str(ctx.exception))

    def test_invalid_list_item_shape_degrades_to_defaults(self) -> None:
        """A non-dict item inside a typed list becomes a default item."""
        with self.assertLogs("resume.schema", level=logging.WARNING):
            resume = Resume.from_dict({"name": "Malformed Placeholder", "education": [42]})

        self.assertEqual(1, len(resume.education))
        self.assertIsInstance(resume.education[0], Education)
        self.assertEqual("", resume.education[0].degree)

    def test_valid_section_shape_emits_no_warning(self) -> None:
        """Happy-path pair: clean input must stay silent.

        Without this, a change that warned on *everything* would still pass
        every assertion above.
        """
        logger = logging.getLogger("resume.schema")
        with self.assertNoLogs(logger, level=logging.WARNING):
            resume = Resume.from_dict(make_full_candidate())
        self.assertEqual(2, len(resume.experience))


class CorruptedCandidateJsonTests(E2EPipelineTestBase):
    """The one genuine error path: the JSON parser, not schema validation.

    Schema validation is advisory and never raises. A truncated file is a
    different failure entirely -- it never reaches the schema, because
    ``json.loads`` cannot produce a document at all. Conflating the two would
    assert that malformed *content* raises, which it must not.
    """

    def _write(self, name: str, text: str) -> str:
        path = Path(self.out_dir) / name
        path.write_text(text, encoding="utf-8")
        return str(path)

    def test_rejects_truncated_candidate_json_with_a_clear_error(self) -> None:
        path = self._write("truncated.json", f'{{"name": "{NAME}", "experience": [{{"title": "Eng"')

        with self.assertRaises(json.JSONDecodeError) as ctx:
            read_yaml_or_json(path)

        error = ctx.exception
        self.assertIsInstance(error, ValueError)
        # A clear error locates the fault, so the file can actually be fixed.
        # Only the structured attributes are asserted: the message text itself
        # is a CPython implementation detail and varies across versions.
        self.assertGreater(error.pos, 0)
        self.assertEqual(1, error.lineno)
        self.assertGreater(error.colno, 0)

    def test_rejects_empty_candidate_json(self) -> None:
        path = self._write("empty.json", "")
        with self.assertRaises(json.JSONDecodeError):
            read_yaml_or_json(path)

    def test_rejects_missing_candidate_file(self) -> None:
        with self.assertRaises(FileNotFoundError):
            read_yaml_or_json(str(Path(self.out_dir) / "absent.json"))

    def test_reads_well_formed_candidate_json(self) -> None:
        """Happy-path pair: the reader loads a good file end to end."""
        candidate = make_full_candidate()
        path = self._write("good.json", json.dumps(candidate))

        loaded = read_yaml_or_json(path)
        self.assertEqual(candidate, loaded)

        resume = Resume.from_dict(loaded)
        paragraphs = render_text(resume, make_template(), self.out_dir)
        self.assertProseIn(NAME, paragraphs)
        self.assertProseIn(CERTIFICATION, paragraphs)


class FixtureIntegrityTests(unittest.TestCase):
    """The fixture itself must keep exercising what it claims to."""

    def test_fixture_covers_every_typed_section_the_schema_models(self) -> None:
        """Guards against a section being added to the schema and missed here."""
        candidate = make_full_candidate()
        typed_sections = {
            "summary",
            "skills_groups",
            "experience",
            "interests",
            "presentations",
            "technologies",
            "languages",
            "coursework",
            "certifications",
            "education",
        }
        missing = typed_sections - set(candidate)
        self.assertEqual(set(), missing, f"fixture is missing sections: {missing}")

    def test_fixture_covers_the_untyped_sections(self) -> None:
        candidate = make_full_candidate()
        for section in ("skills", "teaching", "contact"):
            with self.subTest(section=section):
                self.assertIn(section, candidate)

    def test_template_lists_every_renderable_section(self) -> None:
        """Every registered renderer key appears in the fixture template."""
        from resume.docx_standard import SECTION_RENDERERS

        template_keys = {section["key"] for section in make_template()["sections"]}
        self.assertEqual(set(SECTION_RENDERERS), template_keys)

    def test_keyword_spec_helper_builds_the_expected_shape(self) -> None:
        """Reuses the suite's own spec factory rather than hand-rolling one."""
        spec = make_keyword_spec(required=[SKILL_GROUP_ITEM], preferred=[TECHNOLOGY])
        self.assertEqual([SKILL_GROUP_ITEM], spec["required"])
        self.assertEqual([TECHNOLOGY], spec["preferred"])


if __name__ == "__main__":
    unittest.main()
