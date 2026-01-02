import tempfile
import unittest
import json
import subprocess
from pathlib import Path


CLI = ["python", "-m", "resume"]


SAMPLE_LINKEDIN = """John Doe
Contact
john@example.com | (555) 123-4567 | San Francisco, CA

Summary
Engineer with experience in distributed systems and Python.

Skills
Python, Distributed Systems, AWS; Docker; Kubernetes

Experience
Senior Software Engineer at ExampleCorp (2019-2023) - San Francisco, CA
- Built microservices in Python and Go
- Deployed on Kubernetes and AWS EKS

Education
BS Computer Science, University of Somewhere, 2015
"""


SAMPLE_RESUME = """John Doe
john@example.com | (555) 123-4567 | San Francisco, CA

Profile
Seasoned engineer with focus on reliability and scale.

Technical Skills
Python, Go, Kubernetes, AWS, Terraform

Work History
Staff Engineer at Foo Inc (2023-Now) - Remote
- Led migration to Kubernetes
- Improved reliability SLOs by 20%

Education
MS Computer Science, Institute of Tech, 2017
"""


class TestCLI(unittest.TestCase):
    def test_help(self):
        out = subprocess.run(CLI + ["-h"], capture_output=True, text=True)  # nosec B603 - test code with literal args
        self.assertEqual(out.returncode, 0)
        self.assertIn("resume-assistant", out.stdout)

    def test_extract_and_summarize(self):
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            li = td / "linkedin.txt"
            rs = td / "resume.txt"
            data_out = td / "data.json"
            summary_out = td / "summary.md"
            li.write_text(SAMPLE_LINKEDIN, encoding="utf-8")
            rs.write_text(SAMPLE_RESUME, encoding="utf-8")

            # extract
            r = subprocess.run(CLI + ["extract", "--linkedin", str(li), "--resume", str(rs), "--out", str(data_out)])  # nosec B603 - test code with temp files
            self.assertEqual(r.returncode, 0)
            self.assertTrue(data_out.exists(), "data.json should be created")
            data = json.loads(data_out.read_text(encoding="utf-8"))
            self.assertIn("experience", data)
            self.assertIn("skills", data)

            # summarize
            r = subprocess.run(CLI + [  # nosec B603 - test code with temp files
                "summarize", "--data", str(data_out), "--seed", "keywords=Python Kubernetes AWS", "--out", str(summary_out)
            ])
            self.assertEqual(r.returncode, 0)
            self.assertTrue(summary_out.exists(), "summary.md should be created")
            self.assertIn("Resume Summary", summary_out.read_text(encoding="utf-8"))

    def test_render_docx_skippable(self):
        # Skip test if python-docx not available
        try:
            import docx  # noqa: F401
        except Exception:
            self.skipTest("python-docx not installed")
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            data = td / "data.json"
            template = td / "template.yaml"
            out = td / "out.docx"
            data.write_text(json.dumps({
                "name": "Jane Doe",
                "email": "jane@example.com",
                "phone": "(555) 999-1234",
                "location": "Austin, TX",
                "summary": "Backend engineer with Go and AWS.",
                "skills": ["Go", "AWS", "Docker"],
                "experience": [{
                    "title": "Backend Engineer", "company": "Acme", "start": "2020", "end": "2024", "location": "Remote",
                    "bullets": ["Built REST APIs", "Managed AWS infrastructure"]
                }],
                "education": [{"degree": "BS CS", "institution": "Uni", "year": "2019"}]
            }), encoding="utf-8")
            template.write_text("""
sections:
  - key: summary
    title: Professional Summary
  - key: skills
    title: Core Skills
  - key: experience
    title: Experience
    max_items: 5
  - key: education
    title: Education
""".strip(), encoding="utf-8")

            r = subprocess.run(CLI + [  # nosec B603 - test code with temp files
                "render", "--data", str(data), "--template", str(template), "--seed", "keywords=AWS Go", "--out", str(out)
            ])
            self.assertEqual(r.returncode, 0)
            self.assertTrue(out.exists(), ".docx should be created")
            self.assertGreater(out.stat().st_size, 0)


    def test_structure_inference_skippable(self):
        try:
            from docx import Document  # type: ignore
        except Exception:
            self.skipTest("python-docx not installed")
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            ref = td / "ref.docx"
            out = td / "struct.json"
            # Create reference docx
            doc = Document()
            doc.add_heading("Profile", level=1)
            doc.add_paragraph("...")
            doc.add_heading("Work History", level=1)
            doc.add_paragraph("...")
            doc.add_heading("Technical Skills", level=1)
            doc.add_paragraph("...")
            doc.add_heading("Education", level=1)
            doc.add_paragraph("...")
            doc.save(str(ref))

            r = subprocess.run(CLI + ["structure", "--source", str(ref), "--out", str(out)])  # nosec B603 - test code with temp files
            self.assertEqual(r.returncode, 0)
            self.assertTrue(out.exists())
            data = json.loads(out.read_text(encoding="utf-8"))
            self.assertIn("order", data)
            self.assertIn("experience", data.get("order", []))

    def test_align_and_tailor(self):
        with tempfile.TemporaryDirectory() as td:
            td = Path(td)
            # candidate data
            cand = td / "cand.json"
            cand.write_text(json.dumps({
                "name": "Alex Test",
                "skills": ["Python", "AWS", "Docker", "Kubernetes", "Grafana"],
                "experience": [
                    {"title": "SRE", "company": "Acme", "start": "2020", "end": "2024", "location": "Remote",
                     "bullets": ["Managed Kubernetes clusters", "Built dashboards in Grafana", "Automated deployments with Docker"]},
                    {"title": "DevOps", "company": "Other", "start": "2018", "end": "2020", "location": "Remote",
                     "bullets": ["Provisioned AWS infrastructure", "Wrote Python tooling"]}
                ]
            }), encoding="utf-8")
            # job config
            job = td / "job.yaml"
            job.write_text("""
title: Senior SRE
company: ExampleCo
keywords:
  required:
    - skill: Kubernetes
      weight: 3
    - skill: AWS
      weight: 3
  preferred:
    - skill: Grafana
      weight: 2
  soft_skills:
    - Communication
    - Collaboration
  tech_skills:
    - Python
    - Docker
  synonyms:
    Kubernetes: [k8s, EKS]
""".strip(), encoding="utf-8")
            align_out = td / "align.json"
            tailored = td / "tailored.json"

            r = subprocess.run(CLI + [  # nosec B603 - test code with temp files
                "align", "--data", str(cand), "--job", str(job), "--out", str(align_out), "--tailored", str(tailored), "--max-bullets", "2"
            ])
            self.assertEqual(r.returncode, 0)
            a = json.loads(align_out.read_text(encoding="utf-8"))
            self.assertIn("matched_keywords", a)
            mk = {x["skill"] for x in a["matched_keywords"]}
            self.assertIn("Kubernetes", mk)
            self.assertIn("AWS", mk)
            # category presence
            self.assertIn("missing_by_category", a)
            # Tailored output
            t = json.loads(tailored.read_text(encoding="utf-8"))
            self.assertIn("experience", t)
            # Only bullets with matched keywords or first bullet retained
            self.assertTrue(any("Kubernetes" in b or "Grafana" in b or "AWS" in b for e in t["experience"] for b in e.get("bullets", [])))


if __name__ == "__main__":
    unittest.main(verbosity=2)
