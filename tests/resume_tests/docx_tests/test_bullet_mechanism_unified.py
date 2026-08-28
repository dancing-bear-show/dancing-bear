"""The standard layout emits bullets through exactly ONE mechanism.

WHY THIS FILE IS CROSS-SECTION, NOT PER-SECTION
    Three unrelated bullet mechanisms coexisted in this layout for a long time
    and every per-section test passed the whole while, because each one only
    ever asserted against the mechanism its own section happened to use::

        Summary      style="List Bullet"  left_indent=0     literal glyph: no
        Skills       style="Normal"       left_indent=None  literal glyph: yes
        Experience   style="List Bullet"  left_indent=0     literal glyph: no
        Education    style="Normal"       left_indent=0     (no bullet at all)

    ``List Bullet`` draws its glyph and its indent from a Word numbering
    definition, which is a different system from a paragraph that prints a
    literal glyph and sets its own indent. Sections using one could not be
    aligned with sections using the other at any indent setting. A test suite
    made only of per-section assertions cannot see that, because the defect is
    precisely that the sections disagree with *each other*.

    So the assertions below gather every bulleted paragraph the standard layout
    produces, ACROSS all sections, and require them to share one style and one
    left indent. That is the regression guard. Adding a fourth mechanism, or
    reverting any section to a Word list style, fails here.
"""

from __future__ import annotations

import os
import tempfile
import unittest

from resume.docx_writer import write_resume_docx
from resume.schema import Resume

BULLET_GLYPH = "•"

# Exercises summary, skills, experience and education in one document, plus the
# simple-list sections, so a single render covers every bullet-emitting site.
# Synthetic throughout -- see golden_fixtures for why fixtures never derive from
# a real profile.
_FIXTURE = {
    "name": "Uniform Placeholder",
    "headline": "Placeholder Engineer",
    "email": "uniform@example.com",
    "phone": "+1-555-0111",
    "location": "Uniform City, ZZ",
    "summary": [
        {"text": "First summary bullet."},
        {"text": "Second summary bullet."},
    ],
    "skills_groups": [
        {"title": "Platform", "items": ["Kubernetes", "Terraform"]},
        {"title": "Languages", "items": ["Python", "Go"]},
    ],
    "experience": [
        {
            "title": "Placeholder Engineer",
            "company": "Uniform Corp",
            "location": "Uniform City, ZZ",
            "start": "2020",
            "end": "2024",
            "bullets": ["Experience bullet one.", "Experience bullet two."],
        },
        {
            "title": "Junior Placeholder",
            "company": "Earlier Uniform Ltd",
            "start": "2017",
            "end": "2020",
            "bullets": ["An earlier bullet."],
        },
    ],
    "education": [
        {
            "degree": "BSc Placeholder Studies",
            "institution": "Uniform University",
            "year": "2015",
        }
    ],
    "presentations": [
        {"title": "A Placeholder Talk", "event": "Uniform Conf", "year": 2019}
    ],
    "teaching": ["Intro to Placeholding (Uniform University)"],
    "certifications": [{"name": "Certified Placeholder", "year": 2021}],
    "languages": [{"name": "English", "level": "Native"}],
    "interests": ["Placeholding"],
}

_SECTIONS = [
    {"key": "summary", "title": "Summary"},
    {"key": "skills", "title": "Skills"},
    {"key": "experience", "title": "Experience"},
    {"key": "education", "title": "Education"},
    {"key": "presentations", "title": "Presentations"},
    {"key": "teaching", "title": "Teaching"},
    {"key": "certifications", "title": "Certifications"},
    {"key": "languages", "title": "Languages"},
    {"key": "interests", "title": "Interests"},
]


def _standard_template() -> dict:
    return {
        "page": {"compact": True, "body_pt": 10, "meta_pt": 9},
        "sections": [dict(s) for s in _SECTIONS],
    }


class _Paragraph:
    """One rendered paragraph, reduced to what the bullet contract cares about."""

    __slots__ = ("section", "style", "left_indent", "text")

    def __init__(self, section: str, style: str, left_indent, text: str) -> None:
        self.section = section
        self.style = style
        self.left_indent = left_indent
        self.text = text

    @property
    def is_bulleted(self) -> bool:
        return self.text.strip().startswith(BULLET_GLYPH)

    @property
    def triple(self) -> tuple[str, object, bool]:
        return (self.style, self.left_indent, self.is_bulleted)

    def __repr__(self) -> str:
        return (
            f"<{self.section}: style={self.style!r} "
            f"left_indent={self.left_indent!r} text={self.text[:40]!r}>"
        )


def _render_standard(template: dict | None = None) -> list[_Paragraph]:
    """Render the fixture through the standard layout and flatten it."""
    import docx

    with tempfile.TemporaryDirectory() as tmp:
        out = os.path.join(tmp, "resume.docx")
        write_resume_docx(
            Resume.from_dict(_FIXTURE), template or _standard_template(), out
        )
        doc = docx.Document(out)
        section = "<header>"
        paragraphs: list[_Paragraph] = []
        for p in doc.paragraphs:
            style = getattr(p.style, "name", str(p.style))
            if style.startswith("Heading") or style == "Title":
                section = p.text or style
                continue
            if not p.text.strip():
                continue
            indent = p.paragraph_format.left_indent
            paragraphs.append(
                _Paragraph(
                    section, style, None if indent is None else indent.pt, p.text
                )
            )
        return paragraphs


class StandardLayoutBulletMechanismTests(unittest.TestCase):
    """One bullet mechanism, asserted across every section at once."""

    maxDiff = None

    @classmethod
    def setUpClass(cls) -> None:
        cls.paragraphs = _render_standard()
        cls.bulleted = [p for p in cls.paragraphs if p.is_bulleted]

    def test_fixture_produces_bullets_in_several_sections(self):
        """Guard the guard: a fixture that stopped bulleting would pass vacuously.

        Every assertion below is over the set of bulleted paragraphs, so an
        empty set satisfies all of them. This pins that the set is non-empty
        AND spans more than one section, which is what makes the cross-section
        comparisons meaningful.
        """
        self.assertTrue(self.bulleted, "fixture rendered no bulleted paragraphs")
        sections = {p.section for p in self.bulleted}
        self.assertGreaterEqual(
            len(sections), 3, f"bullets confined to too few sections: {sections}"
        )

    def test_every_bulleted_paragraph_shares_one_style_and_indent(self):
        """THE regression guard: all sections agree on the bullet mechanism.

        Compared across sections, not within them. Three mechanisms coexisted
        for a long time precisely because nothing ever compared them.
        """
        triples = {p.triple for p in self.bulleted}
        self.assertEqual(
            len(triples),
            1,
            "standard layout emits bullets through more than one mechanism: "
            + repr(sorted((p.section, p.style, p.left_indent) for p in self.bulleted)),
        )
        style, indent, _ = next(iter(triples))
        self.assertEqual(style, "Normal")
        self.assertEqual(indent, 0)

    def test_no_bulleted_paragraph_uses_a_word_list_style(self):
        """``List Bullet`` carries a numbering indent that ``flush_left`` cannot reset."""
        offenders = [p for p in self.paragraphs if p.style.startswith("List ")]
        self.assertEqual(offenders, [], f"Word list styles present: {offenders}")

    def test_no_paragraph_contains_more_than_one_glyph(self):
        """A second glyph inside a paragraph means items were joined, not bulleted."""
        offenders = [p for p in self.paragraphs if p.text.count(BULLET_GLYPH) > 1]
        self.assertEqual(offenders, [], f"paragraphs with nested glyphs: {offenders}")

    def test_every_bulleted_paragraph_starts_with_the_glyph_and_a_space(self):
        """The glyph is a literal run prefix, not a style-supplied decoration."""
        for p in self.bulleted:
            self.assertTrue(
                p.text.startswith(f"{BULLET_GLYPH} "),
                f"bullet does not start with a glyph run: {p!r}",
            )

    def test_education_renders_without_bullets(self):
        """Education is a plain paragraph, matching the reference document.

        The reference this layout is styled after renders its Education section
        as plain ``Normal`` paragraphs and puts no glyph on them, so unifying
        the bullet mechanism deliberately does NOT give this section a bullet
        it never had.
        """
        education = [p for p in self.paragraphs if p.section == "Education"]
        self.assertTrue(education, "fixture rendered no Education paragraphs")
        for p in education:
            self.assertFalse(p.is_bulleted, f"Education gained a bullet: {p!r}")

    def test_obsolete_style_config_cannot_reintroduce_a_second_mechanism(self):
        """Templates in the wild still carry the keys that used to switch mechanism.

        ``plain_bullets``, ``bullets.style`` and ``bullet_style`` selected a
        Word list style before unification. They must now be inert: a template
        setting all three still renders through the one mechanism.
        """
        template = _standard_template()
        for sec in template["sections"]:
            sec["plain_bullets"] = False
            sec["bullets"] = {"style": "list"}
            sec["bullet_style"] = "List Bullet"

        paragraphs = _render_standard(template)
        bulleted = [p for p in paragraphs if p.is_bulleted]
        self.assertTrue(bulleted)
        self.assertEqual({p.triple for p in bulleted}, {("Normal", 0, True)})
        self.assertEqual([p for p in paragraphs if p.style.startswith("List ")], [])

    def test_a_configured_glyph_applies_uniformly_across_sections(self):
        """Changing the glyph changes it everywhere, because there is one emitter."""
        template = _standard_template()
        template["page"]["bullets"] = {"glyph": "▸"}

        paragraphs = _render_standard(template)
        bulleted = [p for p in paragraphs if p.text.strip().startswith("▸")]
        self.assertTrue(bulleted)
        self.assertEqual({p.triple[:2] for p in bulleted}, {("Normal", 0)})
        self.assertEqual(
            [p for p in paragraphs if p.text.strip().startswith(BULLET_GLYPH)],
            [],
            "default glyph still emitted somewhere after the page config changed it",
        )


if __name__ == "__main__":
    unittest.main()
