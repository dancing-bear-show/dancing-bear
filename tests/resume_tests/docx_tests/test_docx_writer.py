"""Tests for resume/docx_writer.py DOCX rendering utilities."""

from __future__ import annotations

from tests.fixtures import test_path
from tests.resume_tests.fixtures import make_docx_paragraph_side_effect
import unittest
from unittest.mock import MagicMock, patch


class TestMatchSectionKey(unittest.TestCase):
    """Tests for _match_section_key function."""

    def test_matches_summary_variants(self):
        from resume.docx_writer import _match_section_key
        self.assertEqual(_match_section_key("Summary"), "summary")
        self.assertEqual(_match_section_key("profile"), "summary")
        self.assertEqual(_match_section_key("ABOUT"), "summary")

    def test_matches_skills_variants(self):
        from resume.docx_writer import _match_section_key
        self.assertEqual(_match_section_key("Skills"), "skills")
        self.assertEqual(_match_section_key("TECHNICAL SKILLS"), "skills")

    def test_matches_technologies(self):
        from resume.docx_writer import _match_section_key
        self.assertEqual(_match_section_key("Technologies"), "technologies")
        self.assertEqual(_match_section_key("technology"), "technologies")
        self.assertEqual(_match_section_key("tools"), "technologies")

    def test_matches_experience_variants(self):
        from resume.docx_writer import _match_section_key
        self.assertEqual(_match_section_key("Experience"), "experience")
        self.assertEqual(_match_section_key("work history"), "experience")
        self.assertEqual(_match_section_key("employment"), "experience")

    def test_matches_education_variants(self):
        from resume.docx_writer import _match_section_key
        self.assertEqual(_match_section_key("Education"), "education")
        self.assertEqual(_match_section_key("academics"), "education")

    def test_returns_none_for_unknown(self):
        from resume.docx_writer import _match_section_key
        self.assertIsNone(_match_section_key("Unknown Section"))
        self.assertIsNone(_match_section_key("Certifications"))

    def test_strips_whitespace(self):
        from resume.docx_writer import _match_section_key
        self.assertEqual(_match_section_key("  Summary  "), "summary")


class TestGetHeaderLevel(unittest.TestCase):
    """Tests for _get_header_level function."""

    def test_returns_from_sec_if_present(self):
        from resume.docx_writer import _get_header_level
        sec = {"header_level": 2}
        result = _get_header_level(sec, None)
        self.assertEqual(result, 2)

    def test_returns_from_page_cfg_if_sec_missing(self):
        from resume.docx_writer import _get_header_level
        page_cfg = {"header_level": 3}
        result = _get_header_level(None, page_cfg)
        self.assertEqual(result, 3)

    def test_prefers_sec_over_page_cfg(self):
        from resume.docx_writer import _get_header_level
        sec = {"header_level": 2}
        page_cfg = {"header_level": 3}
        result = _get_header_level(sec, page_cfg)
        self.assertEqual(result, 2)

    def test_defaults_to_1_if_missing(self):
        from resume.docx_writer import _get_header_level
        result = _get_header_level(None, None)
        self.assertEqual(result, 1)

    def test_handles_invalid_header_level(self):
        from resume.docx_writer import _get_header_level
        sec = {"header_level": "not-a-number"}
        result = _get_header_level(sec, None)
        self.assertEqual(result, 1)


class TestExtractExperienceLocations(unittest.TestCase):
    """Tests for _extract_experience_locations function."""

    def test_extracts_unique_locations(self):
        from resume.docx_writer import _extract_experience_locations
        from resume.schema import Resume
        data = {
            "experience": [
                {"title": "Engineer", "location": "Seattle, WA"},
                {"title": "Manager", "location": "Portland, OR"},
                {"title": "Director", "location": "Seattle, WA"},  # Duplicate
            ]
        }
        result = _extract_experience_locations(Resume.from_dict(data))
        self.assertEqual(result, ["Seattle, WA", "Portland, OR"])

    def test_handles_empty_locations(self):
        from resume.docx_writer import _extract_experience_locations
        from resume.schema import Resume
        # `location` is declared str, but an explicit null in the source file
        # survives as None -- the helper must coerce rather than crash.
        data = {
            "experience": [
                {"title": "Engineer", "location": ""},
                {"title": "Manager", "location": None},
                {"title": "Director"},  # No location key
            ]
        }
        result = _extract_experience_locations(Resume.from_dict(data))
        self.assertEqual(result, [])

    def test_handles_missing_experience(self):
        from resume.docx_writer import _extract_experience_locations
        from resume.schema import Resume
        data = {}
        result = _extract_experience_locations(Resume.from_dict(data))
        self.assertEqual(result, [])

    def test_handles_none_experience(self):
        from resume.docx_writer import _extract_experience_locations
        from resume.schema import Resume
        # experience: None is coerced to an empty list by the schema.
        data = {"experience": None}
        result = _extract_experience_locations(Resume.from_dict(data))
        self.assertEqual(result, [])

    def test_preserves_order(self):
        from resume.docx_writer import _extract_experience_locations
        from resume.schema import Resume
        data = {
            "experience": [
                {"location": "A"},
                {"location": "C"},
                {"location": "B"},
            ]
        }
        result = _extract_experience_locations(Resume.from_dict(data))
        self.assertEqual(result, ["A", "C", "B"])


class TestGetContactField(unittest.TestCase):
    """Tests for _get_contact_field function."""

    def test_returns_top_level_field(self):
        from resume.docx_writer import _get_contact_field
        from resume.schema import Resume
        data = {"email": "test@example.com"}
        result = _get_contact_field(Resume.from_dict(data), "email")
        self.assertEqual(result, "test@example.com")

    def test_returns_nested_contact_field(self):
        from resume.docx_writer import _get_contact_field
        from resume.schema import Resume
        # Resume.from_dict promotes contact.email to resume.email when the
        # top-level key is absent, so the field is accessible either way.
        data = {"contact": {"email": "nested@example.com"}}
        result = _get_contact_field(Resume.from_dict(data), "email")
        self.assertEqual(result, "nested@example.com")

    def test_prefers_top_level_over_nested(self):
        from resume.docx_writer import _get_contact_field
        from resume.schema import Resume
        data = {
            "email": "top@example.com",
            "contact": {"email": "nested@example.com"},
        }
        result = _get_contact_field(Resume.from_dict(data), "email")
        self.assertEqual(result, "top@example.com")

    def test_returns_empty_if_missing(self):
        from resume.docx_writer import _get_contact_field
        from resume.schema import Resume
        data = {}
        result = _get_contact_field(Resume.from_dict(data), "email")
        self.assertEqual(result, "")


class TestCollectLinkExtras(unittest.TestCase):
    """Tests for _collect_link_extras function."""

    def test_collects_website_linkedin_github(self):
        from resume.docx_writer import _collect_link_extras
        from resume.schema import Resume
        data = {
            "website": "https://example.com",
            "linkedin": "https://linkedin.com/in/johndoe",
            "github": "https://github.com/johndoe",
        }
        result = _collect_link_extras(Resume.from_dict(data))
        self.assertEqual(len(result), 3)

    def test_collects_from_nested_contact(self):
        from resume.docx_writer import _collect_link_extras
        from resume.schema import Resume
        # website is not in _CONTACT_PROMOTED, so it stays in resume.contact
        # and is read via the contact dict fallback in _get_contact_field.
        data = {
            "contact": {
                "website": "https://example.com",
            }
        }
        result = _collect_link_extras(Resume.from_dict(data))
        self.assertEqual(len(result), 1)

    def test_collects_from_links_list(self):
        from resume.docx_writer import _collect_link_extras
        from resume.schema import Resume
        data = {
            "links": ["https://portfolio.com", "https://blog.com"],
        }
        result = _collect_link_extras(Resume.from_dict(data))
        self.assertEqual(len(result), 2)

    def test_handles_empty_data(self):
        from resume.docx_writer import _collect_link_extras
        from resume.schema import Resume
        data = {}
        result = _collect_link_extras(Resume.from_dict(data))
        self.assertEqual(result, [])


class TestResolveSections(unittest.TestCase):
    """Tests for _resolve_sections function."""

    def test_returns_template_sections_by_default(self):
        from resume.docx_writer import _resolve_sections
        template = {
            "sections": [
                {"key": "summary", "title": "Summary"},
                {"key": "experience", "title": "Experience"},
            ]
        }
        result = _resolve_sections(template, None)
        self.assertEqual(len(result), 2)
        self.assertEqual(result[0]["key"], "summary")

    def test_respects_structure_order(self):
        from resume.docx_writer import _resolve_sections
        template = {
            "sections": [
                {"key": "summary", "title": "Summary"},
                {"key": "experience", "title": "Experience"},
                {"key": "skills", "title": "Skills"},
            ]
        }
        structure = {"order": ["skills", "experience", "summary"]}
        result = _resolve_sections(template, structure)
        self.assertEqual(len(result), 3)
        self.assertEqual(result[0]["key"], "skills")
        self.assertEqual(result[1]["key"], "experience")
        self.assertEqual(result[2]["key"], "summary")

    def test_uses_custom_titles_from_structure(self):
        from resume.docx_writer import _resolve_sections
        template = {"sections": [{"key": "skills", "title": "Skills"}]}
        structure = {
            "order": ["skills"],
            "titles": {"skills": "Technical Expertise"},
        }
        result = _resolve_sections(template, structure)
        # When key is in template, it uses template's config
        self.assertEqual(result[0]["key"], "skills")

    def test_handles_empty_sections(self):
        from resume.docx_writer import _resolve_sections
        template = {"sections": []}
        result = _resolve_sections(template, None)
        self.assertEqual(result, [])


class TestSectionRenderers(unittest.TestCase):
    """Tests for SECTION_RENDERERS registry."""

    def test_has_expected_renderers(self):
        from resume.docx_writer import SECTION_RENDERERS
        expected_keys = [
            "summary",
            "skills",
            "technologies",
            "interests",
            "presentations",
            "languages",
            "coursework",
            "certifications",
            "experience",
            "education",
        ]
        for key in expected_keys:
            self.assertIn(key, SECTION_RENDERERS, f"Missing renderer for {key}")


class TestSectionsWithKeywords(unittest.TestCase):
    """Tests for SECTIONS_WITH_KEYWORDS set."""

    def test_includes_expected_sections(self):
        from resume.docx_writer import SECTIONS_WITH_KEYWORDS
        self.assertIn("summary", SECTIONS_WITH_KEYWORDS)
        self.assertIn("experience", SECTIONS_WITH_KEYWORDS)

    def test_excludes_other_sections(self):
        from resume.docx_writer import SECTIONS_WITH_KEYWORDS
        self.assertNotIn("skills", SECTIONS_WITH_KEYWORDS)
        self.assertNotIn("education", SECTIONS_WITH_KEYWORDS)


class TestWriteResumeDocx(unittest.TestCase):
    """Tests for write_resume_docx function."""

    @patch("resume.docx_writer.safe_import")
    def test_raises_when_docx_unavailable(self, mock_safe_import):
        from resume.docx_writer import write_resume_docx
        from resume.schema import Resume
        mock_safe_import.return_value = None
        with self.assertRaises(RuntimeError) as ctx:
            write_resume_docx(Resume(), {}, test_path("out.docx"))  # nosec B108 - test fixture path
        self.assertIn("python-docx", str(ctx.exception))

    @patch("resume.docx_writer.safe_import")
    def test_creates_document_and_saves(self, mock_safe_import):
        from resume.docx_writer import write_resume_docx
        from resume.schema import Resume
        # Mock the docx module with all necessary components
        mock_docx = MagicMock()
        mock_doc = MagicMock()
        mock_section = MagicMock()
        mock_doc.sections = [mock_section]
        # Mock paragraphs list that grows when add_heading/add_paragraph is called
        mock_paragraphs = []
        mock_doc.paragraphs = mock_paragraphs

        add_side_effect = make_docx_paragraph_side_effect(mock_paragraphs)
        mock_doc.add_heading.side_effect = add_side_effect
        mock_doc.add_paragraph.side_effect = add_side_effect
        mock_doc.styles = {"Normal": MagicMock(), "Heading 1": MagicMock(), "Title": MagicMock()}
        mock_docx.Document.return_value = mock_doc
        mock_safe_import.return_value = mock_docx

        # Patch the import inside the function
        with patch.dict("sys.modules", {"docx": mock_docx}):
            resume = Resume.from_dict({"name": "John Doe"})
            template = {"sections": [], "page": {"compact": False}}
            write_resume_docx(resume, template, test_path("test.docx"))  # nosec B108 - test fixture path

        mock_doc.save.assert_called_once_with(test_path("test.docx"))  # nosec B108 - test fixture path


class TestSectionSynonyms(unittest.TestCase):
    """Tests for SECTION_SYNONYMS mapping."""

    def test_summary_synonyms(self):
        from resume.docx_writer import SECTION_SYNONYMS
        self.assertIn("summary", SECTION_SYNONYMS["summary"])
        self.assertIn("profile", SECTION_SYNONYMS["summary"])
        self.assertIn("about", SECTION_SYNONYMS["summary"])

    def test_skills_synonyms(self):
        from resume.docx_writer import SECTION_SYNONYMS
        self.assertIn("skills", SECTION_SYNONYMS["skills"])
        self.assertIn("technical skills", SECTION_SYNONYMS["skills"])

    def test_experience_synonyms(self):
        from resume.docx_writer import SECTION_SYNONYMS
        self.assertIn("experience", SECTION_SYNONYMS["experience"])
        self.assertIn("work history", SECTION_SYNONYMS["experience"])
        self.assertIn("employment", SECTION_SYNONYMS["experience"])

    def test_education_synonyms(self):
        from resume.docx_writer import SECTION_SYNONYMS
        self.assertIn("education", SECTION_SYNONYMS["education"])
        self.assertIn("academics", SECTION_SYNONYMS["education"])


class TestUsePlainBullets(unittest.TestCase):
    """Tests for BulletRenderer.get_bullet_config (successor to _use_plain_bullets)."""

    def test_returns_tuple(self):
        from tests.resume_tests.fixtures import mock_docx_modules, make_fake_renderer
        with mock_docx_modules:
            from resume.docx_renderers import BulletRenderer
            renderer, _doc = make_fake_renderer(BulletRenderer)
            result = renderer.get_bullet_config(None)
        self.assertIsInstance(result, tuple)


class TestApplyPageStyles(unittest.TestCase):
    """Tests for _apply_page_styles function."""

    @staticmethod
    def _make_mock_doc():
        doc = MagicMock()
        doc.core_properties = MagicMock()
        return doc

    def test_skips_non_compact(self):
        from resume.docx_writer import _apply_page_styles
        doc = self._make_mock_doc()
        _apply_page_styles(doc, {"compact": False})
        doc.sections.__getitem__.assert_not_called()

    def test_skips_empty_config(self):
        from resume.docx_writer import _apply_page_styles
        doc = self._make_mock_doc()
        _apply_page_styles(doc, {})
        doc.sections.__getitem__.assert_not_called()


class TestSetDocumentMetadata(unittest.TestCase):
    """Tests for _set_document_metadata function."""

    def setUp(self):
        self.doc = MagicMock()
        self.doc.core_properties = MagicMock()

    def test_sets_title_from_name(self):
        from resume.docx_writer import _set_document_metadata
        from resume.schema import Resume
        data = {"name": "John Doe", "email": "john@example.com"}
        _set_document_metadata(self.doc, Resume.from_dict(data), {})
        self.assertIn("John Doe", self.doc.core_properties.title)

    def test_sets_author_from_name(self):
        from resume.docx_writer import _set_document_metadata
        from resume.schema import Resume
        data = {"name": "Jane Smith"}
        _set_document_metadata(self.doc, Resume.from_dict(data), {})
        self.assertEqual(self.doc.core_properties.author, "Jane Smith")

    def test_includes_experience_locations_in_keywords(self):
        from resume.docx_writer import _set_document_metadata
        from resume.schema import Resume
        data = {"name": "Test", "experience": [{"location": "NYC"}, {"location": "LA"}]}
        template = {"page": {"metadata_include_locations": True}}
        _set_document_metadata(self.doc, Resume.from_dict(data), template)
        self.assertIn("NYC", self.doc.core_properties.keywords)
        self.assertIn("LA", self.doc.core_properties.keywords)

    def test_excludes_locations_when_disabled(self):
        from resume.docx_writer import _set_document_metadata
        from resume.schema import Resume
        data = {"name": "Test", "experience": [{"location": "NYC"}]}
        template = {"page": {"metadata_include_locations": False}}
        _set_document_metadata(self.doc, Resume.from_dict(data), template)
        self.assertNotIn("NYC", self.doc.core_properties.keywords)

    def test_handles_missing_data(self):
        from resume.docx_writer import _set_document_metadata
        from resume.schema import Resume
        _set_document_metadata(self.doc, Resume.from_dict({}), {})
        self.assertEqual(self.doc.core_properties.title, "Resume")


class TestCenterParagraph(unittest.TestCase):
    """Tests for _center_paragraph function."""

    def test_centers_paragraph(self):
        from resume.docx_writer import _center_paragraph
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        para = MagicMock()
        para.paragraph_format = MagicMock()
        _center_paragraph(para)
        self.assertEqual(para.alignment, WD_ALIGN_PARAGRAPH.CENTER)

    def test_handles_exception_gracefully(self):
        from unittest.mock import PropertyMock
        from resume.docx_writer import _center_paragraph
        para = MagicMock()
        type(para).alignment = PropertyMock(side_effect=Exception("test"))
        # Should not raise
        _center_paragraph(para)


if __name__ == "__main__":
    unittest.main()
