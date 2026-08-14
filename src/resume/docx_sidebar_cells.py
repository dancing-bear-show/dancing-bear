"""Cell-level formatting primitives for DOCX sidebar resume layout.

Extracted from docx_sidebar. Provides:
  - _set_cell_shading: background shading on table cells
  - _remove_cell_borders: strip all borders from a table cell
  - _add_indented_run: indented paragraph with styled run
  - _render_sidebar_section: generic sidebar section with optional bullets
  - _render_main_section_heading: section heading in main column
  - _render_edu_meta: institution and year secondary paragraph
"""
from __future__ import annotations

from typing import Any

from docx.shared import Pt, Inches, RGBColor  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.oxml import OxmlElement  # type: ignore

from .docx_styles import _parse_hex_color, _tight_paragraph
from .render_config import IndentedRunStyle


def _set_cell_shading(cell, hex_color: str) -> None:
    """Set background shading on a table cell."""
    rgb = _parse_hex_color(hex_color)
    if not rgb:
        return
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tc_pr.append(shd)


def _remove_cell_borders(cell) -> None:
    """Remove all borders from a table cell."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def _add_indented_run(cell, text: str, page_cfg: dict[str, Any], style: IndentedRunStyle | None = None):
    """Add an indented paragraph with styled run."""
    s = style or IndentedRunStyle()
    p = cell.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    if s.italic:
        run.italic = True
    run.font.size = Pt(page_cfg.get("meta_pt", 9) + s.size_offset)
    rgb = _parse_hex_color(s.color)
    if rgb:
        run.font.color.rgb = RGBColor(*rgb)
    _tight_paragraph(p, after_pt=s.after_pt)
    return p


def _render_sidebar_section(cell, title: str, items: list[str], page_cfg: dict[str, Any], bulleted: bool = True) -> None:
    """Render a generic section in sidebar with optional bullets."""
    h1_color = page_cfg.get("h1_color", "#D4A84B")
    text_color = page_cfg.get("sidebar_text_color", "#333333")
    bullet_color = page_cfg.get("sidebar_bullet_color", "#4A90A4")

    p = cell.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(page_cfg.get("h1_pt", 14))
    rgb = _parse_hex_color(h1_color)
    if rgb:
        run.font.color.rgb = RGBColor(*rgb)
    _tight_paragraph(p, before_pt=12, after_pt=6)

    for item in items:
        p = cell.add_paragraph()
        if bulleted:
            bullet_run = p.add_run("• ")
            bullet_rgb = _parse_hex_color(bullet_color)
            if bullet_rgb:
                bullet_run.font.color.rgb = RGBColor(*bullet_rgb)
            bullet_run.font.size = Pt(page_cfg.get("body_pt", 10))
        run = p.add_run(item)
        run.font.size = Pt(page_cfg.get("body_pt", 10))
        rgb = _parse_hex_color(text_color)
        if rgb:
            run.font.color.rgb = RGBColor(*rgb)
        _tight_paragraph(p, after_pt=2)


def _render_main_section_heading(cell, title: str, page_cfg: dict[str, Any]) -> None:
    """Render a section heading in main column."""
    h1_color = page_cfg.get("h1_color", "#D4A84B")

    p = cell.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(page_cfg.get("h1_pt", 14))
    rgb = _parse_hex_color(h1_color)
    if rgb:
        run.font.color.rgb = RGBColor(*rgb)
    _tight_paragraph(p, before_pt=10, after_pt=6)


def _render_edu_meta(cell, institution: str, year: str, page_cfg: dict[str, Any]) -> None:
    """Render institution and year as secondary paragraph."""
    if not institution and not year:
        return
    p2 = cell.add_paragraph()
    p2.paragraph_format.left_indent = Inches(0.25)
    inst_run = p2.add_run(institution)
    inst_run.italic = True
    inst_run.font.size = Pt(page_cfg.get("meta_pt", 9))
    inst_rgb = _parse_hex_color("#666666")
    if inst_rgb:
        inst_run.font.color.rgb = RGBColor(*inst_rgb)
    if year:
        year_run = p2.add_run(f"  {year}")
        year_run.font.size = Pt(page_cfg.get("meta_pt", 9))
        if inst_rgb:
            year_run.font.color.rgb = RGBColor(*inst_rgb)
    _tight_paragraph(p2, after_pt=6)
