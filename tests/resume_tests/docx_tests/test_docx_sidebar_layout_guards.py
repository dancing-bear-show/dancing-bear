"""Sidebar layout guards: column-width validation and contact-line parity.

Both behaviours here are regressions of the same class as the body-less render
bug: the renderer produced a plausible-looking document that was silently wrong,
with a zero exit code and no warning. These tests therefore assert on rendered
content, not on return codes.

Sad-path methods are named test_rejects_* / test_invalid_* per the naming
contract in workflows/resume/consolidate-schema.yaml, so the ratio is countable
with grep.
"""

from __future__ import annotations

import os
import tempfile
import unittest

from resume.docx_sidebar_sections import (
    _DEFAULT_USABLE_WIDTH_IN,
    _usable_width_in,
    _validate_column_width,
)
from resume.docx_writer import write_resume_docx


def _sidebar_template(**layout_overrides):
    """A minimal sidebar template; layout keys overridable per test."""
    layout = {"type": "sidebar"}
    layout.update(layout_overrides)
    return {
        "layout": layout,
        "sidebar_sections": [{"key": "skills", "title": "Skills"}],
        "sections": [
            {"key": "summary", "title": "Summary"},
            {"key": "experience", "title": "Experience"},
        ],
    }


def _candidate():
    return {
        "name": "Jane Doe",
        "headline": "Senior Site Reliability Engineer",
        "email": "jane@example.com",
        "phone": "(555) 010-0202",
        "location": "Austin, TX",
        "linkedin": "linkedin.com/in/janedoe",
        "summary": "Builds reliable systems.",
        "skills": ["Python", "Kubernetes"],
        "experience": [
            {
                "title": "Principal SRE",
                "company": "Northwind",
                "start": "2021",
                "end": "Present",
                "bullets": ["Cut p99 latency 43%"],
            }
        ],
    }


class TestColumnWidthValidation(unittest.TestCase):
    """layout.sidebar_width / main_width are INCHES, not percent."""

    def test_accepts_default_widths(self):
        """The documented defaults (2.3 + 5.2 = 7.5) are in range."""
        self.assertEqual(_validate_column_width(2.3, "sidebar_width"), 2.3)
        self.assertEqual(_validate_column_width(5.2, "main_width"), 5.2)

    def test_accepts_integer_and_string_numerics(self):
        """YAML may deliver an int or a quoted number; both coerce to float."""
        self.assertEqual(_validate_column_width(3, "sidebar_width"), 3.0)
        self.assertEqual(_validate_column_width("2.5", "sidebar_width"), 2.5)

    def test_rejects_percent_looking_value(self):
        """34 reads as a percentage but would render a 34-inch column.

        This is the actual mistake: on an 8.5" page the layout collapses into an
        unreadable sliver, with no error and a plausible file size.
        """
        with self.assertRaises(ValueError) as ctx:
            _validate_column_width(34, "sidebar_width")
        msg = str(ctx.exception)
        self.assertIn("inches", msg.lower())
        self.assertIn("34", msg)

    def test_rejects_width_exactly_over_usable_page(self):
        """Boundary: the usable width itself passes, just above it does not."""
        self.assertEqual(
            _validate_column_width(_DEFAULT_USABLE_WIDTH_IN, "main_width"),
            _DEFAULT_USABLE_WIDTH_IN,
        )
        with self.assertRaises(ValueError):
            _validate_column_width(_DEFAULT_USABLE_WIDTH_IN + 0.01, "main_width")

    def test_rejects_zero_and_negative_widths(self):
        """A non-positive column cannot render."""
        for bad in (0, -1, -2.5):
            with self.subTest(width=bad):
                with self.assertRaises(ValueError):
                    _validate_column_width(bad, "sidebar_width")

    def test_rejects_non_numeric_width(self):
        """A string or None must fail loudly, not coerce to a default."""
        for bad in ("wide", None, [], {}):
            with self.subTest(width=bad):
                with self.assertRaises(ValueError) as ctx:
                    _validate_column_width(bad, "sidebar_width")
                self.assertIn("sidebar_width", str(ctx.exception))

    def test_invalid_combined_widths_exceed_page(self):
        """Each column fits alone, but together they overflow the page."""
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, "resume.docx")
            template = _sidebar_template(sidebar_width=5.0, main_width=5.0)
            with self.assertRaises(ValueError) as ctx:
                write_resume_docx(
                    data=_candidate(), template=template, out_path=out
                )
            self.assertIn("exceeds", str(ctx.exception))

    def test_invalid_width_raises_through_the_writer(self):
        """The guard fires on the real render path, not only when called directly."""
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, "resume.docx")
            with self.assertRaises(ValueError):
                write_resume_docx(
                    data=_candidate(),
                    template=_sidebar_template(sidebar_width=34),
                    out_path=out,
                )


class TestUsableWidthFromSection(unittest.TestCase):
    """The width limit follows the real page geometry, not a hardcoded 7.5"."""

    def test_derives_usable_width_from_margins(self):
        """page.margins_in is configurable, so the limit must move with it.

        Note python-docx's own default margin is 1.25", not the 0.5" this
        renderer applies via _apply_page_styles — so a bare Document() is 6.0"
        usable, and 7.5" only holds once those styles have run.
        """
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        self.assertAlmostEqual(_usable_width_in(doc), 6.0, places=2)

        sec = doc.sections[0]
        sec.left_margin = Inches(0.5)
        sec.right_margin = Inches(0.5)
        self.assertAlmostEqual(_usable_width_in(doc), 7.5, places=2)

        sec.left_margin = Inches(1.5)
        sec.right_margin = Inches(1.5)
        self.assertAlmostEqual(_usable_width_in(doc), 5.5, places=2)

    def test_rejects_width_valid_at_default_margins_but_not_at_wide_ones(self):
        """A 6" column fits a 0.5"-margin page but overflows a 1.5"-margin one.

        This is the case a hardcoded 7.5" limit accepted incorrectly.
        """
        self.assertEqual(_validate_column_width(6.0, "main_width", 7.5), 6.0)
        with self.assertRaises(ValueError) as ctx:
            _validate_column_width(6.0, "main_width", 5.5)
        self.assertIn("5.5", str(ctx.exception))

    def test_accepts_shipped_defaults_on_a_non_compact_template(self):
        """The documented 2.3 + 5.2 defaults must render without page.compact.

        Margins are only narrowed to 0.5" when page.compact is set, so a plain
        template leaves python-docx's 1.25" default (6.0" usable) while the
        column defaults total 7.5". The guard takes the wider of measured and
        documented width so it targets percent-shaped values, not the defaults.
        """
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, "resume.docx")
            write_resume_docx(
                data=_candidate(), template=_sidebar_template(), out_path=out
            )
            self.assertTrue(os.path.exists(out))

    def test_invalid_geometry_falls_back_to_default(self):
        """An unreadable section must not make the guard itself raise."""

        class _Broken:
            @property
            def sections(self):
                raise RuntimeError("no sections")

        self.assertEqual(_usable_width_in(_Broken()), _DEFAULT_USABLE_WIDTH_IN)


class TestSidebarContactParity(unittest.TestCase):
    """The sidebar header must carry the same link extras as the standard layout."""

    def _render_and_read_header(self, template):
        from docx import Document

        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, "resume.docx")
            write_resume_docx(data=_candidate(), template=template, out_path=out)
            doc = Document(out)
            return " | ".join(
                p.text for p in doc.sections[0].header.paragraphs if p.text.strip()
            )

    def test_sidebar_header_includes_linkedin(self):
        """Regression: the sidebar header dropped linkedin/website/github.

        The standard layout rendered them via _collect_link_extras while the
        sidebar hardcoded [phone, email, location], so the same profile lost its
        links depending on which layout was chosen.
        """
        header = self._render_and_read_header(_sidebar_template())
        self.assertIn("linkedin.com/in/janedoe", header)
        # The originally-present fields must survive the change.
        self.assertIn("jane@example.com", header)
        self.assertIn("(555) 010-0202", header)
        self.assertIn("Austin, TX", header)

    def test_sidebar_header_omits_absent_links(self):
        """A profile with no linkedin must not emit an empty separator."""
        from docx import Document

        data = _candidate()
        del data["linkedin"]
        with tempfile.TemporaryDirectory() as tmpdir:
            out = os.path.join(tmpdir, "resume.docx")
            write_resume_docx(
                data=data, template=_sidebar_template(), out_path=out
            )
            header = " | ".join(
                p.text
                for p in Document(out).sections[0].header.paragraphs
                if p.text.strip()
            )
        self.assertNotIn("linkedin", header.lower())
        self.assertIn("jane@example.com", header)


if __name__ == "__main__":
    unittest.main()
