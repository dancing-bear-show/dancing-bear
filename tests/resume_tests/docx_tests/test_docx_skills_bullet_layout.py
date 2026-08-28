"""Rendered-output contract for skills groups and sidebar entry alignment.

These assert properties of the *rendered* .docx rather than of a fake document,
because both properties under test are only observable once paragraphs exist:
"how many bullet glyphs ended up inside one paragraph" is not a call the
renderer makes, it is a shape the output has.

The style target is a single-column resume that uses ``Normal`` paragraphs with
a literal "• " prefix and never nests a glyph inside a paragraph. No content
from that document appears here -- every fixture value below is invented.
"""

from __future__ import annotations

import os
import tempfile
import unittest
from typing import TYPE_CHECKING

if TYPE_CHECKING:  # pragma: no cover - typing only, never imported at runtime
    from docx.text.paragraph import Paragraph


def _docx_available() -> bool:
    """Check if python-docx is installed."""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


# Guarded at module scope, not merely on the test classes: ``resume.docx_styles``
# imports ``docx.*`` eagerly, so importing the resume DOCX modules below is
# itself what raises ImportError when python-docx is absent.
if _docx_available():
    from resume.docx_writer import write_resume_docx
    from resume.schema import Resume

    from tests.resume_tests.golden.golden_fixtures import (
        sidebar_template,
        standard_template,
    )

BULLET = "•"

# The reference document's longest bullet paragraph. A group collapsed into one
# paragraph blows past this immediately, so it doubles as a ceiling check.
REFERENCE_MAX_PARAGRAPH_CHARS = 310


def _all_paragraphs(path: str) -> list[Paragraph]:
    """Every paragraph in the document, including those inside table cells.

    The sidebar layout puts its content in a table, so a body-only walk would
    silently skip that entire layout.
    """
    import docx

    doc = docx.Document(path)
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _render(resume_data: dict, template: dict) -> list[Paragraph]:
    """Render a fixture through the public writer and return its paragraphs."""
    with tempfile.TemporaryDirectory() as tmp:
        out = os.path.join(tmp, "resume.docx")
        write_resume_docx(Resume.from_dict(resume_data), template, out)
        return _all_paragraphs(out)


def _multi_group_skills(group_count: int = 7, items_per_group: int = 6) -> dict:
    """A skills-group resume shaped like a real one: several groups, long descs.

    Sized deliberately: the collapsing bug is proportional to items-per-group,
    so a two-item fixture would still fit on one line and hide the defect.
    """
    groups = []
    for gi in range(group_count):
        groups.append({
            "title": f"Group {gi}",
            "items": [
                {
                    "name": f"Capability {gi}-{i}",
                    "desc": "a representative description of the capability "
                            "and the context where it was applied",
                }
                for i in range(items_per_group)
            ],
        })
    return {"name": "Example Candidate", "headline": "Example Headline",
            "skills_groups": groups}


@unittest.skipUnless(_docx_available(), "python-docx not installed")
class SkillsGroupBulletLayoutTests(unittest.TestCase):
    """Each skill item renders as its own bullet paragraph."""

    def test_no_paragraph_contains_more_than_one_bullet_glyph(self):
        """The regression that would silently return: items joined into one run.

        Asserted directly on glyph counts rather than on paragraph count,
        because a renderer could emit the right number of paragraphs and still
        bury separators inside them.
        """
        paragraphs = _render(_multi_group_skills(), standard_template())

        offenders = [p.text for p in paragraphs if p.text.count(BULLET) > 1]
        self.assertEqual(
            offenders, [],
            f"{len(offenders)} paragraph(s) hold more than one '{BULLET}'; "
            "skill items were joined into a single paragraph instead of "
            "getting one paragraph each.",
        )

    def test_each_skill_item_is_its_own_paragraph(self):
        """Every item appears as a bullet paragraph carrying only that item."""
        data = _multi_group_skills(group_count=2, items_per_group=3)
        paragraphs = _render(data, standard_template())

        bullet_texts = [p.text for p in paragraphs if p.text.strip().startswith(BULLET)]
        for gi in range(2):
            for i in range(3):
                needle = f"Capability {gi}-{i}"
                matches = [t for t in bullet_texts if needle in t]
                self.assertEqual(
                    len(matches), 1,
                    f"expected exactly one bullet paragraph for {needle}, got {matches}",
                )

    def test_group_title_renders_unprefixed(self):
        """The group name is its own line and carries no bullet glyph."""
        data = _multi_group_skills(group_count=2, items_per_group=2)
        paragraphs = _render(data, standard_template())

        titles = [p.text for p in paragraphs if p.text.strip() in ("Group 0", "Group 1")]
        self.assertEqual(
            sorted(titles), ["Group 0", "Group 1"],
            "group titles must render on their own line, unprefixed and "
            "not folded into the first item's paragraph",
        )

    def test_no_paragraph_exceeds_reference_maximum_length(self):
        """A collapsed group produces one oversized prose paragraph."""
        paragraphs = _render(_multi_group_skills(), standard_template())

        longest = max((len(p.text) for p in paragraphs), default=0)
        self.assertLessEqual(
            longest, REFERENCE_MAX_PARAGRAPH_CHARS,
            f"longest paragraph is {longest} chars, over the "
            f"{REFERENCE_MAX_PARAGRAPH_CHARS}-char reference maximum",
        )

    def test_sidebar_layout_also_avoids_multi_bullet_paragraphs(self):
        """The same property holds for the two-column layout."""
        paragraphs = _render(_multi_group_skills(), sidebar_template())

        offenders = [p.text for p in paragraphs if p.text.count(BULLET) > 1]
        self.assertEqual(offenders, [])


@unittest.skipUnless(_docx_available(), "python-docx not installed")
class SidebarEntryIndentTests(unittest.TestCase):
    """A sidebar entry's bullet line and its continuation lines agree.

    The sidebar is a table layout with its own hanging-indent convention: the
    glyph is a run inside the headline paragraph (so the headline text already
    starts past it) and continuation lines carry an explicit 0.25" indent to
    line up underneath that text. Education, experience and presentations all
    use it, so it is pinned here as the shared local contract rather than
    per-section.
    """

    CONTINUATION_INDENT_EMU = 228600  # 0.25"

    def _entry_fixture(self) -> dict:
        return {
            "name": "Example Candidate",
            "education": [
                {"degree": "BSc Example", "institution": "Example University",
                 "year": "2010"},
            ],
            "experience": [
                {"title": "Example Role", "company": "Example Co",
                 "start": "2020", "end": "2024",
                 "bullets": [{"text": "an example accomplishment line"}]},
            ],
            "presentations": [
                {"title": "Example Talk Title", "authors": "A. Author",
                 "event": "Example Conference 2019", "note": "an example note"},
            ],
        }

    def test_presentation_continuations_share_the_sidebar_indent(self):
        """Authors, event and note all sit at the shared continuation indent."""
        paragraphs = _render(self._entry_fixture(), sidebar_template())
        by_text = {p.text.strip(): p for p in paragraphs if p.text.strip()}

        for line in ("A. Author", "Example Conference 2019", "an example note"):
            self.assertIn(line, by_text)
            self.assertEqual(
                by_text[line].paragraph_format.left_indent,
                self.CONTINUATION_INDENT_EMU,
                f"{line!r} must sit at the sidebar continuation indent so it "
                "aligns under the bullet line's text",
            )

    def test_presentation_bullet_line_matches_other_sidebar_sections(self):
        """The glyph-bearing headline is unindented in every sidebar section."""
        paragraphs = _render(self._entry_fixture(), sidebar_template())

        headlines = [
            p for p in paragraphs if p.text.strip().startswith(BULLET)
        ]
        self.assertGreaterEqual(len(headlines), 3, "expected edu/exp/pres headlines")
        for p in headlines:
            self.assertIsNone(
                p.paragraph_format.left_indent,
                f"{p.text!r} carries an explicit indent; sidebar bullet lines "
                "hang the glyph at the margin instead",
            )


def _group_with_items(items: list[dict]) -> dict:
    """A single-group skills resume carrying exactly the given raw items."""
    return {
        "name": "Example Candidate",
        "skills_groups": [{"title": "Cloud", "items": items}],
    }


def _bullet_texts(paragraphs: list[Paragraph]) -> list[str]:
    """Text of every paragraph that starts with the bullet glyph."""
    return [p.text for p in paragraphs if p.text.strip().startswith(BULLET)]


def _bullet_bodies(paragraphs: list[Paragraph]) -> list[str]:
    """The text of each bullet paragraph with its leading glyph removed."""
    return [t.strip()[len(BULLET):].strip() for t in _bullet_texts(paragraphs)]


@unittest.skipUnless(_docx_available(), "python-docx not installed")
class SkillsGroupEmptyItemTests(unittest.TestCase):
    """An item that normalizes to empty produces no paragraph at all.

    A ``SkillGroupItem`` built from a dict carrying no recognised name key
    (``name|title|label``) and no ``desc`` resolves to the empty string.
    Emitted anyway it became a bare "• " line -- a visible empty bullet.

    These assert paragraph *counts* rather than merely the absence of the
    glyph, because an empty bullet still starts with the glyph: only counting
    distinguishes "skipped" from "emitted with no text".
    """

    def test_rejects_item_with_unknown_name_key(self):
        """An unrecognised key yields no paragraph, not a bare glyph."""
        data = _group_with_items(
            [{"name": "AWS"}, {"bogus": "ignored"}, {"name": "K8s"}]
        )
        paragraphs = _render(data, standard_template())

        self.assertEqual(
            _bullet_bodies(paragraphs), ["AWS", "K8s"],
            "an item normalizing to empty must emit no paragraph at all",
        )

    def test_rejects_item_whose_name_is_only_whitespace(self):
        """A whitespace-only name normalizes to empty and is skipped."""
        data = _group_with_items(
            [{"name": "AWS"}, {"name": "   "}, {"name": "K8s"}]
        )

        self.assertEqual(
            _bullet_bodies(_render(data, standard_template())), ["AWS", "K8s"]
        )

    def test_surrounding_items_render_in_order(self):
        """Skipping an empty item does not disturb the items around it."""
        data = _group_with_items(
            [
                {"name": "AWS"},
                {"bogus": "ignored"},
                {"name": "K8s"},
                {"name": "Terraform"},
            ]
        )

        self.assertEqual(
            _bullet_bodies(_render(data, standard_template())),
            ["AWS", "K8s", "Terraform"],
        )

    def test_alias_keyed_item_still_renders(self):
        """``title`` and ``label`` alias onto ``name`` and must not be skipped.

        Guards against over-skipping: these resolve to real text, so a guard
        keyed off the raw dict rather than the normalized string would drop
        them silently.
        """
        data = _group_with_items(
            [{"title": "Aliased Title"}, {"label": "Aliased Label"}]
        )

        self.assertEqual(
            _bullet_bodies(_render(data, standard_template())),
            ["Aliased Title", "Aliased Label"],
        )

    def test_rejects_empty_item_in_the_bullets_branch_too(self):
        """The other branch of ``_render_groups`` shares the same chokepoint.

        ``_render_groups`` routes to ``_render_bullet_items`` or
        ``_render_inline_items`` purely on the ``bullets`` config flag, so a
        filter applied to only one branch would leave the bare glyph reachable
        by flipping one config value. Asserted through the renderer directly
        because the standard template does not enable the plain-bullet branch.
        """
        import docx

        from resume.docx_sections_skills import SkillsSectionRenderer
        from resume.schema import SkillGroupItem

        renderer = SkillsSectionRenderer(docx.Document())
        raw = [
            SkillGroupItem.from_dict({"name": "AWS"}),
            SkillGroupItem.from_dict({"bogus": "ignored"}),
            SkillGroupItem.from_dict({"name": "K8s"}),
        ]

        self.assertEqual(
            renderer._normalize_group_items(raw, True, " - "), ["AWS", "K8s"],
            "the empty item must be dropped before either render branch",
        )

    def test_description_only_item_still_renders(self):
        """An item carrying only ``desc`` has real content and must render.

        ``_labeled_item_text`` joins the empty name to the desc with the
        separator, so this normalizes to non-empty text rather than "".
        """
        data = _group_with_items(
            [{"name": "AWS"}, {"desc": "a real description"}, {"name": "K8s"}]
        )
        bodies = _bullet_bodies(_render(data, standard_template()))

        self.assertEqual(len(bodies), 3, f"desc-only item was dropped: {bodies}")
        self.assertTrue(
            any("a real description" in b for b in bodies),
            f"desc-only item lost its text: {bodies}",
        )
