"""DOCX section renderers for resume content.

Provides reusable components for rendering resume sections.
"""
from __future__ import annotations

from typing import Any, Dict, List, Optional

from docx.shared import Pt, RGBColor  # type: ignore

from .docx_styles import StyleManager, TextFormatter


class BulletRenderer:
    """Renders bullet lists with various styles."""

    def __init__(self, doc, page_cfg: Optional[Dict[str, Any]] = None):
        self.doc = doc
        self.page_cfg = page_cfg or {}
        self.styles = StyleManager()
        self.text = TextFormatter()

    def get_bullet_config(self, sec: Optional[Dict[str, Any]]) -> tuple:
        """Determine bullet style and glyph from config.

        Returns:
            Tuple of (use_plain: bool, glyph: str)
        """
        glyph = "•"
        style = None
        if sec:
            bul = sec.get("bullets") if isinstance(sec.get("bullets"), dict) else {}
            if bul:
                style = bul.get("style") or style
                glyph = bul.get("glyph") or glyph
            if sec.get("plain_bullets") is True:
                style = "plain"
        if not style and isinstance(self.page_cfg.get("bullets"), dict):
            bulp = self.page_cfg.get("bullets") or {}
            style = bulp.get("style") or style
            glyph = bulp.get("glyph") or glyph
        return (style == "plain" or (sec and sec.get("plain_bullets") is True), glyph)

    def add_bullet_line(
        self,
        text: str,
        *,
        sec: Optional[Dict[str, Any]] = None,
        keywords: Optional[List[str]] = None,
        glyph: str = "•",
    ):
        """Add a plain bullet line (glyph + text)."""
        p = self.doc.add_paragraph()
        self.styles.tight_paragraph(p, after_pt=0)
        self.styles.flush_left(p)
        p.add_run(f"{glyph} ")
        if keywords:
            self._bold_keywords(p, text, keywords)
        else:
            p.add_run(text)
        return p

    def add_named_bullet(
        self,
        name: str,
        desc: str,
        *,
        sec: Optional[Dict[str, Any]] = None,
        glyph: str = "•",
        sep: str = ": ",
    ):
        """Add a bullet with bold name and description."""
        p = self.doc.add_paragraph()
        self.styles.tight_paragraph(p, after_pt=0)
        self.styles.flush_left(p)
        p.add_run(f"{glyph} ")

        cfg = sec or {}
        name_color = cfg.get("name_color") or cfg.get("item_color") or cfg.get("title_color")

        r_name = p.add_run(name)
        r_name.bold = True
        self.styles.apply_run_color(r_name, name_color)

        p.add_run(sep)
        p.add_run(desc)
        return p

    def add_bullets(
        self,
        items: List[str],
        *,
        sec: Optional[Dict[str, Any]] = None,
        keywords: Optional[List[str]] = None,
        plain: bool = True,
        glyph: str = "•",
        list_style: str = "List Bullet",
    ):
        """Render a list of bullet items."""
        if plain:
            for it in items:
                self.add_bullet_line(it, sec=sec, keywords=keywords, glyph=glyph)
            return

        for it in items:
            p = self.doc.add_paragraph(style=list_style)
            self.styles.tight_paragraph(p, after_pt=0)
            self.styles.compact_bullet(p)
            if keywords:
                self._bold_keywords(p, it, keywords)
            else:
                p.add_run(it)

    def _bold_keywords(self, paragraph, text: str, keywords: List[str]):
        """Add text with keywords bolded."""
        lowered = text.lower()
        idx = 0
        found_any = False

        while idx < len(text):
            match_pos = None
            match_kw = None
            for kw in keywords:
                if not kw:
                    continue
                pos = lowered.find(kw.lower(), idx)
                if pos != -1 and (match_pos is None or pos < match_pos):
                    match_pos = pos
                    match_kw = text[pos:pos + len(kw)]

            if match_pos is None:
                paragraph.add_run(text[idx:])
                break

            if match_pos > idx:
                paragraph.add_run(text[idx:match_pos])

            br = paragraph.add_run(match_kw or "")
            br.bold = True
            found_any = True
            idx = match_pos + len(match_kw or "")

        if not found_any and idx == 0:
            paragraph.add_run(text)


class HeaderRenderer:
    """Renders header lines for experience and education entries."""

    def __init__(self, doc):
        self.doc = doc
        self.styles = StyleManager()

    def add_header_line(
        self,
        *,
        title_text: str = "",
        company_text: str = "",
        loc_text: str = "",
        span_text: str = "",
        sec: Optional[Dict[str, Any]] = None,
        style: str = "Normal",
    ):
        """Add a formatted header line.

        Format: Title at Company — [Location] — (Duration)
        """
        cfg = sec or {}
        p = self.doc.add_paragraph(style=style)
        self.styles.tight_paragraph(p, after_pt=0)
        self.styles.flush_left(p)

        item_color = cfg.get("item_color") or cfg.get("header_color")
        loc_color = cfg.get("location_color") or item_color
        dur_color = cfg.get("duration_color") or cfg.get("location_color") or item_color
        loc_brackets = bool(cfg.get("location_brackets", True))
        dur_brackets = bool(cfg.get("duration_brackets", True))

        meta_pt = None
        try:
            meta_pt = float(cfg.get("meta_pt")) if cfg.get("meta_pt") else None
        except Exception:
            pass

        # Title
        if title_text:
            r_title = p.add_run(title_text)
            r_title.bold = True
            self.styles.apply_run_color(r_title, item_color)

        # Company
        if title_text and company_text:
            p.add_run(" at ")
        if company_text:
            r_comp = p.add_run(company_text)
            r_comp.bold = True
            self.styles.apply_run_color(r_comp, item_color)

        # Location
        if loc_text:
            p.add_run(" — ")
            if loc_brackets:
                p.add_run("[")
            r_loc = p.add_run(loc_text)
            r_loc.italic = True
            self.styles.apply_run_size(r_loc, meta_pt)
            self.styles.apply_run_color(r_loc, loc_color)
            if loc_brackets:
                p.add_run("]")

        # Duration
        if span_text:
            p.add_run(" — ")
            if dur_brackets:
                p.add_run("(")
            r_span = p.add_run(span_text)
            self.styles.apply_run_size(r_span, meta_pt)
            self.styles.apply_run_color(r_span, dur_color)
            if dur_brackets:
                p.add_run(")")

        return p

    def add_group_title(
        self,
        title: str,
        sec: Optional[Dict[str, Any]] = None,
    ):
        """Add a group/category title with optional background."""
        title = (title or "").strip()
        if not title:
            return None

        cfg = sec or {}
        p = self.doc.add_paragraph()
        self.styles.tight_paragraph(p, after_pt=0)
        self.styles.flush_left(p)

        gt_color = cfg.get("group_title_color")
        gt_bg = cfg.get("group_title_bg") or cfg.get("title_bg")

        r = p.add_run(title)
        r.bold = True

        # Apply background shading
        bg_rgb = self.styles.parse_hex_color(gt_bg)
        if bg_rgb:
            self.styles.apply_shading(p, bg_rgb)
            if not gt_color:
                gt_color = self.styles.auto_contrast_color(bg_rgb)

        # Apply text color
        txt_color = gt_color or cfg.get("item_color") or cfg.get("title_color")
        self.styles.apply_run_color(r, txt_color)

        return p


class ListSectionRenderer:
    """Renders simple list sections (interests, languages, etc.)."""

    def __init__(self, doc, page_cfg: Optional[Dict[str, Any]] = None):
        self.doc = doc
        self.bullets = BulletRenderer(doc, page_cfg)
        self.text = TextFormatter()

    def render_simple_list(
        self,
        items: List[Any],
        sec: Optional[Dict[str, Any]] = None,
        *,
        name_keys: tuple = ("name", "title", "label", "text"),
        desc_key: Optional[str] = None,
        desc_sep: str = " — ",
    ) -> List[str]:
        """Normalize and render a simple list section.

        Args:
            items: Raw items (strings or dicts).
            sec: Section config.
            name_keys: Keys to try for item name.
            desc_key: Optional key for description (e.g., "level" for languages).
            desc_sep: Separator between name and description.

        Returns:
            List of normalized string items.
        """
        cfg = sec or {}
        lines: List[str] = []

        for it in items:
            if isinstance(it, dict):
                name = ""
                for key in name_keys:
                    name = str(it.get(key) or "").strip()
                    if name:
                        break
                if desc_key:
                    desc = str(it.get(desc_key) or "").strip()
                    if desc:
                        name = f"{name}{desc_sep}{desc}"
                if name:
                    lines.append(self.text.clean_inline(name))
            else:
                s = str(it).strip()
                if s:
                    lines.append(self.text.clean_inline(s))

        if lines:
            as_bullets = bool(cfg.get("bullets", True))
            sep = cfg.get("separator") or " • "

            if as_bullets:
                plain, glyph = self.bullets.get_bullet_config(sec)
                self.bullets.add_bullets(lines, sec=sec, plain=plain, glyph=glyph)
            else:
                from .docx_styles import StyleManager
                p = self.doc.add_paragraph(sep.join(lines))
                StyleManager.tight_paragraph(p, after_pt=2)

        return lines


# Section-specific renderers

class InterestsSectionRenderer(ListSectionRenderer):
    """Renders interests section."""

    def render(self, data: Dict[str, Any], sec: Optional[Dict[str, Any]] = None):
        items = data.get("interests") or []
        return self.render_simple_list(items, sec)


class LanguagesSectionRenderer(ListSectionRenderer):
    """Renders languages section with proficiency levels."""

    def render(self, data: Dict[str, Any], sec: Optional[Dict[str, Any]] = None):
        items = data.get("languages") or []
        return self.render_simple_list(
            items,
            sec,
            name_keys=("name", "language", "title"),
            desc_key="level",
            desc_sep=" — ",
        )


class CourseworkSectionRenderer(ListSectionRenderer):
    """Renders coursework section."""

    def render(self, data: Dict[str, Any], sec: Optional[Dict[str, Any]] = None):
        items = data.get("coursework") or []
        return self.render_simple_list(
            items,
            sec,
            name_keys=("name", "course", "title"),
            desc_key="desc",
            desc_sep=" — ",
        )


class CertificationsSectionRenderer(ListSectionRenderer):
    """Renders certifications section."""

    def render(self, data: Dict[str, Any], sec: Optional[Dict[str, Any]] = None):
        items = data.get("certifications") or []
        return self.render_simple_list(
            items,
            sec,
            name_keys=("name", "title", "cert"),
            desc_key="year",
            desc_sep=" — ",
        )


class PresentationsSectionRenderer(ListSectionRenderer):
    """Renders presentations/talks section."""

    def render(self, data: Dict[str, Any], sec: Optional[Dict[str, Any]] = None):
        items_raw = data.get("presentations") or []
        lines: List[str] = []

        for it in items_raw:
            if isinstance(it, dict):
                title = str(it.get("title") or it.get("name") or "").strip()
                event = str(it.get("event") or "").strip()
                year = str(it.get("year") or "").strip()
                link = str(it.get("link") or "").strip()

                parts = [p for p in [title or event, event if title else "", year] if p]
                line = " — ".join(parts)
                if link:
                    line = f"{line} ({link})" if line else link
                if line:
                    lines.append(self.text.clean_inline(line))
            else:
                s = str(it).strip()
                if s:
                    lines.append(self.text.clean_inline(s))

        if lines:
            plain, glyph = self.bullets.get_bullet_config(sec)
            self.bullets.add_bullets(lines, sec=sec, plain=plain, glyph=glyph)

        return lines
