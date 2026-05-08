"""DOCX resume writer.

Renders resume data to DOCX format using templates and styling configuration.

This module provides backward-compatible entry points for resume generation.
For new code, prefer using the class-based API:

    from resume.docx_base import create_resume_writer
    writer = create_resume_writer(data, template)
    writer.write(out_path)
"""
from __future__ import annotations

from typing import Any, Dict, List, Optional

from .io_utils import safe_import
from .docx_styles import (
    _parse_hex_color,
    _tight_paragraph,
    _flush_left,
    _apply_paragraph_shading,
    _format_phone_display,
    _format_link_display,
)
from .docx_sections import BulletRenderer, HeaderRenderer
from .docx_standard import SECTION_RENDERERS, SECTIONS_WITH_KEYWORDS  # re-export
from .render_config import HeaderLineConfig, BulletConfig


SECTION_SYNONYMS = {
    "summary": {"summary", "profile", "about"},
    "skills": {"skills", "technical skills"},
    "technologies": {"technologies", "technology", "tools"},
    "experience": {"experience", "work history", "employment"},
    "education": {"education", "academics"},
}


def _match_section_key(title: str) -> Optional[str]:
    t = title.strip().lower()
    for key, names in SECTION_SYNONYMS.items():
        if t in names:
            return key
    return None


# Backward-compatible function aliases that delegate to renderers
def _bold_keywords(paragraph, text: str, keywords: List[str]):
    """Bold keywords in paragraph text."""
    renderer = BulletRenderer.__new__(BulletRenderer)
    renderer._bold_keywords(paragraph, text, keywords)


def _add_bullet_line(doc, text: str, *, keywords: List[str] | None = None, glyph: str = "•"):
    renderer = BulletRenderer(doc)
    return renderer.add_bullet_line(text, keywords=keywords, glyph=glyph)


def _add_plain_bullet(doc, text: str, keywords: List[str] | None = None):
    return _add_bullet_line(doc, text, keywords=keywords, glyph="•")


def _add_bullets(
    doc,
    items: List[str],
    *,
    keywords: List[str] | None = None,
    cfg: BulletConfig | None = None,
):
    c = cfg or BulletConfig()
    renderer = BulletRenderer(doc)
    renderer.add_bullets(items, keywords=keywords, plain=c.plain, glyph=c.glyph, list_style=c.list_style)


def _render_group_title(doc, title: str, sec: Dict[str, Any] | None = None):
    renderer = HeaderRenderer(doc)
    return renderer.add_group_title(title, sec)


def _add_header_line(
    doc,
    cfg: HeaderLineConfig | None = None,
    *,
    sec: Dict[str, Any] | None = None,
):
    c = cfg or HeaderLineConfig()
    renderer = HeaderRenderer(doc)
    return renderer.add_header_line(
        title_text=c.title_text, company_text=c.company_text,
        loc_text=c.loc_text, span_text=c.span_text,
        sec=sec, style=c.style,
    )


def _add_named_bullet(
    doc,
    name_text: str,
    desc_text: str,
    *,
    sec: Dict[str, Any] | None = None,
    cfg: BulletConfig | None = None,
):
    c = cfg or BulletConfig()
    renderer = BulletRenderer(doc)
    return renderer.add_named_bullet(name_text, desc_text, sec=sec, glyph=c.glyph, sep=c.sep)


def _get_header_level(sec: Dict[str, Any] | None, page_cfg: Dict[str, Any] | None) -> int:
    try:
        if sec and isinstance(sec.get("header_level"), int):
            return int(sec.get("header_level"))
        if page_cfg and isinstance(page_cfg.get("header_level"), int):
            return int(page_cfg.get("header_level"))
    except Exception:  # nosec B110 - invalid header_level
        pass
    return 1


def _use_plain_bullets(sec: Dict[str, Any] | None, page_cfg: Dict[str, Any] | None) -> tuple:
    renderer = BulletRenderer.__new__(BulletRenderer)
    renderer.page_cfg = page_cfg or {}
    return renderer.get_bullet_config(sec)


def _extract_experience_locations(data: Dict[str, Any]) -> List[str]:
    """Extract unique location strings from experience entries."""
    locs = [str(e.get("location") or "").strip() for e in (data.get("experience") or [])]
    return list(dict.fromkeys([loc for loc in locs if loc]))


def _get_contact_field(data: Dict[str, Any], field: str) -> str:
    """Get a contact field from data or nested contact dict."""
    contact = data.get("contact") or {}
    return data.get(field) or contact.get(field) or ""


def _collect_link_extras(data: Dict[str, Any]) -> List[str]:
    """Collect formatted link extras (website, linkedin, github, links list)."""
    extras = []
    for field in ["website", "linkedin", "github"]:
        val = _get_contact_field(data, field)
        if val:
            extras.append(_format_link_display(val))
    links_list = _get_contact_field(data, "links") or []
    for val in (links_list if isinstance(links_list, list) else []):
        if isinstance(val, str) and val.strip():
            extras.append(_format_link_display(val))
    return extras


def _apply_page_styles(doc, page_cfg: Dict[str, Any]) -> None:
    """Apply compact page styles (margins and fonts)."""
    from .docx_base import apply_page_styles_to_doc
    apply_page_styles_to_doc(doc, page_cfg)


def _set_document_metadata(doc, data: Dict[str, Any], template: Dict[str, Any]) -> None:
    """Set document core properties (title, author, keywords)."""
    from .docx_base import set_document_metadata_on_doc
    page_cfg = template.get("page") or {}
    set_document_metadata_on_doc(doc, data, page_cfg)


def _center_paragraph(para) -> None:
    """Center a paragraph and remove indents."""
    from .docx_styles import StyleManager
    StyleManager.center_paragraph(para)


def _render_document_header(doc, data: Dict[str, Any]) -> None:
    """Render the name, headline, and contact line at the top of the resume."""
    name = _get_contact_field(data, "name")
    headline = _get_contact_field(data, "headline")
    email = _get_contact_field(data, "email")
    phone = _get_contact_field(data, "phone")
    display_phone = _format_phone_display(phone) if phone else ""
    location = _get_contact_field(data, "location")

    if name:
        doc.add_heading(name, level=0)
        _tight_paragraph(doc.paragraphs[-1], after_pt=2)
        _center_paragraph(doc.paragraphs[-1])

    if headline:
        p_head = doc.add_paragraph(str(headline))
        _tight_paragraph(p_head, after_pt=2)
        _center_paragraph(p_head)

    subtitle_parts = [p for p in [email, display_phone, location] if p]
    subtitle_parts.extend(_collect_link_extras(data))

    if subtitle_parts:
        p = doc.add_paragraph(" | ".join(subtitle_parts))
        _tight_paragraph(p, after_pt=6)
        _center_paragraph(p)


def _resolve_sections(template: Dict[str, Any], structure: Optional[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Resolve section order and configuration from template/structure."""
    sections = template.get("sections") or []

    if structure and isinstance(structure.get("order"), list):
        order_keys: List[str] = structure.get("order", [])
        key_to_title: Dict[str, str] = structure.get("titles", {})
        tpl_by_key = {s.get("key"): s for s in sections if s.get("key")}
        sections = [
            {**tpl_by_key.get(k, {"key": k, "title": key_to_title.get(k, k.title())})}
            for k in order_keys
            if k in tpl_by_key or key_to_title.get(k)
        ]

    return sections


def _render_section_heading(doc, title: str, template: Dict[str, Any]) -> None:
    """Render a section heading with optional shading."""
    if not title:
        return
    doc.add_heading(str(title), level=1)
    _tight_paragraph(doc.paragraphs[-1], before_pt=6, after_pt=2)
    _flush_left(doc.paragraphs[-1])
    page_h1_bg = (template.get("page") or {}).get("h1_bg") or (template.get("page") or {}).get("heading_bg")
    bg_rgb = _parse_hex_color(page_h1_bg)
    if bg_rgb:
        _apply_paragraph_shading(doc.paragraphs[-1], bg_rgb)


def write_resume_docx(
    data: Dict[str, Any],
    template: Dict[str, Any],
    out_path: str,
    seed: Optional[Dict[str, Any]] = None,
    structure: Optional[Dict[str, Any]] = None,
) -> None:
    """Write resume to DOCX format.

    This is the main entry point for backward compatibility.
    For new code, prefer using create_resume_writer() from docx_base.

    Args:
        data: Resume data (name, experience, education, etc.)
        template: Template configuration (sections, page styles, etc.)
        out_path: Output file path
        seed: Optional seed data (keywords, etc.)
        structure: Optional structure override for section order
    """
    # Check for sidebar layout and delegate to appropriate writer
    layout_cfg = template.get("layout") or {}
    if layout_cfg.get("type") == "sidebar":
        from .docx_sidebar import write_resume_docx_sidebar
        return write_resume_docx_sidebar(data, template, out_path, seed)

    # Standard single-column layout
    docx = safe_import("docx")
    if not docx:
        raise RuntimeError("Rendering DOCX requires python-docx; install python-docx.")

    from docx import Document  # type: ignore

    doc = Document()
    page_cfg = template.get("page") or {}

    # Apply page styles, metadata, and header
    _apply_page_styles(doc, page_cfg)
    _set_document_metadata(doc, data, template)
    _render_document_header(doc, data)

    # Extract keywords from seed
    keywords = []
    if seed and isinstance(seed.get("keywords"), list):
        keywords = [str(k) for k in seed.get("keywords", [])]

    sections = _resolve_sections(template, structure)
    _render_sections(doc, template, data, sections, keywords)
    doc.save(out_path)


def _render_sections(doc, template, data, sections, keywords) -> None:
    """Render all sections into the document."""
    page_cfg = template.get("page") or {}
    for sec in sections:
        key = sec.get("key")
        if not key:
            continue
        title = sec.get("title") or (key.title() if isinstance(key, str) else "")
        _render_section_heading(doc, title, template)
        renderer_class = SECTION_RENDERERS.get(key)
        if renderer_class:
            renderer = renderer_class(doc, page_cfg)
            if key in SECTIONS_WITH_KEYWORDS:
                renderer.render(data, sec, keywords)
            else:
                renderer.render(data, sec)
