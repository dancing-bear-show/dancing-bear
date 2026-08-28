"""Empty-section suppression on the ``write_resume_docx`` path.

``cmd_render`` renders through ``write_resume_docx`` -> ``_render_sections``,
while the class-based API renders through ``create_resume_writer`` ->
``StandardResumeWriter._render_section``. Both build the *same* standard
single-column layout, so any behavioural difference between them is a bug in
whichever one the tests do not cover.

Only the class-based path used to consult ``_section_has_data``. The CLI path
emitted a heading for every configured section unconditionally, so a resume
with no ``experience`` still rendered an "Experience" heading over nothing.
That is why the empty-scalar-summary fix could pass its own tests while the
production CLI kept rendering the bare heading it was meant to remove.

The equivalence test below is the load-bearing one: it pins the two entry
points to each other, so a guard added to one and not the other fails here
rather than shipping.
"""
from __future__ import annotations

import os
import tempfile
import unittest


def _docx_available() -> bool:
    """Check if python-docx is installed."""
    try:
        import docx  # noqa: F401

        return True
    except ImportError:
        return False


# Sections the templates below configure. Every one is a registered renderer,
# so anything omitted from the rendered output was dropped by the data guard
# rather than by the missing-renderer check.
_TEMPLATE = {
    "sections": [
        {"key": "summary", "title": "Summary"},
        {"key": "skills", "title": "Skills"},
        {"key": "experience", "title": "Experience"},
        {"key": "education", "title": "Education"},
        {"key": "interests", "title": "Interests"},
    ]
}


class _RenderHelpers:
    """Rendering helpers shared by the cases below.

    Kept off ``TestCase`` deliberately: subclassing a populated ``TestCase`` to
    reuse its helpers would re-run every one of its test methods under the
    subclass name too.
    """

    def _headings(self, path: str) -> list[str]:
        """Return the text of every heading paragraph in the rendered file."""
        from docx import Document

        return [
            p.text.strip()
            for p in Document(path).paragraphs
            if p.style.name.startswith("Heading") and p.text.strip()
        ]

    def _render_via_writer(self, data: dict, template: dict) -> list[str]:
        """Render through the CLI path and return its headings."""
        from resume.docx_writer import write_resume_docx
        from resume.schema import Resume

        out = os.path.join(tempfile.mkdtemp(), "writer.docx")
        write_resume_docx(Resume.from_dict(data), template, out)
        return self._headings(out)

    def _render_via_factory(self, data: dict, template: dict) -> list[str]:
        """Render through the class-based path and return its headings."""
        from resume.docx_base import create_resume_writer
        from resume.schema import Resume

        out = os.path.join(tempfile.mkdtemp(), "factory.docx")
        create_resume_writer(Resume.from_dict(data), template).write(out)
        return self._headings(out)


@unittest.skipUnless(_docx_available(), "python-docx not installed")
class TestWriteResumeDocxSkipsEmptySections(_RenderHelpers, unittest.TestCase):
    """``write_resume_docx`` must not emit headings for sections with no data."""

    def test_name_only_resume_renders_no_section_headings(self):
        """A resume carrying only a name renders the name and nothing else.

        Every configured section is empty, so every heading must be suppressed.
        Before the fix this returned all five headings over empty bodies.
        """
        headings = self._render_via_writer({"name": "Casey Synthetic"}, _TEMPLATE)
        self.assertEqual(headings, [])

    def test_populated_sections_still_render_heading_and_body(self):
        """Sections with data are untouched: heading *and* body both survive.

        Guards the obvious over-correction — suppressing empty sections must
        not suppress populated ones.
        """
        from docx import Document
        from resume.docx_writer import write_resume_docx
        from resume.schema import Resume

        data = {
            "name": "Casey Synthetic",
            "summary": ["Reliability engineer."],
            "skills": ["Python", "Go"],
            "experience": [
                {
                    "company": "Example Corp",
                    "title": "Site Reliability Engineer",
                    "bullets": ["Reduced page volume."],
                }
            ],
        }
        out = os.path.join(tempfile.mkdtemp(), "populated.docx")
        write_resume_docx(Resume.from_dict(data), _TEMPLATE, out)

        self.assertEqual(
            self._headings(out), ["Summary", "Skills", "Experience"]
        )

        # The bodies must be present too, not just the headings.
        body = "\n".join(p.text for p in Document(out).paragraphs)
        self.assertIn("Reliability engineer", body)
        self.assertIn("Python", body)
        self.assertIn("Reduced page volume", body)

    def test_empty_sections_are_dropped_from_a_partially_populated_resume(self):
        """Only the empty sections drop; ordering of the survivors is preserved."""
        data = {
            "name": "Casey Synthetic",
            "summary": ["Reliability engineer."],
            "education": [{"school": "Example University", "degree": "BSc"}],
        }
        self.assertEqual(
            self._render_via_writer(data, _TEMPLATE), ["Summary", "Education"]
        )


@unittest.skipUnless(_docx_available(), "python-docx not installed")
class TestEntryPointsAgree(_RenderHelpers, unittest.TestCase):
    """The two entry points must render the same headings for the same input.

    This is the regression that let the bug through: the guard existed on one
    path only, and no test compared the two. Asserting equality rather than a
    hardcoded list means a future guard change has to be applied to both paths
    to keep this green.
    """

    def _assert_entry_points_agree(self, data: dict, template: dict = _TEMPLATE):
        via_writer = self._render_via_writer(data, template)
        via_factory = self._render_via_factory(data, template)
        self.assertEqual(
            via_writer,
            via_factory,
            "write_resume_docx and create_resume_writer disagree: "
            f"{via_writer!r} != {via_factory!r}",
        )
        return via_writer

    def test_entry_points_agree_on_a_name_only_resume(self):
        """Both paths suppress every empty section."""
        self.assertEqual(self._assert_entry_points_agree({"name": "Casey Synthetic"}), [])

    def test_entry_points_agree_on_a_populated_resume(self):
        """Both paths render the same populated sections in the same order."""
        data = {
            "name": "Casey Synthetic",
            "summary": ["Reliability engineer."],
            "skills": ["Python"],
            "interests": ["Cycling"],
        }
        self.assertEqual(
            self._assert_entry_points_agree(data), ["Summary", "Skills", "Interests"]
        )

    def test_entry_points_agree_on_an_empty_scalar_summary(self):
        """The #290 case, now pinned on the path the CLI actually uses.

        A scalar ``summary: ""`` normalizes to ``[PriorityItem(text='')]``,
        which is truthy; ``_data_key_is_truthy`` is what still reports it empty.
        """
        data = {"name": "Casey Synthetic", "summary": "", "skills": ["Python"]}
        self.assertEqual(self._assert_entry_points_agree(data), ["Skills"])


@unittest.skipUnless(_docx_available(), "python-docx not installed")
class TestUnrecognizedSectionKeysStillRender(unittest.TestCase):
    """``_section_has_data`` returns True for keys outside its data-key map.

    Every registered renderer is currently mapped in ``_SECTION_DATA_KEYS``, so
    this fallback is unreachable through the shipped registry. It exists for
    renderers added later without a matching data-key entry, and the intended
    behaviour is to *render* them rather than silently swallow their content.
    Registering a temporary renderer is what makes that reachable, so this pins
    real behaviour instead of asserting a vacuous truth.
    """

    def test_section_without_a_data_key_mapping_still_renders(self):
        from unittest.mock import patch

        from resume.docx_standard import SECTION_RENDERERS
        from resume.docx_writer import write_resume_docx
        from resume.schema import Resume

        rendered: list[str] = []

        class _ProjectsRenderer:
            def __init__(self, doc, page_cfg):
                self.doc = doc

            def render(self, resume, sec):
                rendered.append("called")
                self.doc.add_paragraph("Synthetic project body")

        patched = dict(SECTION_RENDERERS)
        patched["projects"] = _ProjectsRenderer

        template = {"sections": [{"key": "projects", "title": "Projects"}]}
        out = os.path.join(tempfile.mkdtemp(), "unmapped.docx")

        with patch.dict(
            "resume.docx_writer.SECTION_RENDERERS", patched, clear=True
        ):
            write_resume_docx(Resume.from_dict({"name": "Casey Synthetic"}), template, out)

        from docx import Document

        text = "\n".join(p.text for p in Document(out).paragraphs)
        self.assertEqual(rendered, ["called"], "unmapped section renderer was skipped")
        self.assertIn("Projects", text)
        self.assertIn("Synthetic project body", text)

    def test_section_with_no_registered_renderer_is_skipped(self):
        """An unknown key with no renderer emits no heading, matching the factory.

        This is the other half of the divergence: the old CLI path emitted the
        heading *before* looking the renderer up, leaving a heading with no
        body for a key like "projects" that nothing can draw.
        """
        from resume.docx_writer import write_resume_docx
        from resume.schema import Resume

        template = {"sections": [{"key": "projects", "title": "Projects"}]}
        out = os.path.join(tempfile.mkdtemp(), "norenderer.docx")
        write_resume_docx(Resume.from_dict({"name": "Casey Synthetic"}), template, out)

        from docx import Document

        text = "\n".join(p.text for p in Document(out).paragraphs)
        self.assertNotIn("Projects", text)


if __name__ == "__main__":
    unittest.main()
