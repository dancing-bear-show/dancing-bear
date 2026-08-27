"""Base class for DOCX resume writers.

Provides common functionality for resume rendering.
"""
from __future__ import annotations

from abc import ABC, abstractmethod
from typing import TYPE_CHECKING, Any

if TYPE_CHECKING:  # pragma: no cover - typing only
    from docx.document import Document

from docx.shared import Pt, Inches, RGBColor  # type: ignore

from .io_utils import safe_import
from .docx_styles import (
    StyleManager,
    TextFormatter,
    _parse_hex_color,
    _is_dark,
    _format_link_display,
)
from .docx_links import normalize_link_url
from .schema import Resume


STYLE_HEADING_1 = "Heading 1"


def _apply_h1_color(doc, h1_color, h1_bg) -> None:
    """Apply color to Heading 1 style, auto-contrasting if needed."""
    rgb = _parse_hex_color(h1_color)
    bg = _parse_hex_color(h1_bg)
    if (not rgb) and bg:
        rgb = (255, 255, 255) if _is_dark(bg) else (0, 0, 0)
    if rgb:
        doc.styles[STYLE_HEADING_1].font.color.rgb = RGBColor(*rgb)


def _apply_font_styles(doc, page_cfg: dict[str, Any]) -> None:
    """Apply font sizes and colors to Normal, Heading 1, and Title styles."""
    body_pt = float(page_cfg.get("body_pt", 10.5))
    h1_pt = float(page_cfg.get("h1_pt", 12))
    title_pt = float(page_cfg.get("title_pt", 14))
    h1_color = page_cfg.get("h1_color") or page_cfg.get("heading_color")
    h1_bg = page_cfg.get("h1_bg") or page_cfg.get("heading_bg")
    title_color = page_cfg.get("title_color")

    doc.styles["Normal"].font.size = Pt(body_pt)

    if STYLE_HEADING_1 in doc.styles:
        doc.styles[STYLE_HEADING_1].font.size = Pt(h1_pt)
        doc.styles[STYLE_HEADING_1].font.bold = True
        _apply_h1_color(doc, h1_color, h1_bg)

    if "Title" in doc.styles:
        doc.styles["Title"].font.size = Pt(title_pt)
        doc.styles["Title"].font.bold = True
        rgbt = _parse_hex_color(title_color)
        if rgbt:
            doc.styles["Title"].font.color.rgb = RGBColor(*rgbt)


def apply_page_styles_to_doc(doc, page_cfg: dict[str, Any]) -> None:
    """Apply compact page styles (margins and fonts) to a document object.

    Shared implementation used by both ResumeWriterBase and legacy module helpers.
    """
    if not page_cfg.get("compact"):
        return

    try:
        sec = doc.sections[0]
        m = float(page_cfg.get("margins_in", 0.5))
        sec.top_margin = Inches(m)
        sec.bottom_margin = Inches(m)
        sec.left_margin = Inches(m)
        sec.right_margin = Inches(m)
    except Exception:  # nosec B110 - non-critical margin setting failure
        pass

    try:
        _apply_font_styles(doc, page_cfg)
    except Exception:  # nosec B110 - non-critical font/style setting failure
        pass


def _extract_locations(data: dict[str, Any]) -> list[str]:
    """Extract unique non-empty location strings from experience entries."""
    locs = [str(e.get("location") or "").strip() for e in (data.get("experience") or [])]
    return list(dict.fromkeys([loc for loc in locs if loc]))


def _set_category(cp, locs: list[str]) -> None:
    """Set category on core properties, silently ignoring failures."""
    try:
        cp.category = "; ".join(locs)
    except Exception:  # nosec B110 - non-critical category metadata setting
        pass


_OPC_PROPERTY_MAX = 255


def _fit_property(items: list[str], sep: str = "; ") -> str:
    """Join items into an OPC core property, dropping any that overflow.

    OPC caps core properties at 255 characters and python-docx raises
    ValueError rather than truncating, so an over-long keyword list would
    abort metadata assignment entirely.
    """
    out = ""
    for item in items:
        candidate = f"{out}{sep}{item}" if out else str(item)
        if len(candidate) > _OPC_PROPERTY_MAX:
            break
        out = candidate
    return out


def _metadata_title(
    page_cfg: dict[str, Any], name: str, contact_parts: list[str], include_pii: bool
) -> str:
    """Build the core-properties title, honoring an override and the PII flag."""
    override = page_cfg.get("metadata_title")
    if override:
        return str(override)
    if include_pii:
        contact_line = " | ".join([p for p in contact_parts if p])
        return " - ".join([p for p in [name, contact_line] if p]) or "Resume"
    return " - ".join([p for p in [name, "Resume"] if p]) or "Resume"


def _metadata_keywords(
    cp,
    data: dict[str, Any],
    page_cfg: dict[str, Any],
    identity_parts: list[str],
    include_pii: bool,
) -> list[str]:
    """Build the keyword list, honoring an override and the PII flag.

    Also sets the category property when locations are included.
    """
    override = page_cfg.get("metadata_keywords")
    if override:
        if isinstance(override, list):
            return [str(k).strip() for k in override]
        return [str(override)]

    name = identity_parts[0] if identity_parts else ""
    kw = [k for k in identity_parts if k] if include_pii else [k for k in [name] if k]
    if bool(page_cfg.get("metadata_include_locations", True)):
        uniq_locs = _extract_locations(data)
        kw.extend(uniq_locs)
        if uniq_locs:
            _set_category(cp, uniq_locs)
    return kw


def _identity_fields(data: dict[str, Any]) -> tuple[str, str, str, str]:
    """Derive (name, email, phone, location), falling back to the nested contact dict."""
    contact = data.get("contact") or {}
    name = data.get("name") or ""
    email = data.get("email") or contact.get("email") or ""
    phone = data.get("phone") or contact.get("phone") or ""
    location = data.get("location") or contact.get("location") or ""
    return name, email, phone, location


def _stamp_metadata_dates(cp, page_cfg: dict[str, Any], name: str) -> None:
    """Stamp created/modified/last_modified_by to now, unless the caller opts out.

    python-docx ships a 2013 default timestamp in its template. Left alone, a
    freshly generated resume advertises a decade-old created date.
    """
    if not bool(page_cfg.get("metadata_stamp_dates", True)):
        return
    from datetime import datetime, timezone
    now = datetime.now(timezone.utc).replace(tzinfo=None)
    cp.created = now
    cp.modified = now
    cp.last_modified_by = name or cp.last_modified_by


def set_document_metadata_on_doc(
    doc, data: dict[str, Any], page_cfg: dict[str, Any]
) -> None:
    """Set document core properties (title, author, keywords).

    Shared implementation used by both ResumeWriterBase and legacy module helpers.
    """
    try:
        name, email, phone, location = _identity_fields(data)
        cp = doc.core_properties

        # metadata_pii: when False, keep contact details out of core properties.
        # Metadata survives PDF export and is read by ATS and anyone opening
        # Properties, so duplicating a phone number there exposes it in a field
        # the author cannot see. The body already carries the contact line.
        include_pii = bool(page_cfg.get("metadata_pii", True))

        cp.title = _metadata_title(page_cfg, name, [email, phone, location], include_pii)
        cp.subject = str(page_cfg.get("metadata_subject") or "Resume")
        if name:
            cp.author = name

        comments = page_cfg.get("metadata_comments")
        if comments is not None:
            cp.comments = str(comments)

        kw = _metadata_keywords(
            cp, data, page_cfg, [name, email, phone, location], include_pii
        )
        # OPC caps core properties at 255 chars and python-docx raises rather
        # than truncating. Trim on a separator boundary so the field degrades
        # to fewer whole keywords instead of being dropped entirely — the
        # enclosing except would otherwise swallow every later assignment too.
        cp.keywords = _fit_property(kw)

        _stamp_metadata_dates(cp, page_cfg, name)
    except Exception:  # nosec B110 - non-critical metadata setting failure
        pass


class ResumeWriterBase(ABC):
    """Base class for DOCX resume writers."""

    def __init__(self, data: dict[str, Any] | Resume, template: dict[str, Any]):
        """Initialize writer with resume data and template config.

        Accepts either a typed ``Resume`` or a raw dict. Both are kept: the
        migrated section renderers read ``self.resume``, while the modules that
        still consume candidate data as a mapping read ``self.data``. A dict
        argument is lifted through ``Resume.from_dict``, which is safe because
        the conversion is idempotent on data that has already been through the
        schema -- and everything reaching a writer has, since the entry point
        normalizes on the way in.

        Args:
            data: Resume data, typed or as a dict.
            template: Template configuration (sections, page styles, etc.)
        """
        self.resume = data if isinstance(data, Resume) else Resume.from_dict(data)
        self.data = self.resume.to_dict()
        self.template = template
        self.page_cfg = template.get("page") or {}
        self.layout_cfg = template.get("layout") or {}
        # Annotated because `write()` reassigns this to a docx Document. Without
        # an annotation mypy infers the attribute's type from this assignment
        # alone — None — and then reports every later
        # `self.doc.add_heading(...)` in the subclasses as an attribute error on
        # None. The import is TYPE_CHECKING-only because `docx.document` is the
        # class's real home, while runtime code constructs it via the
        # `docx.Document()` factory.
        self.doc: Document | None = None
        self.styles = StyleManager()
        self.text = TextFormatter()

    def write(self, out_path: str, seed: dict[str, Any] | None = None) -> None:
        """Write resume to DOCX file.

        Args:
            out_path: Output file path
            seed: Optional seed data (keywords, etc.)
        """
        docx = safe_import("docx")
        if not docx:
            raise RuntimeError("Rendering DOCX requires python-docx; install python-docx.")

        from docx import Document  # type: ignore
        self.doc = Document()

        self._apply_page_styles()
        self._set_document_metadata()
        self._render_content(seed)
        self.doc.save(out_path)

    @abstractmethod
    def _render_content(self, seed: dict[str, Any] | None = None) -> None:
        """Render the main document content. Subclasses must implement."""
        pass

    # -------------------------------------------------------------------------
    # Page setup and metadata
    # -------------------------------------------------------------------------

    def _apply_page_styles(self) -> None:
        """Apply compact page styles (margins and fonts)."""
        apply_page_styles_to_doc(self.doc, self.page_cfg)

    def _set_document_metadata(self) -> None:
        """Set document core properties (title, author, keywords)."""
        set_document_metadata_on_doc(self.doc, self.data, self.page_cfg)

    def _extract_experience_locations(self) -> list[str]:
        """Extract unique location strings from experience entries."""
        locs = [str(e.get("location") or "").strip() for e in (self.data.get("experience") or [])]
        return list(dict.fromkeys([loc for loc in locs if loc]))

    # -------------------------------------------------------------------------
    # Contact field helpers
    # -------------------------------------------------------------------------

    def _get_contact_field(self, field: str) -> str:
        """Get a contact field from data or nested contact dict."""
        contact = self.data.get("contact") or {}
        return self.data.get(field) or contact.get(field) or ""

    def _collect_link_extras(self) -> list[str]:
        """Collect formatted link extras (website, linkedin, github, links list)."""
        return [display for display, _url in self._collect_link_extra_items()]

    def _collect_link_extra_items(self) -> list[tuple[str, str]]:
        """Collect link extras as (display, url) pairs for hyperlink rendering."""
        items: list[tuple[str, str]] = []
        for field in ["website", "linkedin", "github"]:
            val = self._get_contact_field(field)
            if val:
                items.append((_format_link_display(val), normalize_link_url(val)))
        links_list = self._get_contact_field("links") or []
        for val in (links_list if isinstance(links_list, list) else []):
            if isinstance(val, str) and val.strip():
                items.append((_format_link_display(val), normalize_link_url(val)))
        return items

    # -------------------------------------------------------------------------
    # Paragraph helpers
    # -------------------------------------------------------------------------

    def _center_paragraph(self, para) -> None:
        """Center a paragraph and remove indents."""
        self.styles.center_paragraph(para)

    def _add_colored_run(self, paragraph, text: str, hex_color: str | None, **kwargs) -> Any:
        """Add a run with optional color and formatting."""
        run = paragraph.add_run(text)
        if hex_color:
            rgb = _parse_hex_color(hex_color)
            if rgb:
                run.font.color.rgb = RGBColor(*rgb)
        for key, val in kwargs.items():
            setattr(run, key, val) if hasattr(run, key) else setattr(run.font, key, val)
        return run


def create_resume_writer(
    data: dict[str, Any] | Resume,
    template: dict[str, Any],
) -> ResumeWriterBase:
    """Factory function to create the appropriate resume writer.

    Args:
        data: Resume data (name, experience, education, etc.)
        template: Template configuration (sections, page styles, layout type)

    Returns:
        ResumeWriterBase subclass instance (StandardResumeWriter or SidebarResumeWriter)

    Examples:
        >>> # Standard single-column layout
        >>> data = {"name": "John Doe", "experience": [...]}
        >>> template = {"sections": [...], "page": {"compact": True}}
        >>> writer = create_resume_writer(data, template)
        >>> writer.write("resume.docx")

        >>> # Sidebar layout (two-column with repeating header)
        >>> template = {
        ...     "layout": {"type": "sidebar", "sidebar_width": 2.5},
        ...     "sections": [...],
        ...     "page": {"compact": True}
        ... }
        >>> writer = create_resume_writer(data, template)
        >>> writer.write("resume_sidebar.docx")
    """
    layout_cfg = template.get("layout") or {}
    layout_type = layout_cfg.get("type", "standard")

    if layout_type == "sidebar":
        from .docx_sidebar import SidebarResumeWriter
        return SidebarResumeWriter(data, template)
    else:
        from .docx_standard import StandardResumeWriter
        return StandardResumeWriter(data, template)
