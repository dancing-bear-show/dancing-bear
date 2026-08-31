"""Section-level rendering for DOCX sidebar resume layout.

Extracted from docx_sidebar. Provides:
  - _render_main_education, _render_exp_entry, _render_main_experience
  - _render_main_teaching, _render_pres_entry, _render_main_presentations
  - write_resume_docx_sidebar: backward-compatible function
  - SidebarResumeWriter: two-column sidebar layout writer class
"""
from __future__ import annotations

from typing import Any

from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

from .docx_base import ResumeWriterBase
from .schema import ExperienceEntry, PriorityItem, Presentation, Resume, _Item
from .docx_styles import (
    TextFormatter,
    _apply_paragraph_shading,
    _parse_hex_color,
    _tight_paragraph,
)
from .render_config import CenteredHeaderLineStyle, IndentedRunStyle
from .docx_sidebar_cells import (
    _add_indented_run,
    _remove_cell_borders,
    _render_edu_meta,
    _render_main_section_heading,
    _render_sidebar_section,
    _set_cell_shading,
)


def _add_colored_bullet_run(p, bullet_color: str) -> None:
    """Add a bullet glyph run to a paragraph, colored if a valid hex color is given."""
    bullet_run = p.add_run("• ")
    bullet_rgb = _parse_hex_color(bullet_color)
    if bullet_rgb:
        bullet_run.font.color.rgb = RGBColor(*bullet_rgb)


def _item_text(item: _Item) -> str:
    """Return a sidebar item's display text: its own primary field.

    These renderers do not share one primary field -- summary entries and
    experience bullets are ``PriorityItem`` and carry their prose in ``text``,
    while skill items are ``SkillGroupItem`` and carry it in ``name`` -- so the
    field is read off the item's own type rather than hardcoded.

    Reading the resolved field is what makes alias-keyed entries render. The
    pre-migration code was ``x.get("text", x)``: it honoured ONE literal key and
    otherwise fell back to the whole mapping, which python-docx stringified by
    iterating -- yielding the concatenation of exactly the keys that entry
    carried, in input order. So the garbage string identifies the whole input
    shape, not just the alias::

        {"line": "prose", "priority": 2}     -> "linepriority"
        {"line": "prose"}                    -> "line"
        {"label": "Cloud", "desc": "x"}      -> "labeldesc"
        {"label": "Cloud"}                   -> "label"

    Canonical spellings broke identically once the keyed read missed --
    ``{"name": "Cloud", "desc": "x"}`` on a skill item rendered "namedesc" --
    so this was never alias-specific. What the aliases add is that the schema
    already resolves them (``line``/``name`` onto ``PriorityItem.text``,
    ``title``/``label`` onto ``SkillGroupItem.name``) and the old replay threw
    that resolution away. An item with no resolved text yields ``""`` rather
    than a repr.
    """
    primary = type(item)._primary_field()
    if not primary:
        return ""
    return str(getattr(item, primary, "") or "")


def _render_main_education(cell, resume: Resume, page_cfg: dict[str, Any], sec: dict[str, Any] | None = None) -> None:  # nosec - sec kept for API compatibility
    """Render education in main column."""
    bullet_color = page_cfg.get("main_bullet_color", "#4A90A4")

    for edu in resume.education:
        degree = edu.degree
        institution = edu.institution
        year = edu.year

        p = cell.add_paragraph()
        _add_colored_bullet_run(p, bullet_color)

        deg_run = p.add_run(degree)
        deg_run.bold = True
        deg_run.font.size = Pt(page_cfg.get("body_pt", 10))
        _tight_paragraph(p, after_pt=0)
        _render_edu_meta(cell, institution, year, page_cfg)


def _render_exp_entry(cell, exp: ExperienceEntry, page_cfg: dict[str, Any], bullet_color: str, max_bullets: int) -> None:
    """Render a single experience entry."""
    title = exp.title
    company = exp.company
    start = exp.start
    end = exp.end
    # Open-ended roles read "Start – Present", matching the rest of the render
    # path. This used to hardcode the Spanish "presente": the sidebar's Spanish
    # section titles are overridable defaults, but that word was not, so an
    # English resume rendered one untranslatable Spanish token. Reusing
    # TextFormatter also normalizes end values like "now"/"current".
    span = TextFormatter.format_date_span(start, end)
    bullets = exp.bullets

    p = cell.add_paragraph()
    _add_colored_bullet_run(p, bullet_color)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(page_cfg.get("body_pt", 10))
    if span:
        span_run = p.add_run(f"  {span}")
        span_run.font.size = Pt(page_cfg.get("meta_pt", 9))
        span_rgb = _parse_hex_color("#666666")
        if span_rgb:
            span_run.font.color.rgb = RGBColor(*span_rgb)
    _tight_paragraph(p, after_pt=0)

    if company:
        p2 = cell.add_paragraph()
        p2.paragraph_format.left_indent = Inches(0.25)
        comp_run = p2.add_run(company)
        comp_run.italic = True
        comp_run.font.size = Pt(page_cfg.get("meta_pt", 9))
        comp_rgb = _parse_hex_color("#666666")
        if comp_rgb:
            comp_run.font.color.rgb = RGBColor(*comp_rgb)
        _tight_paragraph(p2, after_pt=2)

    for b in bullets[:max_bullets]:
        text = _item_text(b)
        p3 = cell.add_paragraph()
        p3.paragraph_format.left_indent = Inches(0.25)
        run = p3.add_run(text)
        run.font.size = Pt(page_cfg.get("body_pt", 10) - 1)
        _tight_paragraph(p3, after_pt=1)


def _render_main_experience(cell, resume: Resume, page_cfg: dict[str, Any], sec: dict[str, Any]) -> None:
    """Render experience in main column."""
    bullet_color = page_cfg.get("main_bullet_color", "#4A90A4")
    max_bullets = sec.get("recent_max_bullets", 3)
    for exp in resume.experience:
        _render_exp_entry(cell, exp, page_cfg, bullet_color, max_bullets)


def _render_main_teaching(cell, resume: Resume, page_cfg: dict[str, Any], sec: dict[str, Any] | None = None) -> None:  # nosec - sec kept for API compatibility
    """Render teaching in main column (uppercase titles with institution below)."""
    teaching = resume.teaching

    for item in teaching:
        text = item.get("text", item) if isinstance(item, dict) else item
        if "(" in text and text.endswith(")"):
            parts = text.rsplit("(", 1)
            title = parts[0].strip()
            institution = parts[1].rstrip(")")
        else:
            title = text
            institution = ""

        p = cell.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(page_cfg.get("body_pt", 10))
        _tight_paragraph(p, after_pt=0)

        if institution:
            p2 = cell.add_paragraph()
            run2 = p2.add_run(institution)
            run2.italic = True
            run2.font.size = Pt(page_cfg.get("meta_pt", 9))
            inst_rgb = _parse_hex_color("#666666")
            if inst_rgb:
                run2.font.color.rgb = RGBColor(*inst_rgb)
            _tight_paragraph(p2, after_pt=6)


def _render_pres_entry(cell, pres: Presentation, page_cfg: dict[str, Any], bullet_color: str) -> None:
    """Render a single presentation entry."""
    title = pres.title
    authors = pres.authors
    event = pres.event
    note = pres.note

    p = cell.add_paragraph()
    _add_colored_bullet_run(p, bullet_color)
    title_run = p.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(page_cfg.get("body_pt", 10))
    _tight_paragraph(p, after_pt=0)

    p2 = _add_indented_run(cell, authors, page_cfg, IndentedRunStyle(italic=True)) if authors else None
    p3 = _add_indented_run(cell, event, page_cfg, IndentedRunStyle(size_offset=-1, color="#888888")) if event else None

    if note:
        _add_indented_run(cell, note, page_cfg, IndentedRunStyle(italic=True, size_offset=-1, after_pt=4))
    else:
        if p3:
            p3.paragraph_format.space_after = Pt(4)
        elif p2:
            p2.paragraph_format.space_after = Pt(4)


def _render_main_presentations(cell, resume: Resume, page_cfg: dict[str, Any], sec: dict[str, Any] | None = None) -> None:  # nosec - sec kept for API compatibility
    """Render presentations/publications in main column."""
    bullet_color = page_cfg.get("main_bullet_color", "#4A90A4")
    for pres in resume.presentations:
        _render_pres_entry(cell, pres, page_cfg, bullet_color)


# Fallback usable width: US Letter (8.5") minus the default 0.5" margins.
# Only used when the real section geometry cannot be read; page size and
# margins (page.margins_in) are both configurable, so prefer the measured
# value from _usable_width_in().
_DEFAULT_USABLE_WIDTH_IN = 7.5


def _usable_width_in(doc: Any) -> float:
    """Return the section's usable width in inches (page width minus margins).

    Falls back to the US Letter default when the geometry is unreadable — a
    width guard that raised on its own measurement step would be worse than one
    that checks against a sane default.
    """
    try:
        from docx.shared import Emu

        sec = doc.sections[0]
        # Length arithmetic returns a bare int of EMUs, not a Length, so it has
        # no .inches — wrap it back up before converting.
        inches = Emu(sec.page_width - sec.left_margin - sec.right_margin).inches
    except Exception:  # nosec B110 - fall back to the documented default
        return _DEFAULT_USABLE_WIDTH_IN
    return inches if inches and inches > 0 else _DEFAULT_USABLE_WIDTH_IN


def _validate_column_width(value: Any, name: str, max_width: float | None = None) -> float:
    """Return a column width in inches, rejecting out-of-range values.

    Raises ValueError rather than clamping: a caller passing 34 means percent,
    and silently rendering a 3.4"- or 7.5"-wide column would hide the mistake
    inside a document that looks plausible until someone opens it.

    max_width defaults to US Letter's usable width; callers with a real
    document should pass _usable_width_in(doc) so configurable margins and
    page sizes are honoured.
    """
    limit = _DEFAULT_USABLE_WIDTH_IN if max_width is None else max_width
    # bool is a subclass of int, so float(True) == 1.0 — a YAML typo like
    # `sidebar_width: true` would otherwise pass validation and quietly render
    # a 1-inch column. Reject it before coercion.
    if isinstance(value, bool):
        raise ValueError(
            f"layout.{name} must be a number of INCHES, got boolean {value!r}"
        )
    try:
        width = float(value)
    except (TypeError, ValueError):
        raise ValueError(
            f"layout.{name} must be a number of INCHES, got {value!r}"
        ) from None
    if not 0 < width <= limit:
        raise ValueError(
            f"layout.{name}={width} is out of range: expected inches in "
            f"(0, {limit:g}]. Widths are inches, not percent — the usable page "
            f"width (page size minus margins) is {limit:g}\"."
        )
    return width


class SidebarResumeWriter(ResumeWriterBase):
    """Two-column sidebar layout resume writer."""

    def _render_content(self, seed: dict[str, Any] | None = None) -> None:
        """Render two-column sidebar resume content."""
        # Default widths for US Letter (8.5" wide) with 0.5" margins = 7.5" usable
        # Sidebar: 2.3" (~30%), Main: 5.2" (~70%). Measured from the section
        # rather than assumed: page size and page.margins_in are configurable,
        # and page styles are already applied by the time this runs.
        # Which width to measure against depends on whether the template
        # actually configured its margins:
        #   - page.compact set -> _apply_page_styles has already written
        #     page.margins_in, so the measured geometry is authoritative and a
        #     wide-margin template must reject columns that overflow it.
        #   - otherwise -> margins are still python-docx's 1.25" default (6.0"
        #     usable) which the template never asked for, while the documented
        #     column defaults total 7.5". Measuring there would reject the
        #     shipped defaults, so fall back to the documented width.
        if self.page_cfg.get("compact"):
            usable = _usable_width_in(self.doc)
        else:
            usable = max(_usable_width_in(self.doc), _DEFAULT_USABLE_WIDTH_IN)
        sidebar_width = _validate_column_width(
            self.layout_cfg.get("sidebar_width", 2.3), "sidebar_width", usable
        )
        main_width = _validate_column_width(
            self.layout_cfg.get("main_width", 5.2), "main_width", usable
        )
        if sidebar_width + main_width > usable:
            raise ValueError(
                f"layout.sidebar_width ({sidebar_width:g}) + layout.main_width "
                f"({main_width:g}) = {sidebar_width + main_width:g} exceeds the "
                f"{usable:g}\" usable page width (page size minus margins)."
            )
        sidebar_bg = self.layout_cfg.get("sidebar_bg")

        # Render page header (repeats on all pages)
        self._render_page_header()

        # Create two-column table for body
        table = self.doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        from docx.shared import Inches as _Inches
        table.columns[0].width = _Inches(sidebar_width)
        table.columns[1].width = _Inches(main_width)

        sidebar_cell = table.rows[0].cells[0]
        main_cell = table.rows[0].cells[1]

        _remove_cell_borders(sidebar_cell)
        _remove_cell_borders(main_cell)

        if sidebar_bg:
            _set_cell_shading(sidebar_cell, sidebar_bg)

        # Clear default paragraphs
        if sidebar_cell.paragraphs:
            sidebar_cell.paragraphs[0].clear()
        if main_cell.paragraphs:
            main_cell.paragraphs[0].clear()

        # Render sidebar content (Profile + Skills)
        self._render_sidebar_content(sidebar_cell)

        # Render main content
        self._render_main_content(main_cell)

    def _render_centered_header_line(self, p, text: str, style: CenteredHeaderLineStyle) -> None:
        """Render one centered, colored, shaded header line onto an existing paragraph."""
        run = p.add_run(text)
        run.bold = style.bold
        run.font.size = Pt(style.size_pt)
        rgb = _parse_hex_color(style.color)
        if rgb:
            run.font.color.rgb = RGBColor(*rgb)
        _tight_paragraph(p, after_pt=style.after_pt)
        self._center_paragraph(p)
        if style.bg_rgb:
            _apply_paragraph_shading(p, style.bg_rgb)

    def _render_page_header(self) -> None:
        """Add name, headline, and contact as centered header (repeats on each page)."""
        section = self.doc.sections[0]
        header = section.header

        name = self._get_contact_field("name")
        headline = self._get_contact_field("headline")
        email = self._get_contact_field("email")
        phone = self._get_contact_field("phone")
        location = self._get_contact_field("location")

        name_color = self.page_cfg.get("sidebar_name_color", "#1A365D")
        text_color = self.page_cfg.get("sidebar_text_color", "#333333")
        bg_rgb = _parse_hex_color(self.page_cfg.get("header_bg", "#F7F9FC"))

        # Name in header (centered)
        if header.paragraphs:
            p = header.paragraphs[0]
            p.clear()
        else:
            p = header.add_paragraph()
        self._render_centered_header_line(
            p, name,
            CenteredHeaderLineStyle(
                size_pt=self.page_cfg.get("sidebar_name_pt", 20),
                color=name_color, bold=True, after_pt=0, bg_rgb=bg_rgb,
            ),
        )

        # Headline (centered)
        if headline:
            self._render_centered_header_line(
                header.add_paragraph(), headline,
                CenteredHeaderLineStyle(
                    size_pt=self.page_cfg.get("sidebar_headline_pt", 10),
                    color=text_color, bold=False, after_pt=2, bg_rgb=bg_rgb,
                ),
            )

        # Contact line (centered). Link extras (website/linkedin/github/links)
        # come from the inherited ResumeWriterBase helper, so a profile that
        # shows a LinkedIn URL in one layout shows it in both.
        contact_parts = [x for x in [phone, email, location] if x]
        contact_parts.extend(self._collect_link_extras())
        if contact_parts:
            self._render_centered_header_line(
                header.add_paragraph(), " | ".join(contact_parts),
                CenteredHeaderLineStyle(
                    size_pt=self.page_cfg.get("body_pt", 10) - 1,
                    color="#666666", bold=False, after_pt=6, bg_rgb=bg_rgb,
                ),
            )

    @staticmethod
    def _normalize_summary_items(
        summary: list[PriorityItem], is_scalar: bool = False
    ) -> list[str]:
        """Normalize summary to a flat list of strings.

        Load-time normalization already collapsed the scalar-vs-list and
        str-vs-dict shapes into ``list[PriorityItem]``, so this mostly has to
        read the text off each item -- via ``_item_text``, which resolves
        alias-keyed entries to the prose behind them.

        The one shape normalization cannot speak for is an empty *scalar*
        summary. It arrives as ``[PriorityItem(text='')]``, so the naive read
        returns ``['']`` -- a truthy list -- and the caller renders a heading
        above a single empty bullet. Before the migration the scalar branch
        returned ``[]`` for an empty string and the section was suppressed, so
        ``is_scalar`` restores that.

        Only the scalar origin is special-cased. A list-form ``[{"text": ""}]``
        still returns ``['']`` and still renders, because that is what it did
        before the migration; normalization makes the two indistinguishable
        from the items alone, which is why the origin has to be passed in.
        """
        if is_scalar and not (summary and summary[0].text):
            return []
        return [_item_text(summary_item) for summary_item in summary]

    def _render_sidebar_summary(self, cell) -> None:
        """Render summary section in sidebar."""
        for sec in (self.template.get("sections") or []):
            if sec.get("key") == "summary":
                summary_items = self._normalize_summary_items(
                    self.resume.summary, self.resume.summary_is_scalar
                )
                if summary_items:
                    _render_sidebar_section(
                        cell,
                        sec.get("title", "Perfil profesional"),
                        summary_items[:6],
                        self.page_cfg,
                        bulleted=True
                    )
                break

    def _render_sidebar_skills(self, cell) -> None:
        """Render skills section in sidebar."""
        for sec in (self.template.get("sections") or []):
            if sec.get("key") == "skills":
                skill_items = [
                    _item_text(item)
                    for group in self.resume.skills_groups
                    for item in group.items
                ]
                if skill_items:
                    _render_sidebar_section(cell, sec.get("title", "Habilidades claves"), skill_items[:8], self.page_cfg)
                break

    def _render_sidebar_content(self, cell) -> None:
        """Render sidebar content (profile + skills)."""
        self._render_sidebar_summary(cell)
        self._render_sidebar_skills(cell)

    # Maps a section key to its main-column body renderer. Every renderer takes
    # (cell, resume, page_cfg, sec) so the table can call them uniformly; the ones
    # that don't need `sec` accept it as an ignored optional argument, which
    # lets them be referenced directly rather than wrapped in a pass-through
    # lambda.
    _MAIN_SECTION_RENDERERS: dict[str, Any] = {
        "education": _render_main_education,
        "experience": _render_main_experience,
        "teaching": _render_main_teaching,
        "presentations": _render_main_presentations,
    }

    def _render_main_content(self, cell) -> None:
        """Render main column content (education, experience, teaching, presentations)."""
        sections = self.template.get("sections") or []

        for sec in sections:
            renderer = self._MAIN_SECTION_RENDERERS.get(sec.get("key"))
            if renderer is None:
                continue
            _render_main_section_heading(cell, sec.get("title", ""), self.page_cfg)
            renderer(cell, self.resume, self.page_cfg, sec)


