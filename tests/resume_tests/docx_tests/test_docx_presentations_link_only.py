"""Run-level contract for the link-only presentations bullet.

A presentation entry carrying a ``link`` but no title, event or year renders
as a bare hyperlink on a bullet line. That paragraph must contain the glyph run
and the hyperlink and nothing else.

WHY THESE ASSERT ON RUNS AND NOT ON TEXT
    A zero-length run contributes no characters. ``paragraph.text`` is
    byte-identical whether or not one is present, so a text-based assertion
    passes just as happily against a paragraph carrying an empty run and pins
    nothing. The empty run is real in the document XML, and the only place it
    is observable is the run list -- which is what every assertion here reads.

    The empty run came from routing this path through ``add_bullet_line("")``,
    which unconditionally calls ``add_run(text)``. ``new_bullet_paragraph`` is
    the same shared bullet mechanism without that unconditional run.

No content from any real document appears here; every fixture value is
invented.
"""

from __future__ import annotations

import os
import tempfile
import unittest

from resume.docx_writer import write_resume_docx
from resume.schema import Resume

GLYPH = "• "

_TEMPLATE = {
    "page": {"compact": True, "body_pt": 10},
    "sections": [{"key": "presentations", "title": "Presentations"}],
}


def _resume(presentations: list) -> dict:
    return {
        "name": "Ada Placeholder",
        "headline": "Staff Reliability Engineer",
        "presentations": presentations,
    }


def _render(presentations: list) -> list:
    """Render through the public writer and return the body paragraphs."""
    import docx

    with tempfile.TemporaryDirectory() as tmp:
        out = os.path.join(tmp, "resume.docx")
        write_resume_docx(
            Resume.from_dict(_resume(presentations)), _TEMPLATE, out
        )
        return list(docx.Document(out).paragraphs)


def _bullet_paragraphs(paragraphs) -> list:
    """Every paragraph that starts with the bullet glyph."""
    return [p for p in paragraphs if p.text.startswith(GLYPH)]


def _empty_runs(paragraph) -> list[int]:
    """Indices of zero-length runs in a paragraph, for a readable failure."""
    return [i for i, r in enumerate(paragraph.runs) if r.text == ""]


class LinkOnlyPresentationRunTests(unittest.TestCase):
    """A link-only presentation emits no zero-length run."""

    def test_link_only_entry_paragraph_has_no_empty_run(self):
        paragraphs = _render([{"link": "https://example.invalid/talk"}])
        bullets = _bullet_paragraphs(paragraphs)
        self.assertEqual(
            len(bullets), 1, f"expected one bullet, got {[p.text for p in paragraphs]}"
        )

        bullet = bullets[0]
        self.assertEqual(
            _empty_runs(bullet),
            [],
            f"link-only bullet carries zero-length run(s): "
            f"{[r.text for r in bullet.runs]}",
        )

    def test_link_only_entry_paragraph_carries_only_the_glyph_run(self):
        """The glyph is the paragraph's sole run; the link lives outside them.

        ``add_hyperlink`` appends a ``w:hyperlink`` element rather than a run
        on the paragraph, so ``paragraph.runs`` sees only what the bullet
        mechanism itself added. That makes the run count an exact assertion:
        one glyph run, and nothing else.
        """
        paragraphs = _render([{"link": "https://example.invalid/talk"}])
        bullet = _bullet_paragraphs(paragraphs)[0]

        self.assertEqual(
            [r.text for r in bullet.runs],
            [GLYPH],
            "link-only bullet should carry the glyph run and nothing else",
        )

    def test_link_only_entry_still_renders_its_hyperlink(self):
        """Dropping the empty run must not drop the link with it."""
        paragraphs = _render([{"link": "https://example.invalid/talk"}])
        bullet = _bullet_paragraphs(paragraphs)[0]

        self.assertIn(
            "example.invalid/talk",
            bullet.text,
            f"hyperlink text missing from bullet: {bullet.text!r}",
        )
        xml = bullet._p.xml
        self.assertIn(
            "hyperlink", xml, "expected a w:hyperlink element on the bullet"
        )


class NormalPresentationRunTests(unittest.TestCase):
    """A presentation with real text is unaffected by the link-only fix."""

    def test_titled_entry_with_link_keeps_its_text_and_link_runs(self):
        paragraphs = _render(
            [
                {
                    "title": "Invented Talk",
                    "event": "Imaginary Conf",
                    "year": "2024",
                    "link": "https://example.invalid/talk",
                }
            ]
        )
        bullets = _bullet_paragraphs(paragraphs)
        self.assertEqual(len(bullets), 1)
        bullet = bullets[0]

        self.assertEqual(
            _empty_runs(bullet),
            [],
            f"titled bullet carries zero-length run(s): "
            f"{[r.text for r in bullet.runs]}",
        )
        self.assertEqual(
            [r.text for r in bullet.runs],
            [GLYPH, "Invented Talk — Imaginary Conf — 2024", " "],
            "titled bullet should carry glyph, display line, and link spacer",
        )
        self.assertIn("example.invalid/talk", bullet.text)

    def test_titled_entry_without_link_carries_glyph_and_text_only(self):
        paragraphs = _render(
            [{"title": "Invented Talk", "event": "Imaginary Conf", "year": "2024"}]
        )
        bullet = _bullet_paragraphs(paragraphs)[0]

        self.assertEqual(
            [r.text for r in bullet.runs],
            [GLYPH, "Invented Talk — Imaginary Conf — 2024"],
        )


if __name__ == "__main__":
    unittest.main()
