"""DOCX resume writer.

Renders resume data to DOCX format using templates and styling configuration.

This module provides backward-compatible entry points for resume generation.
For new code, prefer using the class-based API:

    from resume.docx_base import create_resume_writer
    writer = create_resume_writer(data, template)
    writer.write(out_path)
"""
from __future__ import annotations

from typing import Any

from .io_utils import safe_import
from .docx_links import (
    normalize_link_url,
    render_contact_runs as _render_contact_runs,
)
from .docx_styles import (
    _parse_hex_color,
    _tight_paragraph,
    _flush_left,
    _apply_paragraph_shading,
    _format_phone_display,
    _format_link_display,
)
from .docx_renderers import BulletRenderer
from .docx_standard import SECTION_RENDERERS, SECTIONS_WITH_KEYWORDS  # re-export


SECTION_SYNONYMS = {
    "summary": {"summary", "profile", "about"},
    "skills": {"skills", "technical skills"},
    "technologies": {"technologies", "technology", "tools"},
    "experience": {"experience", "work history", "employment"},
    "education": {"education", "academics"},
}


def _match_section_key(title: str) -> str | None:
    t = title.strip().lower()
    for key, names in SECTION_SYNONYMS.items():
        if t in names:
            return key
    return None


# Backward-compatible function aliases that delegate to renderers
def _bold_keywords(paragraph, text: str, keywords: list[str]):
    """Bold keywords in paragraph text."""
    renderer = BulletRenderer.__new__(BulletRenderer)
    renderer._bold_keywords(paragraph, text, keywords)


def _add_bullet_line(doc, text: str, *, keywords: list[str] | None = None, glyph: str = "•"):
    renderer = BulletRenderer(doc)
    return renderer.add_bullet_line(text, keywords=keywords, glyph=glyph)


def _get_header_level(sec: dict[str, Any] | None, page_cfg: dict[str, Any] | None) -> int:
    try:
        if sec and isinstance(sec.get("header_level"), int):
            return int(sec.get("header_level"))
        if page_cfg and isinstance(page_cfg.get("header_level"), int):
            return int(page_cfg.get("header_level"))
    except Exception:  # nosec B110 - invalid header_level
        pass
    return 1


def _use_plain_bullets(sec: dict[str, Any] | None, page_cfg: dict[str, Any] | None) -> tuple:
    renderer = BulletRenderer.__new__(BulletRenderer)
    renderer.page_cfg = page_cfg or {}
    return renderer.get_bullet_config(sec)


def _extract_experience_locations(data: dict[str, Any]) -> list[str]:
    """Extract unique location strings from experience entries."""
    locs = [str(e.get("location") or "").strip() for e in (data.get("experience") or [])]
    return list(dict.fromkeys([loc for loc in locs if loc]))


def _get_contact_field(data: dict[str, Any], field: str) -> str:
    """Get a contact field from data or nested contact dict."""
    contact = data.get("contact") or {}
    return data.get(field) or contact.get(field) or ""


def _collect_link_extras(data: dict[str, Any]) -> list[str]:
    """Collect formatted link extras (website, linkedin, github, links list)."""
    return [display for display, _url in _collect_link_extra_items(data)]


def _collect_link_extra_items(data: dict[str, Any]) -> list[tuple[str, str]]:
    """Collect link extras as (display, url) pairs for hyperlink rendering."""
    items: list[tuple[str, str]] = []
    for field in ["website", "linkedin", "github"]:
        val = _get_contact_field(data, field)
        if val:
            items.append((_format_link_display(val), normalize_link_url(val)))
    links_list = _get_contact_field(data, "links") or []
    for val in (links_list if isinstance(links_list, list) else []):
        if isinstance(val, str) and val.strip():
            items.append((_format_link_display(val), normalize_link_url(val)))
    return items


def _apply_page_styles(doc, page_cfg: dict[str, Any]) -> None:
    """Apply compact page styles (margins and fonts)."""
    from .docx_base import apply_page_styles_to_doc
    apply_page_styles_to_doc(doc, page_cfg)


def _set_document_metadata(doc, data: dict[str, Any], template: dict[str, Any]) -> None:
    """Set document core properties (title, author, keywords)."""
    from .docx_base import set_document_metadata_on_doc
    page_cfg = template.get("page") or {}
    set_document_metadata_on_doc(doc, data, page_cfg)


def _center_paragraph(para) -> None:
    """Center a paragraph and remove indents."""
    from .docx_styles import StyleManager
    StyleManager.center_paragraph(para)


def _render_document_header(doc, data: dict[str, Any]) -> None:
    """Render the name, headline, and contact line at the top of the resume."""
    name = _get_contact_field(data, "name")
    headline = _get_contact_field(data, "headline")
    email = _get_contact_field(data, "email")
    phone = _get_contact_field(data, "phone")
    display_phone = _format_phone_display(phone) if phone else ""
    location = _get_contact_field(data, "location")

    if name:
        doc.add_heading(name, level=0)
        _tight_paragraph(doc.paragraphs[-1], after_pt=2)
        _center_paragraph(doc.paragraphs[-1])

    if headline:
        p_head = doc.add_paragraph(str(headline))
        _tight_paragraph(p_head, after_pt=2)
        _center_paragraph(p_head)

    # Contact line: email (mailto: hyperlink) | phone | location | link extras
    plain_parts = [p for p in [display_phone, location] if p]
    link_items = _collect_link_extra_items(data)

    has_content = email or plain_parts or link_items
    if has_content:
        p = doc.add_paragraph()
        _tight_paragraph(p, after_pt=6)
        _center_paragraph(p)
        _render_contact_runs(p, email, plain_parts, link_items)


def _resolve_sections(template: dict[str, Any], structure: dict[str, Any] | None) -> list[dict[str, Any]]:
    """Resolve section order and configuration from template/structure."""
    sections = template.get("sections") or []

    if structure and isinstance(structure.get("order"), list):
        order_keys: list[str] = structure.get("order", [])
        key_to_title: dict[str, str] = structure.get("titles", {})
        tpl_by_key = {s.get("key"): s for s in sections if s.get("key")}
        sections = [
            {**tpl_by_key.get(k, {"key": k, "title": key_to_title.get(k, k.title())})}
            for k in order_keys
            if k in tpl_by_key or key_to_title.get(k)
        ]

    return sections


def _render_section_heading(doc, title: str, template: dict[str, Any]) -> None:
    """Render a section heading with optional shading."""
    if not title:
        return
    doc.add_heading(str(title), level=1)
    _tight_paragraph(doc.paragraphs[-1], before_pt=6, after_pt=2)
    _flush_left(doc.paragraphs[-1])
    page_h1_bg = (template.get("page") or {}).get("h1_bg") or (template.get("page") or {}).get("heading_bg")
    bg_rgb = _parse_hex_color(page_h1_bg)
    if bg_rgb:
        _apply_paragraph_shading(doc.paragraphs[-1], bg_rgb)


def write_resume_docx(
    data: dict[str, Any],
    template: dict[str, Any],
    out_path: str,
    seed: dict[str, Any] | None = None,
    structure: dict[str, Any] | None = None,
) -> None:
    """Write resume to DOCX format.

    This is the main entry point for backward compatibility.
    For new code, prefer using create_resume_writer() from docx_base.

    Args:
        data: Resume data (name, experience, education, etc.)
        template: Template configuration (sections, page styles, etc.)
        out_path: Output file path
        seed: Optional seed data (keywords, etc.)
        structure: Optional structure override for section order
    """
    # Check for sidebar layout and delegate to appropriate writer
    layout_cfg = template.get("layout") or {}
    if layout_cfg.get("type") == "sidebar":
        from .docx_sidebar import write_resume_docx_sidebar
        return write_resume_docx_sidebar(data, template, out_path, seed)

    # Standard single-column layout
    docx = safe_import("docx")
    if not docx:
        raise RuntimeError("Rendering DOCX requires python-docx; install python-docx.")

    from docx import Document  # type: ignore

    doc = Document()
    page_cfg = template.get("page") or {}

    # Apply page styles, metadata, and header
    _apply_page_styles(doc, page_cfg)
    _set_document_metadata(doc, data, template)
    _render_document_header(doc, data)

    # Extract keywords from seed
    keywords = []
    if seed and isinstance(seed.get("keywords"), list):
        keywords = [str(k) for k in seed.get("keywords", [])]

    sections = _resolve_sections(template, structure)
    _render_sections(doc, template, data, sections, keywords)
    doc.save(out_path)


def _render_sections(doc, template, data, sections, keywords) -> None:
    """Render all sections into the document."""
    page_cfg = template.get("page") or {}
    for sec in sections:
        key = sec.get("key")
        if not key:
            continue
        title = sec.get("title") or (key.title() if isinstance(key, str) else "")
        _render_section_heading(doc, title, template)
        renderer_class = SECTION_RENDERERS.get(key)
        if renderer_class:
            renderer = renderer_class(doc, page_cfg)
            if key in SECTIONS_WITH_KEYWORDS:
                renderer.render(data, sec, keywords)
            else:
                renderer.render(data, sec)
