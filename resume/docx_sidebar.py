"""DOCX sidebar layout renderer for resumes.

Provides two-column sidebar layout with repeating header.
"""
from __future__ import annotations

from typing import Any, Dict, List, Optional

from docx.shared import Pt, Inches, RGBColor  # type: ignore
from docx.enum.table import WD_TABLE_ALIGNMENT  # type: ignore
from docx.oxml.ns import qn  # type: ignore
from docx.oxml import OxmlElement  # type: ignore

from .docx_styles import (
    _parse_hex_color,
    _tight_paragraph,
    _apply_paragraph_shading,
)


# Sidebar layout constants
SIDEBAR_SECTIONS = {"summary", "skills", "contact"}
MAIN_SECTIONS = {"education", "experience", "teaching", "presentations", "certifications", "languages"}


def _set_cell_shading(cell, hex_color: str) -> None:
    """Set background shading on a table cell."""
    rgb = _parse_hex_color(hex_color)
    if not rgb:
        return
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def _remove_cell_borders(cell) -> None:
    """Remove all borders from a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _render_sidebar_header(cell, data: Dict[str, Any], page_cfg: Dict[str, Any]) -> None:
    """Render name and headline in sidebar cell."""
    from .docx_writer import _get_contact_field

    name = _get_contact_field(data, "name")
    headline = _get_contact_field(data, "headline")

    name_color = page_cfg.get("sidebar_name_color", "#333333")
    text_color = page_cfg.get("sidebar_text_color", "#333333")

    if name:
        p = cell.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(page_cfg.get("sidebar_name_pt", 24))
        rgb = _parse_hex_color(name_color)
        if rgb:
            run.font.color.rgb = RGBColor(*rgb)
        _tight_paragraph(p, after_pt=2)

    if headline:
        p = cell.add_paragraph()
        run = p.add_run(headline)
        run.font.size = Pt(page_cfg.get("sidebar_headline_pt", 11))
        rgb = _parse_hex_color(text_color)
        if rgb:
            run.font.color.rgb = RGBColor(*rgb)
        _tight_paragraph(p, after_pt=12)


def _render_sidebar_contact(cell, data: Dict[str, Any], page_cfg: Dict[str, Any], title: str = "Contacto") -> None:
    """Render contact section in sidebar."""
    from .docx_writer import _get_contact_field

    email = _get_contact_field(data, "email")
    phone = _get_contact_field(data, "phone")
    location = _get_contact_field(data, "location")

    h1_color = page_cfg.get("h1_color", "#D4A84B")
    text_color = page_cfg.get("sidebar_text_color", "#333333")

    p = cell.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(page_cfg.get("h1_pt", 14))
    rgb = _parse_hex_color(h1_color)
    if rgb:
        run.font.color.rgb = RGBColor(*rgb)
    _tight_paragraph(p, before_pt=12, after_pt=6)

    for item in [phone, email, location]:
        if item:
            p = cell.add_paragraph()
            run = p.add_run(item)
            run.font.size = Pt(page_cfg.get("body_pt", 10))
            rgb = _parse_hex_color(text_color)
            if rgb:
                run.font.color.rgb = RGBColor(*rgb)
            _tight_paragraph(p, after_pt=2)


def _render_sidebar_section(cell, title: str, items: List[str], page_cfg: Dict[str, Any], bulleted: bool = True) -> None:
    """Render a generic section in sidebar with optional bullets."""
    h1_color = page_cfg.get("h1_color", "#D4A84B")
    text_color = page_cfg.get("sidebar_text_color", "#333333")
    bullet_color = page_cfg.get("sidebar_bullet_color", "#4A90A4")

    p = cell.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(page_cfg.get("h1_pt", 14))
    rgb = _parse_hex_color(h1_color)
    if rgb:
        run.font.color.rgb = RGBColor(*rgb)
    _tight_paragraph(p, before_pt=12, after_pt=6)

    for item in items:
        p = cell.add_paragraph()
        if bulleted:
            bullet_run = p.add_run("• ")
            bullet_rgb = _parse_hex_color(bullet_color)
            if bullet_rgb:
                bullet_run.font.color.rgb = RGBColor(*bullet_rgb)
            bullet_run.font.size = Pt(page_cfg.get("body_pt", 10))
        run = p.add_run(item)
        run.font.size = Pt(page_cfg.get("body_pt", 10))
        rgb = _parse_hex_color(text_color)
        if rgb:
            run.font.color.rgb = RGBColor(*rgb)
        _tight_paragraph(p, after_pt=2)


def _render_main_section_heading(cell, title: str, page_cfg: Dict[str, Any]) -> None:
    """Render a section heading in main column."""
    h1_color = page_cfg.get("h1_color", "#D4A84B")

    p = cell.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(page_cfg.get("h1_pt", 14))
    rgb = _parse_hex_color(h1_color)
    if rgb:
        run.font.color.rgb = RGBColor(*rgb)
    _tight_paragraph(p, before_pt=10, after_pt=6)


def _render_main_education(cell, data: Dict[str, Any], page_cfg: Dict[str, Any], sec: Dict[str, Any]) -> None:
    """Render education in main column."""
    education = data.get("education") or []
    bullet_color = page_cfg.get("main_bullet_color", "#4A90A4")

    for edu in education:
        degree = edu.get("degree", "")
        institution = edu.get("institution", "")
        year = edu.get("year", "")

        p = cell.add_paragraph()
        bullet_run = p.add_run("• ")
        bullet_rgb = _parse_hex_color(bullet_color)
        if bullet_rgb:
            bullet_run.font.color.rgb = RGBColor(*bullet_rgb)

        deg_run = p.add_run(degree)
        deg_run.bold = True
        deg_run.font.size = Pt(page_cfg.get("body_pt", 10))
        _tight_paragraph(p, after_pt=0)

        if institution or year:
            p2 = cell.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.25)
            inst_run = p2.add_run(institution)
            inst_run.italic = True
            inst_run.font.size = Pt(page_cfg.get("meta_pt", 9))
            inst_rgb = _parse_hex_color("#666666")
            if inst_rgb:
                inst_run.font.color.rgb = RGBColor(*inst_rgb)
            if year:
                year_run = p2.add_run(f"  {year}")
                year_run.font.size = Pt(page_cfg.get("meta_pt", 9))
                if inst_rgb:
                    year_run.font.color.rgb = RGBColor(*inst_rgb)
            _tight_paragraph(p2, after_pt=6)


def _render_main_experience(cell, data: Dict[str, Any], page_cfg: Dict[str, Any], sec: Dict[str, Any]) -> None:
    """Render experience in main column."""
    experience = data.get("experience") or []
    bullet_color = page_cfg.get("main_bullet_color", "#4A90A4")

    for exp in experience:
        title = exp.get("title", "")
        company = exp.get("company", "")
        start = exp.get("start", "")
        end = exp.get("end", "")
        span = f"{start} – {end}" if end else f"{start} – presente"
        bullets = exp.get("bullets") or []

        p = cell.add_paragraph()
        bullet_run = p.add_run("• ")
        bullet_rgb = _parse_hex_color(bullet_color)
        if bullet_rgb:
            bullet_run.font.color.rgb = RGBColor(*bullet_rgb)

        title_run = p.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(page_cfg.get("body_pt", 10))

        if span:
            span_run = p.add_run(f"  {span}")
            span_run.font.size = Pt(page_cfg.get("meta_pt", 9))
            span_rgb = _parse_hex_color("#666666")
            if span_rgb:
                span_run.font.color.rgb = RGBColor(*span_rgb)
        _tight_paragraph(p, after_pt=0)

        if company:
            p2 = cell.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.25)
            comp_run = p2.add_run(company)
            comp_run.italic = True
            comp_run.font.size = Pt(page_cfg.get("meta_pt", 9))
            comp_rgb = _parse_hex_color("#666666")
            if comp_rgb:
                comp_run.font.color.rgb = RGBColor(*comp_rgb)
            _tight_paragraph(p2, after_pt=2)

        for b in bullets[:sec.get("recent_max_bullets", 3)]:
            text = b.get("text", b) if isinstance(b, dict) else b
            p3 = cell.add_paragraph()
            p3.paragraph_format.left_indent = Inches(0.25)
            run = p3.add_run(text)
            run.font.size = Pt(page_cfg.get("body_pt", 10) - 1)
            _tight_paragraph(p3, after_pt=1)


def _render_main_teaching(cell, data: Dict[str, Any], page_cfg: Dict[str, Any], sec: Dict[str, Any]) -> None:
    """Render teaching in main column (uppercase titles with institution below)."""
    teaching = data.get("teaching") or []

    for item in teaching:
        text = item.get("text", item) if isinstance(item, dict) else item
        if "(" in text and text.endswith(")"):
            parts = text.rsplit("(", 1)
            title = parts[0].strip()
            institution = parts[1].rstrip(")")
        else:
            title = text
            institution = ""

        p = cell.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(page_cfg.get("body_pt", 10))
        _tight_paragraph(p, after_pt=0)

        if institution:
            p2 = cell.add_paragraph()
            run2 = p2.add_run(institution)
            run2.italic = True
            run2.font.size = Pt(page_cfg.get("meta_pt", 9))
            inst_rgb = _parse_hex_color("#666666")
            if inst_rgb:
                run2.font.color.rgb = RGBColor(*inst_rgb)
            _tight_paragraph(p2, after_pt=6)


def _render_main_presentations(cell, data: Dict[str, Any], page_cfg: Dict[str, Any], sec: Dict[str, Any]) -> None:
    """Render presentations/publications in main column."""
    presentations = data.get("presentations") or []
    bullet_color = page_cfg.get("main_bullet_color", "#4A90A4")

    for pres in presentations:
        title = pres.get("title", "")
        authors = pres.get("authors", "")
        event = pres.get("event", "")
        note = pres.get("note", "")

        p = cell.add_paragraph()
        bullet_run = p.add_run("• ")
        bullet_rgb = _parse_hex_color(bullet_color)
        if bullet_rgb:
            bullet_run.font.color.rgb = RGBColor(*bullet_rgb)

        title_run = p.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(page_cfg.get("body_pt", 10))
        _tight_paragraph(p, after_pt=0)

        if authors:
            p2 = cell.add_paragraph()
            p2.paragraph_format.left_indent = Inches(0.25)
            auth_run = p2.add_run(authors)
            auth_run.italic = True
            auth_run.font.size = Pt(page_cfg.get("meta_pt", 9))
            auth_rgb = _parse_hex_color("#666666")
            if auth_rgb:
                auth_run.font.color.rgb = RGBColor(*auth_rgb)
            _tight_paragraph(p2, after_pt=0)

        if event:
            p3 = cell.add_paragraph()
            p3.paragraph_format.left_indent = Inches(0.25)
            ev_run = p3.add_run(event)
            ev_run.font.size = Pt(page_cfg.get("meta_pt", 9) - 1)
            ev_rgb = _parse_hex_color("#888888")
            if ev_rgb:
                ev_run.font.color.rgb = RGBColor(*ev_rgb)
            _tight_paragraph(p3, after_pt=0)

        if note:
            p4 = cell.add_paragraph()
            p4.paragraph_format.left_indent = Inches(0.25)
            note_run = p4.add_run(note)
            note_run.italic = True
            note_run.font.size = Pt(page_cfg.get("meta_pt", 9) - 1)
            _tight_paragraph(p4, after_pt=4)
        else:
            if event:
                p3.paragraph_format.space_after = Pt(4)
            elif authors:
                p2.paragraph_format.space_after = Pt(4)


def _render_page_header(doc, data: Dict[str, Any], page_cfg: Dict[str, Any]) -> None:
    """Add name, headline, and contact as centered header (repeats on each page)."""
    from .docx_writer import _get_contact_field, _center_paragraph

    section = doc.sections[0]
    header = section.header

    name = _get_contact_field(data, "name")
    headline = _get_contact_field(data, "headline")
    email = _get_contact_field(data, "email")
    phone = _get_contact_field(data, "phone")
    location = _get_contact_field(data, "location")

    name_color = page_cfg.get("sidebar_name_color", "#1A365D")
    text_color = page_cfg.get("sidebar_text_color", "#333333")
    header_bg = page_cfg.get("header_bg", "#F7F9FC")

    # Name in header (centered)
    if header.paragraphs:
        p = header.paragraphs[0]
        p.clear()
    else:
        p = header.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(page_cfg.get("sidebar_name_pt", 20))
    rgb = _parse_hex_color(name_color)
    if rgb:
        run.font.color.rgb = RGBColor(*rgb)
    _tight_paragraph(p, after_pt=0)
    _center_paragraph(p)
    bg_rgb = _parse_hex_color(header_bg)
    if bg_rgb:
        _apply_paragraph_shading(p, bg_rgb)

    # Headline (centered)
    if headline:
        p2 = header.add_paragraph()
        run2 = p2.add_run(headline)
        run2.font.size = Pt(page_cfg.get("sidebar_headline_pt", 10))
        rgb2 = _parse_hex_color(text_color)
        if rgb2:
            run2.font.color.rgb = RGBColor(*rgb2)
        _tight_paragraph(p2, after_pt=2)
        _center_paragraph(p2)
        if bg_rgb:
            _apply_paragraph_shading(p2, bg_rgb)

    # Contact line (centered)
    contact_parts = [x for x in [phone, email, location] if x]
    if contact_parts:
        p3 = header.add_paragraph()
        run3 = p3.add_run(" | ".join(contact_parts))
        run3.font.size = Pt(page_cfg.get("body_pt", 10) - 1)
        rgb3 = _parse_hex_color("#666666")
        if rgb3:
            run3.font.color.rgb = RGBColor(*rgb3)
        _tight_paragraph(p3, after_pt=6)
        _center_paragraph(p3)
        if bg_rgb:
            _apply_paragraph_shading(p3, bg_rgb)


def write_resume_docx_sidebar(
    data: Dict[str, Any],
    template: Dict[str, Any],
    out_path: str,
    seed: Optional[Dict[str, Any]] = None,
    structure: Optional[Dict[str, Any]] = None,
) -> None:
    """Write resume with two-column sidebar layout."""
    from .io_utils import safe_import
    from .docx_writer import _apply_page_styles, _set_document_metadata

    docx = safe_import("docx")
    if not docx:
        raise RuntimeError("Rendering DOCX requires python-docx; install python-docx.")

    from docx import Document  # type: ignore

    doc = Document()
    page_cfg = template.get("page") or {}
    layout_cfg = template.get("layout") or {}

    # Apply page styles and metadata
    _apply_page_styles(doc, page_cfg)
    _set_document_metadata(doc, data, template)

    # Layout dimensions
    sidebar_width = layout_cfg.get("sidebar_width", 2.3)
    main_width = layout_cfg.get("main_width", 5.2)
    sidebar_bg = layout_cfg.get("sidebar_bg")

    # Render page header (repeats on all pages)
    _render_page_header(doc, data, page_cfg)

    # Create two-column table for body
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    table.columns[0].width = Inches(sidebar_width)
    table.columns[1].width = Inches(main_width)

    sidebar_cell = table.rows[0].cells[0]
    main_cell = table.rows[0].cells[1]

    _remove_cell_borders(sidebar_cell)
    _remove_cell_borders(main_cell)

    if sidebar_bg:
        _set_cell_shading(sidebar_cell, sidebar_bg)

    # Clear default paragraphs
    if sidebar_cell.paragraphs:
        sidebar_cell.paragraphs[0].clear()
    if main_cell.paragraphs:
        main_cell.paragraphs[0].clear()

    # Page 1 sidebar: Profile + Skills
    for sec in (template.get("sections") or []):
        if sec.get("key") == "summary":
            summary_items = data.get("summary") or []
            if isinstance(summary_items, str):
                summary_items = [summary_items]
            elif isinstance(summary_items, list):
                summary_items = [s.get("text", s) if isinstance(s, dict) else s for s in summary_items]
            if summary_items:
                _render_sidebar_section(
                    sidebar_cell,
                    sec.get("title", "Perfil profesional"),
                    summary_items[:6],
                    page_cfg,
                    bulleted=True
                )
            break

    for sec in (template.get("sections") or []):
        if sec.get("key") == "skills":
            skills_groups = data.get("skills_groups") or []
            skill_items = []
            for group in skills_groups:
                for item in (group.get("items") or []):
                    name = item.get("name", item) if isinstance(item, dict) else item
                    skill_items.append(name)
            if skill_items:
                _render_sidebar_section(sidebar_cell, sec.get("title", "Habilidades claves"), skill_items[:8], page_cfg)
            break

    # Main content (natural flow)
    sections = template.get("sections") or []

    for sec in sections:
        key = sec.get("key")
        title = sec.get("title", "")

        if key == "education":
            _render_main_section_heading(main_cell, title, page_cfg)
            _render_main_education(main_cell, data, page_cfg, sec)
        elif key == "experience":
            _render_main_section_heading(main_cell, title, page_cfg)
            _render_main_experience(main_cell, data, page_cfg, sec)
        elif key == "teaching":
            _render_main_section_heading(main_cell, title, page_cfg)
            _render_main_teaching(main_cell, data, page_cfg, sec)
        elif key == "presentations":
            _render_main_section_heading(main_cell, title, page_cfg)
            _render_main_presentations(main_cell, data, page_cfg, sec)

    doc.save(out_path)
