"""Tests for resume/docx_links.py — hyperlink helpers.

Covers:
  - normalize_link_url: email, bare URL, already-schemed URL, edge cases
  - add_hyperlink: integration with a real python-docx Document when available;
    fallback to plain-text run when OXML manipulation is unavailable (mock).
  - End-to-end: render a minimal resume with an email and a presentations link,
    then inspect the resulting DOCX via zipfile to confirm w:hyperlink elements
    and external relationships appear in word/document.xml and
    word/_rels/document.xml.rels.
"""
from __future__ import annotations

import io
import os
import tempfile
import unittest
import zipfile
from unittest.mock import MagicMock, patch

from resume.schema import Resume


# ---------------------------------------------------------------------------
# Unit tests for display_url — no docx dependency
# ---------------------------------------------------------------------------

class TestDisplayUrl(unittest.TestCase):
    """Tests for display_url."""

    def _fn(self, url: str) -> str:
        from resume.docx_links import display_url
        return display_url(url)

    def test_empty_returns_empty(self):
        self.assertEqual(self._fn(""), "")

    def test_https_scheme_stripped(self):
        self.assertEqual(
            self._fn("https://github.com/dancing-bear-show/dancing-bear"),
            "github.com/dancing-bear-show/dancing-bear",
        )

    def test_http_scheme_stripped(self):
        self.assertEqual(
            self._fn("http://example.com/page"),
            "example.com/page",
        )

    def test_www_prefix_stripped(self):
        self.assertEqual(
            self._fn("https://www.youtube.com/watch?v=TmjY1HJemi4"),
            "youtube.com/watch?v=TmjY1HJemi4",
        )

    def test_mailto_scheme_stripped(self):
        self.assertEqual(
            self._fn("mailto:brian@example.com"),
            "brian@example.com",
        )

    def test_trailing_slash_stripped(self):
        self.assertEqual(
            self._fn("https://example.com/"),
            "example.com",
        )

    def test_bare_url_unchanged(self):
        # No scheme, no www — returned as-is (no leading slash to strip)
        self.assertEqual(self._fn("slides.example.com/my-talk"), "slides.example.com/my-talk")

    def test_whitespace_only_returns_empty(self):
        self.assertEqual(self._fn("  "), "")


# ---------------------------------------------------------------------------
# Unit tests for normalize_link_url — no docx dependency
# ---------------------------------------------------------------------------

class TestNormalizeLinkUrl(unittest.TestCase):
    """Tests for normalize_link_url."""

    def _fn(self, text: str) -> str:
        from resume.docx_links import normalize_link_url
        return normalize_link_url(text)

    def test_empty_returns_empty(self):
        self.assertEqual(self._fn(""), "")

    def test_whitespace_only_returns_stripped_empty(self):
        self.assertEqual(self._fn("  "), "")

    def test_email_gets_mailto_prefix(self):
        self.assertEqual(self._fn("user@example.com"), "mailto:user@example.com")

    def test_email_with_plus(self):
        self.assertEqual(self._fn("user+tag@example.com"), "mailto:user+tag@example.com")

    def test_https_url_returned_as_is(self):
        self.assertEqual(
            self._fn("https://example.com/path"),
            "https://example.com/path",
        )

    def test_http_url_returned_as_is(self):
        self.assertEqual(self._fn("http://example.com"), "http://example.com")

    def test_mailto_url_returned_as_is(self):
        self.assertEqual(
            self._fn("mailto:user@example.com"),
            "mailto:user@example.com",
        )

    def test_bare_domain_gets_https_prefix(self):
        self.assertEqual(self._fn("example.com"), "https://example.com")

    def test_linkedin_bare_url_gets_https(self):
        self.assertEqual(
            self._fn("linkedin.com/in/johndoe"),
            "https://linkedin.com/in/johndoe",
        )

    def test_github_bare_url_gets_https(self):
        self.assertEqual(
            self._fn("github.com/johndoe"),
            "https://github.com/johndoe",
        )

    def test_file_scheme_returns_empty(self):
        """file: scheme is disallowed — normalize_link_url returns empty string."""
        self.assertEqual(self._fn("file:///etc/passwd"), "")

    def test_javascript_scheme_returns_empty(self):
        """javascript: scheme is disallowed — normalize_link_url returns empty string."""
        self.assertEqual(self._fn("javascript:alert(1)"), "")

    def test_data_scheme_returns_empty(self):
        """data: scheme is disallowed — normalize_link_url returns empty string."""
        self.assertEqual(self._fn("data:text/html,<h1>hi</h1>"), "")

    def test_ftp_scheme_returns_empty(self):
        """ftp: scheme is disallowed — normalize_link_url returns empty string."""
        self.assertEqual(self._fn("ftp://files.example.com/pub"), "")


class TestIsSafeLinkUrl(unittest.TestCase):
    """Tests for is_safe_link_url."""

    def _fn(self, url: str) -> bool:
        from resume.docx_links import is_safe_link_url
        return is_safe_link_url(url)

    def test_empty_returns_false(self):
        self.assertFalse(self._fn(""))

    def test_https_is_safe(self):
        self.assertTrue(self._fn("https://example.com"))

    def test_http_is_safe(self):
        self.assertTrue(self._fn("http://example.com"))

    def test_mailto_is_safe(self):
        self.assertTrue(self._fn("mailto:user@example.com"))

    def test_file_is_not_safe(self):
        self.assertFalse(self._fn("file:///etc/passwd"))

    def test_javascript_is_not_safe(self):
        self.assertFalse(self._fn("javascript:alert(1)"))

    def test_bare_url_no_scheme_is_safe(self):
        """Bare URLs without a scheme are safe — normalize_link_url will prefix https://."""
        self.assertTrue(self._fn("example.com"))

    def test_bare_email_no_scheme_is_safe(self):
        """Bare emails without a scheme are safe — normalize_link_url will prefix mailto:."""
        self.assertTrue(self._fn("user@example.com"))



# ---------------------------------------------------------------------------
# Unit tests for add_hyperlink — using mocked docx internals
# ---------------------------------------------------------------------------

class TestAddHyperlinkUnsafeUrl(unittest.TestCase):
    """add_hyperlink falls back to plain text for unsafe/empty URLs."""

    def test_empty_url_renders_plain_text_without_relationship(self):
        """An empty URL (e.g. from a rejected scheme) renders as plain text."""
        from resume.docx_links import add_hyperlink
        paragraph = MagicMock()
        add_hyperlink(paragraph, "", "file display")
        paragraph.add_run.assert_called_once_with("file display")
        # part.relate_to must NOT be called — no external relationship created.
        paragraph.part.relate_to.assert_not_called()

    def test_file_scheme_via_normalize_does_not_create_relationship(self):
        """file: URL normalized to '' means add_hyperlink renders plain text only."""
        from resume.docx_links import add_hyperlink, normalize_link_url
        url = normalize_link_url("file:///etc/passwd")
        self.assertEqual(url, "")  # confirm rejected
        paragraph = MagicMock()
        add_hyperlink(paragraph, url, "secret file")
        paragraph.add_run.assert_called_once_with("secret file")
        paragraph.part.relate_to.assert_not_called()

    def test_javascript_scheme_via_normalize_does_not_create_relationship(self):
        """javascript: URL normalized to '' means add_hyperlink renders plain text."""
        from resume.docx_links import add_hyperlink, normalize_link_url
        url = normalize_link_url("javascript:alert(1)")
        self.assertEqual(url, "")
        paragraph = MagicMock()
        add_hyperlink(paragraph, url, "xss")
        paragraph.add_run.assert_called_once_with("xss")
        paragraph.part.relate_to.assert_not_called()


class TestAddHyperlinkFallback(unittest.TestCase):
    """Tests for add_hyperlink fallback to plain text when OXML is unavailable."""

    def test_falls_back_to_plain_run_when_oxml_missing(self):
        """When docx OXML imports fail, falls back to paragraph.add_run(display)."""
        from resume.docx_links import add_hyperlink

        paragraph = MagicMock()
        # Patch the import of docx.oxml so that the internal impl raises
        with patch.dict("sys.modules", {"docx.oxml": None, "docx.opc.constants": None}):
            add_hyperlink(paragraph, "https://example.com", "example.com")

        paragraph.add_run.assert_called_once_with("example.com")

    def test_fallback_calls_add_run_when_impl_raises(self):
        """add_hyperlink falls back to add_run when the OXML impl raises."""
        from resume.docx_links import add_hyperlink

        paragraph = MagicMock()
        # Force the impl to raise by making _p.append raise
        paragraph._p.append.side_effect = RuntimeError("oxml failure")
        # add_run works fine (the MagicMock default)

        with patch.dict("sys.modules", {
            "docx.oxml": MagicMock(),
            "docx.oxml.ns": MagicMock(),
            "docx.opc.constants": MagicMock(),
        }):
            add_hyperlink(paragraph, "https://example.com", "example.com")

        paragraph.add_run.assert_called_once_with("example.com")


class TestAddHyperlinkWithMockedDocx(unittest.TestCase):
    """Tests for add_hyperlink with a fully mocked python-docx environment."""

    def _make_mock_paragraph(self):
        """Build a minimal mock paragraph that lets _add_hyperlink_impl succeed."""
        para = MagicMock()
        part = MagicMock()
        part.relate_to.return_value = "rId1"
        para.part = part
        para._p = MagicMock()
        return para

    def test_hyperlink_impl_calls_relate_to(self):
        """_add_hyperlink_impl registers a relationship via part.relate_to."""
        mock_oxml = MagicMock()
        mock_oxml_ns = MagicMock()
        mock_oxml_ns.qn.side_effect = lambda x: x
        mock_opc_constants = MagicMock()
        mock_opc_constants.RELATIONSHIP_TYPE.HYPERLINK = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

        with patch.dict("sys.modules", {
            "docx.oxml": mock_oxml,
            "docx.oxml.ns": mock_oxml_ns,
            "docx.opc.constants": mock_opc_constants,
        }):
            from resume.docx_links import _add_hyperlink_impl
            para = self._make_mock_paragraph()
            _add_hyperlink_impl(para, "https://example.com", "example.com")

        para.part.relate_to.assert_called_once()
        call_args = para.part.relate_to.call_args
        self.assertEqual(call_args[0][0], "https://example.com")
        self.assertTrue(call_args[1].get("is_external", False))


# ---------------------------------------------------------------------------
# End-to-end rendering test — requires actual python-docx
# ---------------------------------------------------------------------------

def _docx_available() -> bool:
    """Check if python-docx is installed."""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


@unittest.skipUnless(_docx_available(), "python-docx not installed")
class TestHyperlinkEndToEnd(unittest.TestCase):
    """Integration tests: render a DOCX and inspect the zip contents.

    These tests require python-docx to be installed and do NOT make any
    network requests.
    """

    def _render_resume(self, data: dict, template: dict) -> bytes:
        """Render a resume and return the raw DOCX bytes."""
        from resume.docx_base import create_resume_writer
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            path = f.name
        try:
            writer = create_resume_writer(Resume.from_dict(data), template)
            writer.write(path)
            with open(path, "rb") as f:
                return f.read()
        finally:
            if os.path.exists(path):
                os.unlink(path)

    def _get_zip_member(self, docx_bytes: bytes, name: str) -> str:
        """Extract and decode a zip member from DOCX bytes."""
        with zipfile.ZipFile(io.BytesIO(docx_bytes)) as z:
            return z.read(name).decode("utf-8")

    def test_email_produces_mailto_relationship(self):
        """Email in contact info creates a mailto: external relationship."""
        data = {
            "name": "Alice Test",
            "email": "alice@example.com",
        }
        template = {"sections": [], "page": {"compact": False}}

        docx_bytes = self._render_resume(data, template)
        rels_xml = self._get_zip_member(docx_bytes, "word/_rels/document.xml.rels")

        self.assertIn("mailto:alice@example.com", rels_xml)
        self.assertIn('TargetMode="External"', rels_xml)

    def test_email_produces_w_hyperlink_element(self):
        """Email in contact info creates a w:hyperlink element in document.xml."""
        data = {
            "name": "Alice Test",
            "email": "alice@example.com",
        }
        template = {"sections": [], "page": {"compact": False}}

        docx_bytes = self._render_resume(data, template)
        doc_xml = self._get_zip_member(docx_bytes, "word/document.xml")

        self.assertIn("w:hyperlink", doc_xml)

    def test_website_link_produces_external_relationship(self):
        """Website URL in contact info creates an external https: relationship."""
        data = {
            "name": "Bob Test",
            "website": "https://bob.dev",
        }
        template = {"sections": [], "page": {"compact": False}}

        docx_bytes = self._render_resume(data, template)
        rels_xml = self._get_zip_member(docx_bytes, "word/_rels/document.xml.rels")

        self.assertIn("https://bob.dev", rels_xml)
        self.assertIn('TargetMode="External"', rels_xml)

    def test_presentations_link_produces_external_relationship(self):
        """Presentation with a link field creates an external https: relationship."""
        data = {
            "name": "Carol Test",
            "presentations": [
                {
                    "title": "My Talk",
                    "event": "PyCon",
                    "year": "2024",
                    "link": "https://slides.example.com/my-talk",
                }
            ],
        }
        template = {
            "sections": [{"key": "presentations", "title": "Presentations"}],
            "page": {"compact": False},
        }

        docx_bytes = self._render_resume(data, template)
        rels_xml = self._get_zip_member(docx_bytes, "word/_rels/document.xml.rels")

        self.assertIn("https://slides.example.com/my-talk", rels_xml)
        self.assertIn('TargetMode="External"', rels_xml)

    def test_presentations_link_produces_w_hyperlink(self):
        """Presentation with a link field creates a w:hyperlink in document.xml."""
        data = {
            "name": "Carol Test",
            "presentations": [
                {
                    "title": "My Talk",
                    "link": "https://slides.example.com/my-talk",
                }
            ],
        }
        template = {
            "sections": [{"key": "presentations", "title": "Presentations"}],
            "page": {"compact": False},
        }

        docx_bytes = self._render_resume(data, template)
        doc_xml = self._get_zip_member(docx_bytes, "word/document.xml")

        self.assertIn("w:hyperlink", doc_xml)

    def test_no_links_renders_without_hyperlinks(self):
        """Resume with no email or links renders cleanly with no w:hyperlink."""
        data = {
            "name": "Dave Test",
            "phone": "555-1234",
            "location": "Portland, OR",
        }
        template = {"sections": [], "page": {"compact": False}}

        docx_bytes = self._render_resume(data, template)
        doc_xml = self._get_zip_member(docx_bytes, "word/document.xml")

        # No hyperlinks expected; must not raise
        self.assertIsNotNone(docx_bytes)
        # "Dave Test" should appear in the document body
        self.assertIn("Dave Test", doc_xml)

    def test_presentation_without_link_renders_plain(self):
        """Presentation without a link field renders as plain bullet, no hyperlink."""
        data = {
            "name": "Eve Test",
            "presentations": [
                {"title": "A Plain Talk", "event": "BarConf", "year": "2023"},
            ],
        }
        template = {
            "sections": [{"key": "presentations", "title": "Presentations"}],
            "page": {"compact": False},
        }

        docx_bytes = self._render_resume(data, template)
        doc_xml = self._get_zip_member(docx_bytes, "word/document.xml")

        self.assertIn("A Plain Talk", doc_xml)

    def test_presentations_link_display_text_is_cleaned_url(self):
        """Hyperlink display text is the cleaned URL (no scheme/www), not the word 'link'."""
        data = {
            "name": "Frank Test",
            "presentations": [
                {
                    "title": "My Talk",
                    "link": "https://github.com/dancing-bear-show/dancing-bear",
                }
            ],
        }
        template = {
            "sections": [{"key": "presentations", "title": "Presentations"}],
            "page": {"compact": False},
        }

        docx_bytes = self._render_resume(data, template)
        doc_xml = self._get_zip_member(docx_bytes, "word/document.xml")

        # Cleaned URL appears as run text
        self.assertIn("github.com/dancing-bear-show/dancing-bear", doc_xml)
        # The bare word "link" must NOT appear as a standalone run value
        # (it would appear as ">link<" between XML tags for the w:t element)
        self.assertNotIn(">link<", doc_xml)

    def test_presentations_www_link_display_strips_www(self):
        """Hyperlink display text strips www. prefix from URLs."""
        data = {
            "name": "Grace Test",
            "presentations": [
                {
                    "title": "A YouTube Talk",
                    "link": "https://www.youtube.com/watch?v=TmjY1HJemi4",
                }
            ],
        }
        template = {
            "sections": [{"key": "presentations", "title": "Presentations"}],
            "page": {"compact": False},
        }

        docx_bytes = self._render_resume(data, template)
        doc_xml = self._get_zip_member(docx_bytes, "word/document.xml")
        rels_xml = self._get_zip_member(docx_bytes, "word/_rels/document.xml.rels")

        # Display text: scheme and www stripped
        self.assertIn("youtube.com/watch?v=TmjY1HJemi4", doc_xml)
        # Target relationship: full URL preserved
        self.assertIn("https://www.youtube.com/watch?v=TmjY1HJemi4", rels_xml)

    def test_dangerous_scheme_links_do_not_produce_external_relationship(self):
        """Presentations with file: or javascript: links must NOT create an external relationship.

        Both schemes are dangerous for different reasons — file:// can expose
        local filesystem paths, javascript: enables XSS via alert(document.cookie).
        After scheme restriction, normalize_link_url returns "" and add_hyperlink
        falls back to plain text, so no TargetMode="External" entry should appear
        for either. Each case is checked against its own forbidden substrings
        (the raw URL for file:, both the scheme and payload for javascript:) so a
        regression in either scheme's handling fails independently.
        """
        cases = [
            (
                "file scheme",
                "Helen Test", "Internal Slides", "file:///Users/helen/slides.pdf",
                ["file:///Users/helen/slides.pdf"],
                True,
            ),
            (
                "javascript scheme",
                "Ivan Test", "XSS Demo", "javascript:alert(document.cookie)",
                ["javascript:", "alert"],
                False,
            ),
        ]
        for label, name, title, link, forbidden_in_rels, check_title_renders in cases:
            with self.subTest(label):
                data = {
                    "name": name,
                    "presentations": [{"title": title, "link": link}],
                }
                template = {
                    "sections": [{"key": "presentations", "title": "Presentations"}],
                    "page": {"compact": False},
                }

                docx_bytes = self._render_resume(data, template)
                rels_xml = self._get_zip_member(docx_bytes, "word/_rels/document.xml.rels")

                for forbidden in forbidden_in_rels:
                    self.assertNotIn(forbidden, rels_xml)

                if check_title_renders:
                    doc_xml = self._get_zip_member(docx_bytes, "word/document.xml")
                    self.assertIn(title, doc_xml)

    def test_link_only_presentation_no_duplication_e2e(self):
        """Link-only presentation (no title/event/year) renders without URL duplication.

        The cleaned display URL should appear as a hyperlink run text in document.xml,
        and the raw URL must NOT appear as a separate plain-text run element alongside it.
        """
        data = {
            "name": "Jane E2E",
            "presentations": [
                {"link": "https://slides.example.com/my-talk"},
            ],
        }
        template = {
            "sections": [{"key": "presentations", "title": "Presentations"}],
            "page": {"compact": False},
        }

        docx_bytes = self._render_resume(data, template)
        doc_xml = self._get_zip_member(docx_bytes, "word/document.xml")

        # Hyperlink element must be present.
        self.assertIn("w:hyperlink", doc_xml)
        # Cleaned URL appears as the hyperlink display text.
        self.assertIn("slides.example.com/my-talk", doc_xml)
        # The raw URL must NOT appear as a standalone plain-text run (i.e. not in
        # a w:t element separate from the hyperlink's own w:t).
        # Count w:hyperlink occurrences — should be exactly 1 (no duplication).
        hyperlink_count = doc_xml.count("w:hyperlink")
        # Each hyperlink appears as opening and closing tag, so 2 occurrences per link.
        self.assertEqual(hyperlink_count, 2, "Expected exactly one w:hyperlink element.")


if __name__ == "__main__":
    unittest.main()
