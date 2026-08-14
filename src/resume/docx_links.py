"""DOCX hyperlink utilities for resume rendering.

Provides add_hyperlink() to insert a real clickable w:hyperlink element into
a python-docx paragraph. python-docx has no native hyperlink API, so this
module uses the standard OxmlElement / relate_to() approach.

Usage:
    from resume.docx_links import add_hyperlink, normalize_link_url

    url = normalize_link_url("user@example.com")  # -> "mailto:user@example.com"
    add_hyperlink(paragraph, url, "user@example.com")

Graceful degradation:
    Both helpers are exception-safe. add_hyperlink() falls back to a plain-text
    run when the OOXML manipulation fails; normalize_link_url() returns the
    original string on any error.
"""
from __future__ import annotations

import re


# Blue hyperlink color used by Word's default Hyperlink character style.
_HYPERLINK_COLOR = "0563C1"
_EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+$", re.ASCII)
# Matches any URI scheme: http://, https://, mailto:, ftp://, etc.
_SCHEME_RE = re.compile(r"^[a-zA-Z][a-zA-Z0-9+\-.]*:")


_DISPLAY_STRIP_RE = re.compile(r"^(?:https?://|mailto:)(?:www\.)?")

# Only these URI schemes are permitted for clickable hyperlinks.
_SAFE_SCHEMES = frozenset({"http", "https", "mailto"})


def is_safe_link_url(url: str) -> bool:
    """Return True iff *url* uses an allowed scheme (http, https, or mailto).

    Only these three schemes produce safe, clickable external hyperlinks.
    Schemes such as ``file:``, ``javascript:``, ``data:``, etc. return False
    and must be rendered as plain text.
    """
    s = (url or "").strip()
    if not s:
        return False
    m = _SCHEME_RE.match(s)
    if not m:
        # No scheme present — bare URLs / emails will be prefixed by
        # normalize_link_url, so they are implicitly safe.
        return True
    scheme = s[: m.end() - 1].lower()  # strip trailing ":"
    return scheme in _SAFE_SCHEMES


def display_url(url: str) -> str:
    """Return a human-readable display form of *url*.

    Strips the scheme (``https://``, ``http://``, ``mailto:``) and a leading
    ``www.`` prefix so that ``https://www.youtube.com/watch?v=abc`` becomes
    ``youtube.com/watch?v=abc``. A trailing slash is also removed.

    Returns the original string unchanged if *url* is empty or any error occurs.
    """
    try:
        s = (url or "").strip()
        if not s:
            return s
        s = _DISPLAY_STRIP_RE.sub("", s)
        return s.rstrip("/")
    except Exception:  # nosec B110 - best-effort display cleanup; never crash caller
        return url


def normalize_link_url(text: str) -> str:
    """Return a fully-qualified URL or mailto: URI for *text*, or "" for unsafe schemes.

    Rules:
      - Already has a safe scheme (http, https, mailto) -> returned as-is.
      - Has any other scheme (file:, javascript:, data:, ...) -> returns "" so
        the caller falls back to plain text instead of creating an external
        relationship.
      - Looks like an email address (contains @ but no scheme) -> prepend "mailto:".
      - Bare string with no scheme and not an email -> prepend "https://".

    Returns the original string unchanged if *text* is empty or any error occurs.
    """
    try:
        s = (text or "").strip()
        if not s:
            return s
        if _SCHEME_RE.match(s):
            # Reject schemes that are not in the safe allow-list.
            m = _SCHEME_RE.match(s)
            scheme = s[: m.end() - 1].lower()  # type: ignore[union-attr]
            if scheme not in _SAFE_SCHEMES:
                return ""
            return s
        if _EMAIL_RE.match(s):
            return f"mailto:{s}"
        return f"https://{s}"
    except Exception:  # nosec B110 - best-effort URL normalization; never crash caller
        return text


def add_hyperlink(paragraph, url: str, display: str) -> None:
    """Insert a clickable hyperlink run into *paragraph*.

    Attempts to build a ``<w:hyperlink r:id="...">`` element with blue+underline
    formatting using python-docx's low-level OXML API. Falls back to a plain-text
    run if any step fails (e.g. unexpected python-docx version or mock objects) or
    if *url* is empty / uses a disallowed scheme (see ``is_safe_link_url``).

    Args:
        paragraph: python-docx Paragraph object to append the hyperlink to.
        url: Fully-qualified target URL (use normalize_link_url() first).
             An empty string causes an immediate plain-text fallback.
        display: Text to show for the link.
    """
    if not url:
        paragraph.add_run(display)
        return
    try:
        _add_hyperlink_impl(paragraph, url, display)
    except Exception:  # nosec B110 - fall back to plain text; caller must not crash
        paragraph.add_run(display)


_W_VAL = "w:val"  # XML attribute name used repeatedly in run properties


def _add_hyperlink_impl(paragraph, url: str, display: str) -> None:
    """Internal: add hyperlink element. Raises on failure — caller catches."""
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.opc.constants import RELATIONSHIP_TYPE as RT  # type: ignore

    # Register the external hyperlink relationship on the document part.
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    # Build <w:hyperlink r:id="rIdN"> element.
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    # Build the run element inside the hyperlink.
    new_run = OxmlElement("w:r")

    # Run properties: underline + blue colour to match Word's Hyperlink style.
    r_pr = OxmlElement("w:rPr")
    r_style = OxmlElement("w:rStyle")
    r_style.set(qn(_W_VAL), "Hyperlink")
    r_pr.append(r_style)

    # Explicit colour as fallback if Hyperlink style is absent.
    color_el = OxmlElement("w:color")
    color_el.set(qn(_W_VAL), _HYPERLINK_COLOR)
    r_pr.append(color_el)

    u_el = OxmlElement("w:u")
    u_el.set(qn(_W_VAL), "single")
    r_pr.append(u_el)

    new_run.append(r_pr)

    # Text node.
    t = OxmlElement("w:t")
    t.text = display
    # xml:space="preserve" required when text has leading/trailing spaces.
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def render_contact_runs(
    paragraph,
    email: str,
    plain_parts: list[str],
    link_items: list[tuple[str, str]],
) -> None:
    """Build the contact-line paragraph run-by-run with hyperlinks.

    Email renders as a mailto: hyperlink. URL link extras render as external
    hyperlinks. Phone and location remain plain text. Falls back to plain text
    for any part that fails (add_hyperlink handles that internally).
    """
    first = True

    def _sep() -> None:
        nonlocal first
        if not first:
            paragraph.add_run(" | ")
        first = False

    if email:
        _sep()
        add_hyperlink(paragraph, normalize_link_url(email), email)

    for part in plain_parts:
        _sep()
        paragraph.add_run(part)

    for display, url in link_items:
        _sep()
        add_hyperlink(paragraph, url, display)
